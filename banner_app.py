#!/usr/bin/env python3
"""
Banner Tool — Hosted Web App (HF Spaces)
• Load a banner folder OR paste a Ziflow proof URL
• Preview the banner live and pick your own frame times
• Export PNG frames + full ISI (image + copyable HTML)
• Export high-quality MP4 at 238 WPM ISI scroll speed
• Multi-Banner Viewer: load every banner from a Klick staging URL at once,
  scrub them in lockstep, capture frames + bundle as a zip
"""

import asyncio
import json
import mimetypes
import os
import queue
import re
import secrets
import shutil
import subprocess
import sys
import tempfile
import threading
import time
import urllib.parse
import urllib.request
import webbrowser
import zipfile
from dataclasses import dataclass, field
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from typing import Optional

# When running as a PyInstaller bundle, point Playwright at the ms-playwright
# folder that sits next to the .app bundle in the distribution folder.
if getattr(sys, 'frozen', False):
    # sys.executable = .../Banner Frame Capture.app/Contents/MacOS/Banner Frame Capture
    _app_bundle = Path(sys.executable).parent.parent.parent   # the .app itself
    _browsers   = _app_bundle.parent / 'ms-playwright'        # sibling folder
    os.environ.setdefault('PLAYWRIGHT_BROWSERS_PATH', str(_browsers))

from playwright.async_api import async_playwright
import imageio_ffmpeg

# ═══════════════════════════════════════════════════════════════════════════════
# Per-user session state
# ═══════════════════════════════════════════════════════════════════════════════
#
# The app is hosted multi-user (Hugging Face Spaces). Each browser session
# gets its own `Session` keyed by a `bfc_sid` cookie so concurrent users don't
# clobber each other's banner / output / scratch state. A janitor thread
# (started in main()) evicts idle sessions and wipes their work_dir.
#
# Locally-developed flows (single-user, no cookies) still work because every
# request without a cookie creates a fresh session.

@dataclass
class SlotState:
    """Banner Compare (page 3) — one of the two compared banner slots."""
    dir: Optional[Path] = None
    w:   int = 160
    h:   int = 600
    remote: bool = False


@dataclass
class Session:
    id: str
    work_dir: Path                                          # tempfile.mkdtemp() per session

    # Page 1 — Frame Capture
    banner_dir:        Optional[Path] = None                # loaded banner folder (local OR temp)
    banner_w:          int = 160
    banner_h:          int = 600
    banner_is_remote:  bool = False                         # True for Ziflow / DoubleClick (temp)
    banner_studio_url: Optional[str] = None                 # original Studio preview URL (ISI capture)
    output_dir:        Optional[Path] = None                # where the capture writes outputs

    # Page 3 — Banner Compare
    p3a: SlotState = field(default_factory=SlotState)
    p3b: SlotState = field(default_factory=SlotState)

    # Page 4 — Multi-Banner Viewer
    multi_slots: dict = field(default_factory=dict)         # {slot_id: {"dir","w","h","name"}}
    multi_meta:  dict = field(default_factory=dict)         # {"source_url","user","name"}

    # Page 5 — Storyboard
    last_bundle:                Optional[Path] = None
    storyboard_pdf_path:        Optional[Path] = None
    existing_storyboard_pages:  list = field(default_factory=list)
    existing_storyboard_name:   Optional[str] = None
    replaced_storyboard_path:   Optional[Path] = None

    # Page 6 — CPFP (change-proof full-proof)
    cpfp_state: dict = field(default_factory=lambda: {
        "old": {"name": None, "pages": []},
        "new": {"name": None, "pages": []},
    })

    # Per-session SSE: progress events for THIS user only.
    sse_queue: queue.SimpleQueue = field(default_factory=queue.SimpleQueue)
    last_seen: float = field(default_factory=time.monotonic)


SESSIONS: dict[str, Session] = {}
SESSIONS_LOCK = threading.Lock()

# Cap concurrent Playwright/Chromium launches so a busy moment with many users
# doesn't OOM the 16GB Hugging Face Space. Each capture acquires this before
# spawning a browser context.
PLAYWRIGHT_SEM = threading.BoundedSemaphore(4)


def _new_session() -> Session:
    sid = secrets.token_urlsafe(16)
    work = Path(tempfile.mkdtemp(prefix="bfc-"))
    s = Session(id=sid, work_dir=work)
    with SESSIONS_LOCK:
        SESSIONS[sid] = s
    return s

# Default frame capture times (label, seconds) — user can customise in the UI
DEFAULT_FRAMES = [
    ("01_scene1_logo",       1.4),
    ("02_scene2_drug_info",  5.0),
    ("03_scene3_efficacy",   9.5),
    ("04_scene4_safety",    14.0),
    ("05_scene5_final_cta", 19.0),
]

ISI_HIDE = [
    "#clickthru", ".bg", ".circBG",
    ".circ1", ".circ2", ".circ3",
    ".t2", ".t3", ".t4",
    ".logo", ".img1", ".circImg-mask", ".img-circ",
    ".efCirc1", ".efCirc2", ".efCirc3",
    ".cta", ".disclaimer", ".footnotef2",
    "#footer",   # the trailing "Full Prescribing Information" link inside the ISI
]

# JS that detects the banner's natural width/height from inside a loaded page
# or frame. Order: <meta name="ad.size">  →  .banner element  →  body. Used by
# the ISI capture functions so the ISI panel's CSS width matches the actual
# banner size — regardless of what stale value `_banner_w` / `_banner_h` may
# have left over from a previous Frame Capture run.
DETECT_BANNER_SIZE_JS = """() => {
    const meta = document.querySelector('meta[name="ad.size"]');
    if (meta) {
        const m = (meta.content || '').match(/width=(\\d+)[\\s,]+height=(\\d+)/i);
        if (m) return { w: parseInt(m[1], 10), h: parseInt(m[2], 10) };
    }
    const banner = document.querySelector('.banner');
    if (banner) {
        const r = banner.getBoundingClientRect();
        if (r.width) return { w: Math.round(r.width), h: Math.round(r.height) };
    }
    if (document.body && document.body.offsetWidth) {
        return { w: document.body.offsetWidth,
                 h: document.body.offsetHeight };
    }
    return null;
}"""

# JS that detects the ISI panel's NATURAL rendered width — i.e. the width the
# ISI was designed to occupy on the live banner, before our DOM mutations
# stretch it to fill the banner. On a 728×90 leaderboard the ISI is typically
# a 240px side column, not 728px wide; using the banner width as the panel
# width would make the text wrap differently than on the real banner. Returns
# the largest width across `#isi`, `.isi`, `.isiBody`, `.isiScroll`, `.isiFix`.
DETECT_ISI_NATURAL_WIDTH_JS = """() => {
    const candidates = ['#isi', '.isi', '.isiBody', '.isiScroll', '.isiFix'];
    let maxW = 0;
    for (const sel of candidates) {
        const el = document.querySelector(sel);
        if (!el) continue;
        const r = el.getBoundingClientRect();
        if (r.width > maxW) maxW = r.width;
    }
    return Math.round(maxW);
}"""

def _get_banner_size(html_path: Path):
    """Parse width/height from <meta name="ad.size"> or #container CSS."""
    try:
        text = html_path.read_text(encoding="utf-8", errors="replace")
        m = re.search(r'ad\.size[^>]*?width=(\d+)[^,]*,\s*height=(\d+)', text, re.I)
        if m:
            return int(m.group(1)), int(m.group(2))
        css_path = html_path.parent / "style.css"
        if css_path.exists():
            css = css_path.read_text(encoding="utf-8", errors="replace")
            mc = re.search(r'#container\s*\{[^}]*?width:\s*(\d+)px[^}]*?height:\s*(\d+)px', css, re.S)
            if mc:
                return int(mc.group(1)), int(mc.group(2))
    except Exception:
        pass
    return 160, 600


# ── Ziflow proof URL → local folder ────────────────────────────────────────────
ZIFLOW_PROOF_RE = re.compile(r'https?://[^/]+\.ziflow\.io/proof/[A-Za-z0-9]+', re.I)
DOUBLECLICK_PREVIEW_RE = re.compile(
    r'^https?://www\.google\.com/doubleclick/studio/externalpreview/#?/?[A-Za-z0-9_-]+',
    re.I)


def _studio_url_for_creative(base_url: str, creative_id: str) -> str:
    """Rewrite the creativeId inside a DoubleClick Studio preview URL.

    Studio URLs put query params *inside the fragment* (after `#/<previewId>`),
    so urlparse drops them into .fragment, not .query. We have to split the
    fragment ourselves.
    """
    from urllib.parse import urlparse, urlunparse, parse_qsl, urlencode
    u = urlparse(base_url)
    if "?" in u.fragment:
        path, qs = u.fragment.split("?", 1)
        params = dict(parse_qsl(qs, keep_blank_values=True))
    else:
        path, params = u.fragment, {}
    params["creativeId"] = str(creative_id)
    return urlunparse(u._replace(fragment=f"{path}?{urlencode(params)}"))

# Some Ziflow edges 502 when the User-Agent looks like Python's default urllib.
_UA_HEADERS = {
    "User-Agent": ("Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
                   "AppleWebKit/537.36 (KHTML, like Gecko) "
                   "Chrome/124.0.0.0 Safari/537.36"),
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
    "Accept-Language": "en-US,en;q=0.9",
}


def _http_get(url: str, timeout: int = 15) -> bytes:
    """GET a URL with a real-browser User-Agent (Ziflow rejects bare urllib UA)."""
    req = urllib.request.Request(url, headers=_UA_HEADERS)
    with urllib.request.urlopen(req, timeout=timeout) as resp:
        return resp.read()


def _parse_multipart(body: bytes, content_type: str) -> list[dict]:
    """Parse a multipart/form-data body into a list of {name, filename,
    content_type, data} dicts using the stdlib email parser.

    Returns empty list on parse failure rather than raising — handlers
    surface a friendly error to the user.
    """
    from email.parser import BytesParser
    from email.policy import default
    if not content_type or "boundary" not in content_type.lower():
        return []
    pre = f"Content-Type: {content_type}\r\n\r\n".encode()
    try:
        msg = BytesParser(policy=default).parsebytes(pre + body)
    except Exception:
        return []
    parts = []
    for part in msg.iter_parts():
        cd = part.get("Content-Disposition", "") or ""
        name, fname = "", None
        for chunk in cd.split(";"):
            chunk = chunk.strip()
            if chunk.startswith("name="):
                name = chunk[5:].strip().strip('"')
            elif chunk.startswith("filename="):
                fname = chunk[9:].strip().strip('"')
        try:
            data = part.get_payload(decode=True) or b""
        except Exception:
            data = b""
        parts.append({
            "name": name,
            "filename": fname,
            "content_type": part.get_content_type(),
            "data": data,
        })
    return parts


def _safe_subpath(filename: str) -> Path:
    """Sanitize an upload filename for use as a relative path inside work_dir.

    Strips drive letters, leading separators, '..' segments, NULs. Keeps
    forward-slash subfolder hierarchy (used by webkitdirectory uploads).
    """
    # Normalize Windows-style separators
    s = (filename or "").replace("\\", "/")
    parts = []
    for chunk in s.split("/"):
        chunk = chunk.replace("\x00", "")
        if not chunk or chunk == "." or chunk == "..":
            continue
        parts.append(chunk)
    if not parts:
        parts = ["upload.bin"]
    return Path(*parts)


def _ziflow_url_to_local(url: str) -> Path:
    """Download a Ziflow proof's banner assets to a temp folder; return its Path.

    Ziflow proof pages embed an `og:image` meta tag pointing to a thumbnail on
    the public `proof-assets.ziflow.io` CDN. The banner's `index.html` and
    every relative file it references live alongside that thumbnail and are
    fetchable without authentication.
    """
    page = _http_get(url).decode('utf-8', 'replace')
    m = re.search(
        r'og:image"\s+content="(https://proof-assets\.ziflow\.io/Proofs/'
        r'([0-9a-f-]+)/richMedia/([0-9a-f-]+)/[^"]+)"',
        page, re.I)
    if not m:
        raise ValueError("Couldn't find banner asset in this Ziflow page.")
    proof_guid, asset_guid = m.group(2), m.group(3)
    base = f"https://proof-assets.ziflow.io/Proofs/{proof_guid}/richMedia/{asset_guid}/"

    tmp = Path(tempfile.mkdtemp(prefix="ziflow_banner_"))

    # Pull index.html + style.css first (we need them to discover other refs)
    for fname in ("index.html", "style.css"):
        try:
            (tmp / fname).write_bytes(_http_get(base + fname))
        except Exception:
            pass

    if not (tmp / "index.html").exists():
        shutil.rmtree(tmp, ignore_errors=True)
        raise ValueError("index.html not found at the Ziflow asset URL.")

    # Discover relative file references in HTML (href/src — both quote styles)
    # and CSS (url(...)). Also catch DoubleClick Studio's polite-load pattern
    # `data-src="..."` which Enabler swaps to real <img src> at runtime.
    refs = set()
    html_text = (tmp / "index.html").read_text(encoding="utf-8", errors="replace")
    refs.update(re.findall(r'(?:href|src|data-src)\s*=\s*"([^"#:?]+)"', html_text))
    refs.update(re.findall(r"(?:href|src|data-src)\s*=\s*'([^'#:?]+)'", html_text))
    if (tmp / "style.css").exists():
        refs.update(re.findall(
            r"url\(['\"]?([^'\")]+)['\"]?\)",
            (tmp / "style.css").read_text(encoding="utf-8", errors="replace")))

    for ref in refs:
        if ref.startswith(("http://", "https://", "data:", "//")):
            continue
        ref = ref.split("?")[0].split("#")[0]
        if not ref or ref in ("index.html", "style.css"):
            continue
        try:
            data = _http_get(base + ref)
            out = tmp / ref
            out.parent.mkdir(parents=True, exist_ok=True)
            out.write_bytes(data)
        except Exception:
            pass  # 404s on optional/fallback files are harmless

    return tmp


def _parse_ziflow_asset_urls(text: str) -> list:
    """Detect a paste of one or more Ziflow asset CDN URLs and parse them.

    Returns [{name, url, creative_id}, ...] when EVERY non-empty line
    (newline- or comma-separated) matches the public asset CDN pattern
    `https://proof-assets.ziflow.io/Proofs/<proof_guid>/richMedia/<asset_guid>/...`.
    Returns [] otherwise — the caller falls back to other routing.

    Single-URL paste is supported: the bookmarklet sometimes only collects
    one banner (e.g. on a 1-banner proof, or when the click-cycle bails),
    and we still want that path to work without falling through to the
    auth-gate-detecting Playwright route. Distinct enough from a proof URL
    (klickhealth.ziflow.io/proof/<id>) that the all-or-nothing match keeps
    routing safe.
    """
    lines = [l.strip() for l in text.replace(",", "\n").splitlines() if l.strip()]
    if not lines:
        return []
    pattern = re.compile(
        r"^https://proof-assets\.ziflow\.io/Proofs/([0-9a-f-]+)/richMedia/([0-9a-f-]+)/",
        re.I)
    out = []
    seen_guids = set()
    for line in lines:
        m = pattern.match(line)
        if not m:
            return []   # all-or-nothing — partial match falls through
        proof_guid, asset_guid = m.group(1), m.group(2)
        if asset_guid in seen_guids:
            continue
        seen_guids.add(asset_guid)
        out.append({
            "name":        f"banner_{asset_guid[:8]}",
            "url":         (f"https://proof-assets.ziflow.io/Proofs/{proof_guid}"
                            f"/richMedia/{asset_guid}/index.html"),
            "creative_id": asset_guid,
        })
    return out


def _discover_ziflow_banners(proof_url: str) -> list:
    """Walk a Ziflow proof page and return one entry per banner found.

    Ziflow stores each rich-media banner as a folder on `proof-assets.ziflow.io`
    keyed by an asset GUID. The single-banner path (`_ziflow_url_to_local`)
    extracts only the og:image meta tag, which reflects the page's "primary"
    banner. For multi-banner proofs (e.g. all sizes of the same campaign on
    one proof page), every banner's asset GUID appears somewhere in the page
    HTML — typically in iframe src URLs, thumbnail references, or embedded
    JSON. We collect every unique `richMedia/<asset_guid>/...` URL we see.

    Returns a list of {name, url, creative_id} dicts shaped like the other
    discovery helpers (Klick, DoubleClick) so `_multi_load` can treat them
    uniformly. Ordering preserves first-appearance in the HTML.
    """
    page = _http_get(proof_url).decode("utf-8", "replace")

    # Match every richMedia URL on the page. Capture proof_guid + asset_guid +
    # filename so we can build the index.html URL and dedupe by asset_guid.
    pattern = re.compile(
        r'https://proof-assets\.ziflow\.io/Proofs/'
        r'([0-9a-f-]+)/richMedia/([0-9a-f-]+)/([^"\'\s\?#)]+)',
        re.I,
    )
    seen = []   # preserve order; track asset_guids we've already added
    seen_guids = set()
    for proof_guid, asset_guid, filename in pattern.findall(page):
        if asset_guid in seen_guids:
            continue
        seen_guids.add(asset_guid)
        index_url = (f"https://proof-assets.ziflow.io/Proofs/{proof_guid}/"
                     f"richMedia/{asset_guid}/index.html")
        seen.append({
            "name":        f"banner_{asset_guid[:8]}",
            "url":         index_url,
            "creative_id": asset_guid,
        })
    return seen


def _doubleclick_url_to_local(url: str, picked_index: int = 0,
                              cached_banners: list = None):
    """Open a DoubleClick Studio external-preview URL, pick a creative from
    the size dropdown (the first by default, or `picked_index`-th), download
    its files to a temp folder. Returns (folder_Path, info_dict).

    info_dict has:
      - `picked_name`   — the chosen creative's display name (e.g. "160x600")
      - `picked_index`  — its index in the size dropdown
      - `total_count`   — total creatives on the page
      - `all_sizes`     — list of every creative's display name
      - `creatives`     — full discovery list (name + creative_id + url + size).
                          Sent back so the frontend can populate a banner-
                          size dropdown without re-discovering.
      - `studio_url`    — Studio preview URL pinned to the chosen creative
                          (used by capture_isi_png_studio to source the live
                          ISI panel — the bare s0.2mdn.net creative loaded
                          standalone doesn't render its ISI).

    Pass `cached_banners` (the list returned by an earlier call's
    `info["creatives"]`) to skip re-discovery — useful for the
    pick-different-creative flow where the user clicks the dropdown and we
    just need to download a different one.
    """
    import asyncio
    banners = cached_banners or asyncio.run(_discover_doubleclick_banners(url))
    if not banners:
        raise ValueError("No creatives found on this DoubleClick page.")
    if not (0 <= picked_index < len(banners)):
        picked_index = 0
    picked = banners[picked_index]
    folder = _download_banner_authed(picked["url"], "", "")
    info = {
        "picked_name":  picked["name"],
        "picked_index": picked_index,
        "total_count":  len(banners),
        "all_sizes":    [b["name"] for b in banners],
        "creatives":    banners,    # full list for client-side dropdown
        "studio_url":   _studio_url_for_creative(url, picked["creative_id"]),
    }
    return folder, info


# ── Klick staging URL → N local banner folders (Multi-Banner Viewer) ──────────

def _http_basic_get(url: str, user: str = "", password: str = "", timeout: int = 20) -> bytes:
    """GET a URL. Tries plain (no auth) first; if it 401s and credentials were
    given, retries with HTTP Basic Auth. Most AdPiler-hosted banners are on a
    public S3 bucket so no auth is needed for the asset CDN.
    """
    req = urllib.request.Request(url, headers=_UA_HEADERS)
    try:
        with urllib.request.urlopen(req, timeout=timeout) as resp:
            return resp.read()
    except urllib.error.HTTPError as e:
        if e.code != 401 or not (user or password):
            raise
        mgr = urllib.request.HTTPPasswordMgrWithDefaultRealm()
        mgr.add_password(None, url, user or "", password or "")
        handler = urllib.request.HTTPBasicAuthHandler(mgr)
        opener  = urllib.request.build_opener(handler)
        with opener.open(req, timeout=timeout) as resp:
            return resp.read()


def _name_from_url(url: str) -> str:
    """Derive a friendly slot name from a banner URL.

    Looks for an `<digits>x<digits>` substring (e.g. '160x600') first,
    falls back to the last non-empty path segment.
    """
    m = re.search(r'(\d{2,4}x\d{2,4})', url)
    if m:
        return m.group(1)
    # Last meaningful path segment, stripping trailing index.html
    parts = [p for p in url.split("?")[0].rstrip("/").split("/") if p and p != "index.html"]
    return parts[-1] if parts else "banner"


async def _discover_klick_banners(staging_url: str, user: str, password: str):
    """Open the staging page (handling whichever auth method it uses) and
    return a list of {name, url} entries for every banner found.

    Auth strategies handled:
    1. HTTP Basic Auth (passed as Playwright's http_credentials when both
       user and password are supplied).
    2. Form-based password gate (e.g. Klick → AdPiler) — detects an
       <input type="password"> element on the loaded page, dismisses any
       cookie-banner overlay that might intercept clicks, fills it, and
       submits.

    Banner-discovery strategies:
    1. <iframe src="..."> elements (the typical pattern).
    2. <a href> pointing at index.html or a banner-shaped URL.
    """
    from playwright.async_api import async_playwright
    async with async_playwright() as pw:
        browser = await pw.chromium.launch(headless=True)
        try:
            ctx_kwargs = {"user_agent": _UA_HEADERS["User-Agent"]}
            if user:
                ctx_kwargs["http_credentials"] = {"username": user, "password": password}
            ctx  = await browser.new_context(**ctx_kwargs)
            page = await ctx.new_page()

            await page.goto(staging_url, wait_until="domcontentloaded", timeout=45_000)
            await page.wait_for_timeout(1500)

            # Detect a form-based password gate; submit if present.
            has_pw_form = await page.evaluate(
                "() => !!document.querySelector('input[type=password]')"
            )
            if has_pw_form and password:
                # Strip cookie banners / overlays that intercept pointer events.
                await page.evaluate("""() => {
                    const sels = ['#cookiescript_injected_wrapper','#cookiescript_injected',
                                  '.cookiescript', '#cookieconsent', '.cookieconsent',
                                  '#cookies-banner', '.cookies-banner'];
                    sels.forEach(sel => document.querySelectorAll(sel).forEach(el => el.remove()));
                }""")
                await page.wait_for_timeout(150)
                try:
                    await page.fill('input[type="password"]', password)
                    await page.click(
                        'button[type="submit"], input[type="submit"]',
                        timeout=8_000,
                    )
                    await page.wait_for_load_state("networkidle", timeout=30_000)
                    await page.wait_for_timeout(2000)
                except Exception:
                    # If the form submit fails (overlay still there, no submit button),
                    # try pressing Enter as a last resort.
                    try:
                        await page.press('input[type="password"]', "Enter")
                        await page.wait_for_load_state("networkidle", timeout=30_000)
                        await page.wait_for_timeout(2000)
                    except Exception:
                        pass

            srcs = await page.evaluate("""() => {
                const iframes = [...document.querySelectorAll('iframe[src]')]
                    .map(f => f.src)
                    .filter(s => s && !s.startsWith('about:') && !s.includes('youtube'));
                if (iframes.length) return iframes;
                const anchors = [...document.querySelectorAll('a[href]')]
                    .map(a => a.href)
                    .filter(h => /index\\.html|\\d+x\\d+/i.test(h));
                return anchors;
            }""")

            # Sanity check — if we never got past the password page, surface a clear error.
            still_locked = await page.evaluate(
                "() => !!document.querySelector('input[type=password]')"
            )
        finally:
            await browser.close()

    if not srcs and 'still_locked' in locals() and still_locked:
        raise ValueError(
            "Couldn't sign in — the password page is still showing. "
            "Check the password and try again."
        )

    banners, seen, name_counts = [], set(), {}
    for s in srcs:
        if s in seen:
            continue
        seen.add(s)
        if not s.endswith(".html"):
            s = s.rstrip("/") + "/index.html"
        base_name = _name_from_url(s)
        # De-dup banner names (e.g. when the staging page lists 3 different
        # 160x600 variants — append a counter so each gets a unique slot label).
        name_counts[base_name] = name_counts.get(base_name, 0) + 1
        if name_counts[base_name] > 1:
            name = f"{base_name}_{name_counts[base_name]}"
        else:
            name = base_name
        banners.append({"name": name, "url": s})
    return banners


# ── Ziflow Multi-Banner discovery (Klick proofs are SPAs) ────────────────────
#
# Klick's Ziflow proof pages (klickhealth.ziflow.io) ship a 2.6 KB SPA shell
# whose only static richMedia reference is the og:image (the FIRST banner's
# thumbnail). A multi-banner proof shows "1 of M" in the toolbar and
# lazy-loads each banner's index.html into an iframe as the user clicks the
# "Next file" chevron. To enumerate every banner, we drive the SPA the same
# way: render with Playwright, read the current iframe src, click next,
# wait for src to change, repeat.
#
# Selectors verified live on klickhealth.ziflow.io 2026-05-05:
#   - iframe holding the current banner: iframe[name^="ziflow-iframe-"]
#   - "Next file" chevron:               a[data-selector="asset-switcher-next"]
#   - "Prev file" chevron:               a[data-selector="asset-switcher-prev"]
#   - When at first/last banner the chevron has class "disabled".
ZIFLOW_BANNER_IFRAME_SELECTOR = 'iframe[name^="ziflow-iframe-"]'
ZIFLOW_NEXT_SELECTOR          = 'a[data-selector="asset-switcher-next"]'


class ZiflowAuthGateError(Exception):
    """Raised when a Ziflow proof URL hits the 'Verify your identity' gate.

    Klick's proofs at klickhealth.ziflow.io require a logged-in Ziflow
    session to view. Anonymous Playwright sessions get redirected to
    `/public-token` and shown an email-verification form — magic-link based,
    so we can't bypass it server-side. The Multi-Banner Viewer surfaces
    this as a friendly error pointing the user at the new Load-folder flow.
    """


def _parse_cookie_header(cookie_header: str, host: str) -> list:
    """Parse a 'name=value; name=value' Cookie header into Playwright
    cookie dicts. Returns a list of cookies attached to both the host
    itself and `.ziflow.io` so it works regardless of which subdomain
    set them.

    Robust to common paste shapes:
    - Raw `name=value; name2=value2`
    - A `Cookie: name=value;...` line copied from DevTools (we strip the
      `Cookie:` prefix and any wrapping quotes/whitespace)
    """
    if not cookie_header:
        return []
    s = cookie_header.strip()
    # Strip a leading "Cookie:" if pasted from a Request Headers panel.
    if s.lower().startswith("cookie:"):
        s = s[7:].strip()
    # Strip wrapping quotes (curl --cookie 'foo=bar' style).
    if (s.startswith("'") and s.endswith("'")) or (s.startswith('"') and s.endswith('"')):
        s = s[1:-1]
    cookies = []
    for chunk in s.split(";"):
        chunk = chunk.strip()
        if not chunk or "=" not in chunk:
            continue
        name, value = chunk.split("=", 1)
        name, value = name.strip(), value.strip()
        if not name:
            continue
        # Add for the specific host *and* the .ziflow.io parent so cookies
        # set on klickhealth.ziflow.io reach proof-assets.ziflow.io too.
        for domain in {host, ".ziflow.io"}:
            if not domain:
                continue
            cookies.append({"name": name, "value": value,
                            "domain": domain, "path": "/"})
    return cookies


async def _discover_ziflow_banners_async(proof_url: str, cookie_header: str = "") -> list:
    """Drive a Ziflow proof SPA with Playwright to discover every banner.

    Returns a list of {name, url, creative_id} dicts shaped like the other
    discovery helpers so `_multi_load` can treat all sources uniformly.

    Raises ZiflowAuthGateError if the proof requires login (Ziflow redirects
    anonymous viewers to a `/public-token` identity-verification page) AND
    no `cookie_header` was provided (or the cookies didn't bypass the gate).

    Args:
        proof_url:     the Ziflow proof page URL.
        cookie_header: optional `name=value; name=value` Cookie header
                       string. When provided, cookies are attached to the
                       Playwright context before navigation, so users with
                       a logged-in Ziflow session can paste their cookies
                       to access private (Klick) proofs.

    Algorithm:
    1. (Optional) Set cookies on the Playwright context.
    2. Open the proof URL.
    3. Detect the public-token / "Verify your identity" auth gate. If
       present, raise ZiflowAuthGateError.
    4. Wait for the proof viewer's iframe to mount (its src is the current
       banner's index.html on proof-assets.ziflow.io).
    5. Read the "1 of M" counter to get the total count.
    6. Loop up to M times: capture iframe src → extract proof+asset guid →
       click "Next" → wait for iframe src to change.
    7. Return the collected list.
    """
    from playwright.async_api import async_playwright
    import urllib.parse

    host = urllib.parse.urlparse(proof_url).hostname or ""

    async with async_playwright() as pw:
        browser = await pw.chromium.launch(headless=True)
        try:
            ctx  = await browser.new_context(user_agent=_UA_HEADERS["User-Agent"])
            cookies = _parse_cookie_header(cookie_header, host)
            if cookies:
                try:
                    await ctx.add_cookies(cookies)
                except Exception:
                    # Bad cookie shape shouldn't kill the whole flow —
                    # let the auth-gate detector surface a clean error
                    # downstream if the cookies don't unlock the proof.
                    pass
            page = await ctx.new_page()
            await page.goto(proof_url, wait_until="domcontentloaded", timeout=45_000)
            await page.wait_for_timeout(2_000)  # let SPA bootstrap

            # Detect Ziflow's identity-verification gate (Klick proofs land here).
            current_url = page.url
            body_text = await page.evaluate("() => document.body && document.body.innerText || ''")
            if (current_url.endswith("/public-token")
                    or "verify your identity" in body_text.lower()
                    or "already a user? sign in" in body_text.lower()):
                raise ZiflowAuthGateError(
                    "This Ziflow proof requires login (Klick's proofs are "
                    "private). Use the '📌 Banner Tool: Get from Ziflow' "
                    "bookmarklet above — it runs in your logged-in browser "
                    "tab and sends the banner URLs back here. Or, manually: "
                    "paste each banner's iframe-src URL "
                    "(proof-assets.ziflow.io/.../richMedia/<guid>/index.html) "
                    "newline-separated into the Staging URL field; or "
                    "download the bundle from Ziflow and use "
                    "'Storyboard → Load folder…'.")

            # Wait for the proof viewer to mount its outer iframe.
            await page.wait_for_selector(
                ZIFLOW_BANNER_IFRAME_SELECTOR, state="attached", timeout=30_000)

            # ── Mode A: "merged rich media" grid layout ──────────────────────
            # Some Ziflow proofs (multi-file proofs with rich-media banners)
            # render a SINGLE outer iframe pointing at
            # /api/proof/token/<id>/merged-rich-media — which itself contains
            # one nested <iframe src=…/richMedia/<asset_guid>/…> per banner.
            # Detect this layout up front. If present, read every nested
            # iframe src in one pass — no clicking needed.
            outer_src = await page.evaluate(
                f"() => {{ const f = document.querySelector('{ZIFLOW_BANNER_IFRAME_SELECTOR}'); "
                f"return f && f.src ? f.src : ''; }}") or ""
            if "merged-rich-media" in outer_src:
                # Wait for the inner doc + nested iframes to load.
                merged_urls = []
                for _ in range(40):
                    merged_urls = await page.evaluate(
                        "() => { const f = document.querySelector('"
                        + ZIFLOW_BANNER_IFRAME_SELECTOR + "'); "
                        "if (!f) return []; "
                        "let inner = null; try { inner = f.contentDocument; } catch (e) {} "
                        "if (!inner) return []; "
                        "const srcs = [...inner.querySelectorAll('iframe')]"
                        ".map(i => i.src || '').filter(s => s && s.includes('richMedia')); "
                        "return [...new Set(srcs)]; }") or []
                    if merged_urls:
                        break
                    await page.wait_for_timeout(300)
                src_re = re.compile(
                    r"https://proof-assets\.ziflow\.io/Proofs/"
                    r"([0-9a-f-]+)/richMedia/([0-9a-f-]+)/", re.I)
                results = []
                seen_guids = set()
                for u in merged_urls:
                    m = src_re.search(u)
                    if not m:
                        continue
                    proof_guid, asset_guid = m.group(1), m.group(2)
                    if asset_guid in seen_guids:
                        continue
                    seen_guids.add(asset_guid)
                    results.append({
                        "name":        f"banner_{asset_guid[:8]}",
                        "url":         (f"https://proof-assets.ziflow.io/Proofs/{proof_guid}"
                                        f"/richMedia/{asset_guid}/index.html"),
                        "creative_id": asset_guid,
                    })
                if results:
                    return results
                # fall through to Mode B if merged extraction came up empty
            # Give the iframe a moment to set its src to the asset CDN URL
            # (Angular sets it after the wrapper appears).
            for _ in range(20):
                first_src = await page.evaluate(
                    f"() => {{ const f = document.querySelector('{ZIFLOW_BANNER_IFRAME_SELECTOR}'); "
                    f"return f && f.src ? f.src : null; }}")
                if first_src and "richMedia" in first_src:
                    break
                await page.wait_for_timeout(500)

            # Total count from "N of M" in the toolbar; defaults to 1 if no
            # counter (single-banner proofs).
            total = await page.evaluate(
                "() => { const m = document.body.innerText.match(/(\\d+)\\s*of\\s*(\\d+)/i); "
                "return m ? parseInt(m[2], 10) : 1; }"
            ) or 1

            # Iframe src → (proof_guid, asset_guid).
            src_re = re.compile(
                r"https://proof-assets\.ziflow\.io/Proofs/"
                r"([0-9a-f-]+)/richMedia/([0-9a-f-]+)/", re.I)

            results = []
            seen_guids = set()
            for step in range(max(1, total)):
                src = await page.evaluate(
                    f"() => {{ const f = document.querySelector('{ZIFLOW_BANNER_IFRAME_SELECTOR}'); "
                    f"return f && f.src ? f.src : ''; }}") or ""
                m = src_re.search(src)
                if m:
                    proof_guid, asset_guid = m.group(1), m.group(2)
                    if asset_guid not in seen_guids:
                        seen_guids.add(asset_guid)
                        results.append({
                            "name":        f"banner_{asset_guid[:8]}",
                            "url":         (f"https://proof-assets.ziflow.io/Proofs/{proof_guid}"
                                            f"/richMedia/{asset_guid}/index.html"),
                            "creative_id": asset_guid,
                        })
                if len(results) >= total:
                    break

                # Stop if we're on the last banner (Next is "disabled").
                next_state = await page.evaluate(
                    f"() => {{ const a = document.querySelector('{ZIFLOW_NEXT_SELECTOR}'); "
                    "if (!a) return 'missing'; "
                    "return (a.className || '').toLowerCase().includes('disabled') ? 'disabled' : 'enabled'; }")
                if next_state != "enabled":
                    break

                # Click Next and wait for the iframe src to swap.
                prev_src = src
                try:
                    await page.click(ZIFLOW_NEXT_SELECTOR, timeout=4_000)
                except Exception:
                    break
                try:
                    await page.wait_for_function(
                        f"prev => {{ const f = document.querySelector('{ZIFLOW_BANNER_IFRAME_SELECTOR}'); "
                        "return f && f.src && f.src !== prev; }",
                        arg=prev_src, timeout=10_000)
                except Exception:
                    # Try one more wait pass — sometimes Angular swaps the
                    # iframe element entirely (different name suffix) so the
                    # outer page-level src diff doesn't fire on the same
                    # element. Re-query.
                    await page.wait_for_timeout(800)

            return results
        finally:
            await browser.close()


async def _discover_doubleclick_banners(preview_url: str):
    """Render a Google DoubleClick Studio external-preview page, click through
    every creative in the dropdown, and return the underlying banner asset
    URLs (publicly hosted on s0.2mdn.net).

    Returns: [{"name": "<W>x<H>", "url": "<asset-index.html-url>"}, ...]
    """
    from playwright.async_api import async_playwright
    async with async_playwright() as pw:
        browser = await pw.chromium.launch(headless=True)
        try:
            ctx = await browser.new_context(
                viewport={"width": 1400, "height": 900},
                user_agent=_UA_HEADERS["User-Agent"],
            )
            page = await ctx.new_page()

            # Capture every banner index.html that loads (the asset CDN is
            # s0.2mdn.net; URLs follow /pv2/<creativeId>/<timestamp>/index.html).
            observed = {}   # creative_id → asset_url (in load order)
            ordered  = []   # ordered list of creative_ids as they appear
            def _on_response(r):
                u = r.url
                if 's0.2mdn.net' in u and '/pv2/' in u and 'index.html' in u:
                    m = re.search(r'/pv2/(\d+)/', u)
                    if m and m.group(1) not in observed:
                        observed[m.group(1)] = u
                        ordered.append(m.group(1))
            page.on("response", _on_response)

            await page.goto(preview_url, wait_until="domcontentloaded", timeout=45_000)
            try:
                await page.wait_for_selector(
                    'iframe[src*="doubleclick/preview/main"]', timeout=30_000)
            except Exception:
                pass
            await page.wait_for_timeout(3500)

            # Open the size-selector dropdown so we can read all menu items.
            try:
                btn = page.locator('div.goog-flat-menu-button').first
                await btn.wait_for(state="visible", timeout=10_000)
                bb = await btn.bounding_box()
                await page.mouse.click(bb["x"] + bb["width"]/2,
                                       bb["y"] + bb["height"]/2)
                await page.wait_for_timeout(800)

                items = await page.evaluate("""() => {
                    return [...document.querySelectorAll('.goog-menuitem')]
                        .filter(el => el.offsetWidth > 0)
                        .map(el => {
                            const name_el = el.querySelector('.select-name');
                            const rect = el.getBoundingClientRect();
                            return {
                                name: (name_el ? name_el.innerText : '').trim(),
                                cx: rect.x + rect.width/2,
                                cy: rect.y + rect.height/2,
                            };
                        })
                        .filter(it => it.name);
                }""")
            except Exception:
                items = []

            # Read the initial creativeId from the URL fragment — Studio
            # updates the URL hash to reflect the currently-loaded creative.
            # We use this to pair menu-item names with creativeIds robustly
            # (instead of zip-pairing observed-load-order with menu-item-order,
            # which breaks if the initial URL targets a non-first dropdown
            # item or if any click silently fails).
            item_name_to_cid = {}
            async def _current_cid_from_hash():
                try:
                    return await page.evaluate(
                        "() => { const h = location.hash || ''; "
                        "const m = h.match(/[?&]creativeId=([^&]+)/); "
                        "return m ? m[1] : null; }")
                except Exception:
                    return None

            initial_cid = await _current_cid_from_hash()

            # Click each menu item in turn so its banner asset loads. Re-open
            # the dropdown between clicks (it auto-closes on selection). For
            # each click, ADAPTIVELY poll for the new s0.2mdn response — a
            # heavy creative (e.g. 728×90 with lots of assets) can take longer
            # than a small one, and a fixed 2500ms wait would silently miss it.
            for i, it in enumerate(items):
                n_before = len(observed)
                # Re-query the i-th item by name after each iteration so we
                # have a fresh, on-screen element handle (the dropdown list
                # may reflow between openings). Scroll it into view first —
                # if 728×90 sits below the visible portion of a scrollable
                # menu, mouse coords from the initial enumeration miss it.
                target_name = it["name"]
                cx, cy = it["cx"], it["cy"]
                try:
                    handle = await page.evaluate_handle(
                        f"""(name) => {{
                            const items = [...document.querySelectorAll('.goog-menuitem')];
                            return items.find(el => {{
                                const n = el.querySelector('.select-name');
                                return n && n.innerText.trim() === name;
                            }}) || null;
                        }}""", target_name)
                    el = handle.as_element() if handle else None
                    if el:
                        await el.scroll_into_view_if_needed()
                        bb = await el.bounding_box()
                        if bb:
                            cx = bb["x"] + bb["width"] / 2
                            cy = bb["y"] + bb["height"] / 2
                except Exception:
                    pass
                await page.mouse.click(cx, cy)

                # Adaptive wait: poll up to 12s for a new s0.2mdn response.
                # A new creativeId in `observed` means we caught the asset.
                deadline_ms, elapsed_ms = 12_000, 0
                while elapsed_ms < deadline_ms and len(observed) <= n_before:
                    await page.wait_for_timeout(250)
                    elapsed_ms += 250

                # Pair this menu item's name with whichever creativeId is now
                # active in the URL hash — robust against click ordering and
                # against the initial URL targeting a non-first dropdown item.
                cur = await _current_cid_from_hash()
                if cur:
                    item_name_to_cid[target_name] = cur

                if i < len(items) - 1:
                    try:
                        bb = await page.locator('div.goog-flat-menu-button').first.bounding_box()
                        await page.mouse.click(bb["x"] + bb["width"]/2,
                                               bb["y"] + bb["height"]/2)
                        await page.wait_for_timeout(800)
                    except Exception:
                        break

            # If the initial URL's creativeId didn't get paired by any click
            # (e.g., clicking the already-selected item didn't update hash),
            # pair it with the menu item whose name embeds the same size.
            if initial_cid and initial_cid not in item_name_to_cid.values():
                # First menu item that doesn't have an assigned cid yet wins.
                for it in items:
                    if it["name"] not in item_name_to_cid:
                        item_name_to_cid[it["name"]] = initial_cid
                        break
        finally:
            await browser.close()

    # Pair menu-item names with creativeIds. Prefer `item_name_to_cid` (built
    # from the URL hash after each click — robust) over zip-pairing by load
    # order, which fails when click order doesn't match response order or
    # when the initial URL targets a non-first dropdown item.
    out = []
    used_cids = set()
    for it in items:
        size_m = re.search(r'(\d{2,4}x\d{2,4})', it["name"] or "")
        size_name = size_m.group(1) if size_m else (it["name"] or "")
        cid = item_name_to_cid.get(it["name"])
        if cid and cid in observed:
            out.append({"name": size_name, "url": observed[cid], "creative_id": cid})
            used_cids.add(cid)

    # Append any observed cids we didn't pair via the hash mechanism — gives
    # us a result even if the URL-hash trick didn't work for some reason.
    for cid in ordered:
        if cid in used_cids:
            continue
        out.append({"name": f"creative_{cid}",
                    "url": observed[cid], "creative_id": cid})
        used_cids.add(cid)

    if not out and observed:
        # Fallback: at least one asset loaded but we couldn't pair names
        for cid, u in observed.items():
            out.append({"name": f"creative_{cid}", "url": u, "creative_id": cid})
    return out


def _download_banner_authed(index_url: str, user: str, password: str) -> Path:
    """Download a single banner's index.html plus all relative refs to a temp folder."""
    base = index_url.rsplit("/", 1)[0] + "/"
    tmp  = Path(tempfile.mkdtemp(prefix="klick_banner_"))

    (tmp / "index.html").write_bytes(_http_basic_get(index_url, user, password))
    try:
        (tmp / "style.css").write_bytes(_http_basic_get(base + "style.css", user, password))
    except Exception:
        pass

    # Discover relative refs in HTML (href/src — both quote styles, plus
    # `data-src` used by DoubleClick Studio's Enabler polite-load pattern)
    # and CSS (url(...)). Single-quote attribute values (`src='sticker.png'`)
    # used to be missed entirely — that's how a banner could end up with a
    # broken image of its main artwork tile.
    refs = set()
    html_text = (tmp / "index.html").read_text(encoding="utf-8", errors="replace")
    refs.update(re.findall(r'(?:href|src|data-src)\s*=\s*"([^"#:?]+)"', html_text))
    refs.update(re.findall(r"(?:href|src|data-src)\s*=\s*'([^'#:?]+)'", html_text))
    if (tmp / "style.css").exists():
        refs.update(re.findall(r"url\(['\"]?([^'\")]+)['\"]?\)",
                               (tmp / "style.css").read_text(encoding="utf-8", errors="replace")))

    for ref in refs:
        if ref.startswith(("http://", "https://", "data:", "//")):
            continue
        ref = ref.split("?")[0].split("#")[0]
        if not ref or ref in ("index.html", "style.css"):
            continue
        try:
            data = _http_basic_get(base + ref, user, password)
            out  = tmp / ref
            out.parent.mkdir(parents=True, exist_ok=True)
            out.write_bytes(data)
        except Exception:
            pass  # 404s on optional files are harmless
    return tmp


def _js_hide(selectors):
    arr = "[" + ",".join(f'"{s}"' for s in selectors) + "]"
    return f"{arr}.forEach(s=>document.querySelectorAll(s).forEach(e=>e.style.display='none'));"


# ── Document text extraction (for ISI Compare page) ───────────────────────────

def _extract_text_from_pdf(path: Path) -> str:
    import pypdf
    text_parts = []
    with open(path, "rb") as f:
        reader = pypdf.PdfReader(f)
        for page in reader.pages:
            text_parts.append(page.extract_text() or "")
    return "\n".join(text_parts)


def _extract_pdf_pages_text(path: Path):
    """Return one string per PDF page (parallel to _extract_pdf_pages_to_pngs).
    Pages that can't be text-extracted (e.g. fully scanned/raster) yield ''.
    """
    import pypdf
    out = []
    try:
        with open(path, "rb") as f:
            for page in pypdf.PdfReader(f).pages:
                try:
                    out.append(page.extract_text() or "")
                except Exception:
                    out.append("")
    except Exception:
        return []
    return out


def _extract_text_from_docx(path: Path) -> str:
    from docx import Document
    doc = Document(str(path))
    parts = []
    for para in doc.paragraphs:
        parts.append(para.text)
    # Pull text from any tables too
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _extract_text_from_doc(path: Path) -> str:
    """Dispatch to the right extractor based on suffix."""
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _extract_text_from_pdf(path)
    if suffix in (".docx", ".doc"):
        return _extract_text_from_docx(path)
    raise ValueError(f"Unsupported document type: {suffix}")


# ── Word-level text diff (used by both ISI Compare and Banner Compare) ────────

def _normalize_text(s: str) -> str:
    """Collapse whitespace, drop empty lines — what we feed into the diff."""
    lines = [re.sub(r'\s+', ' ', ln).strip() for ln in (s or "").splitlines()]
    return "\n".join(ln for ln in lines if ln)


def _word_diff_html(a: str, b: str) -> dict:
    """Return word-level diff between two texts as left/right HTML segments,
    plus a count of differences. Same words are left plain; words only in `a`
    are wrapped in <del>; words only in `b` are wrapped in <ins>."""
    import difflib, html
    aw = re.findall(r'\S+|\n', a)   # tokens including newlines
    bw = re.findall(r'\S+|\n', b)
    sm = difflib.SequenceMatcher(a=aw, b=bw, autojunk=False)
    left, right = [], []
    diff_count = 0
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        a_chunk = aw[i1:i2]
        b_chunk = bw[j1:j2]
        if tag == "equal":
            left.append(_render_tokens(a_chunk))
            right.append(_render_tokens(b_chunk))
        elif tag == "delete":
            left.append(f'<del>{_render_tokens(a_chunk)}</del>')
            diff_count += sum(1 for t in a_chunk if t != "\n")
        elif tag == "insert":
            right.append(f'<ins>{_render_tokens(b_chunk)}</ins>')
            diff_count += sum(1 for t in b_chunk if t != "\n")
        elif tag == "replace":
            left.append(f'<del>{_render_tokens(a_chunk)}</del>')
            right.append(f'<ins>{_render_tokens(b_chunk)}</ins>')
            diff_count += max(
                sum(1 for t in a_chunk if t != "\n"),
                sum(1 for t in b_chunk if t != "\n"),
            )
    return {
        "left":  "".join(left),
        "right": "".join(right),
        "diffs": diff_count,
        "identical": diff_count == 0,
    }


def _render_tokens(tokens):
    import html
    out = []
    for t in tokens:
        if t == "\n":
            out.append("<br>")
        else:
            out.append(html.escape(t) + " ")
    return "".join(out)


# ── Visual diff (Banner Compare page) ─────────────────────────────────────────

def _visual_diff_image(img_a_bytes: bytes, img_b_bytes: bytes) -> bytes:
    """Return PNG bytes showing pixel-level differences between two screenshots.

    Areas where pixels match → faded grey. Areas where they differ → vivid red
    overlay. Different-sized images are aligned to the larger bounding box.
    """
    import io
    from PIL import Image, ImageChops, ImageEnhance
    a = Image.open(io.BytesIO(img_a_bytes)).convert("RGB")
    b = Image.open(io.BytesIO(img_b_bytes)).convert("RGB")

    # Pad both to the same size (max dims), white background
    w = max(a.width,  b.width)
    h = max(a.height, b.height)
    pad_a = Image.new("RGB", (w, h), "white"); pad_a.paste(a, (0, 0))
    pad_b = Image.new("RGB", (w, h), "white"); pad_b.paste(b, (0, 0))

    diff = ImageChops.difference(pad_a, pad_b).convert("L")
    # Boost the diff so even small changes are visible
    diff = ImageEnhance.Contrast(diff).enhance(2.0)

    # Use banner A as the base, faded; overlay red where pixels differ.
    base = ImageEnhance.Brightness(pad_a).enhance(1.25).convert("RGBA")
    red  = Image.new("RGBA", (w, h), (255, 50, 60, 0))
    red.putalpha(diff)            # alpha = strength of difference
    overlay = Image.alpha_composite(base, red)

    out = io.BytesIO()
    overlay.convert("RGB").save(out, format="PNG")
    return out.getvalue()


# ── Storyboard PDF generator (Page 5) ─────────────────────────────────────────

def _scan_capture_bundle(folder: Path) -> list:
    """Scan a multi-banner capture folder. Returns a list of banner entries
    each with their frame PNGs (sorted) and optional ISI image."""
    out = []
    if not folder or not folder.exists():
        return out
    for sub in sorted(folder.iterdir()):
        if not sub.is_dir():
            continue
        # PNGs in the banner subfolder, excluding the ISI exports
        all_pngs = sorted(sub.glob("*.png"))
        frames   = [p for p in all_pngs if p.name not in ("isi_full.png",)]
        isi_png  = sub / "isi_full.png" if (sub / "isi_full.png").exists() else None
        if not frames and not isi_png:
            continue
        out.append({
            "name":    sub.name,
            "frames":  [str(p) for p in frames],
            "isi_png": str(isi_png) if isi_png else None,
        })
    return out


# Brand colour from the example storyboard (the orange-red dot in the header).
_STORYBOARD_BRAND = "#E45F38"


def _storyboard_page_html(banner: dict, header: str, footer: str,
                          notes: list, page_w: int, page_h: int) -> str:
    """Build the HTML for one storyboard page (one banner). Layout adapts
    based on banner aspect ratio: tall banners get 4 columns, wide banners
    (e.g. 728×90) get 4 stacked rows."""
    import base64
    frame_paths = banner.get("frames", [])[:4]

    # Inline frame images as data URIs so the rendering iframe doesn't need a
    # local file server.
    frame_imgs = []
    for p in frame_paths:
        try:
            data = Path(p).read_bytes()
            frame_imgs.append("data:image/png;base64," + base64.b64encode(data).decode())
        except Exception:
            frame_imgs.append("")

    # Pad notes to length 4 to match the four-frame layout
    notes = (list(notes) + ["", "", "", ""])[:4]

    # Detect orientation by reading the first frame's dimensions.
    # Only TRULY panoramic banners (728×90, 970×90, 320×50 — width ≥ 2× height)
    # get the stacked-rows layout. Square-ish banners like 300×250 still go
    # into the four-columns layout the same as 160×600 / 300×600 do.
    is_wide = False
    if frame_paths:
        try:
            from PIL import Image
            with Image.open(frame_paths[0]) as im:
                is_wide = (im.width / max(im.height, 1)) >= 2.0
        except Exception:
            pass

    # Per-frame block — different layout for wide vs tall banners
    blocks_html = []
    for i in range(4):
        img_src = frame_imgs[i] if i < len(frame_imgs) else ""
        note    = notes[i].replace("\n", "<br>")
        if is_wide:
            blocks_html.append(f"""
            <div class="frame frame-wide">
              <div class="frame-num">Frame {i+1}</div>
              <div class="frame-row">
                <img class="frame-img" src="{img_src}" alt="Frame {i+1}">
                <div class="frame-text">
                  <div class="notes-label">Animation Notes:</div>
                  <div class="notes-body">{note or '&nbsp;'}</div>
                </div>
              </div>
            </div>""")
        else:
            blocks_html.append(f"""
            <div class="frame frame-tall">
              <div class="frame-num">Frame {i+1}</div>
              <img class="frame-img" src="{img_src}" alt="Frame {i+1}">
              <div class="notes-label">Animation Notes:</div>
              <div class="notes-body">{note or '&nbsp;'}</div>
            </div>""")

    layout_class = "wide" if is_wide else "tall"
    full_header = (header or "").strip()
    if banner.get("name"):
        # Append the banner size to the header (e.g. "… – 300x250")
        full_header += (" – " if full_header else "") + banner["name"]

    return f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"><style>
  *{{box-sizing:border-box;margin:0;padding:0;}}
  html,body{{width:{page_w}px;height:{page_h}px;}}
  body{{
    font-family:-apple-system,"Helvetica Neue",Arial,sans-serif;
    background:#fff; color:#222;
    padding:32px 40px;
    display:flex; flex-direction:column;
  }}
  .header{{
    font-size:22px; font-weight:600; letter-spacing:.01em;
    color:#333; padding-bottom:14px;
    border-bottom:1px solid #eee; margin-bottom:18px;
    display:flex; align-items:center;
  }}
  .header .brand{{color:{_STORYBOARD_BRAND}; font-weight:700;}}
  .brand-dot{{
    width:20px; height:20px; border-radius:50%; background:{_STORYBOARD_BRAND};
    margin-left:auto;
  }}
  .frames{{flex:1; display:flex; gap:18px; min-height:0;}}
  .frames.tall  {{flex-direction:row;}}
  .frames.wide  {{flex-direction:column;}}
  .frame{{
    flex:1; display:flex; flex-direction:column; gap:8px;
    min-height:0; min-width:0;
  }}
  .frame-num{{
    font-size:14px; font-weight:700; color:#555;
    letter-spacing:.05em; text-transform:uppercase;
  }}
  .frame-tall .frame-img{{
    width:auto; max-width:100%; max-height:75%;
    object-fit:contain; object-position:top center;
    align-self:center; border:1px solid #d0d0d0;
  }}
  .frame-wide .frame-row{{display:flex; gap:14px; align-items:flex-start;}}
  .frame-wide .frame-img{{
    flex:0 0 70%; max-width:70%; height:auto;
    object-fit:contain; border:1px solid #d0d0d0;
  }}
  .frame-wide .frame-text{{flex:1; min-width:0;}}
  .notes-label{{
    font-size:12px; font-weight:700; color:{_STORYBOARD_BRAND};
    margin-top:8px; letter-spacing:.04em;
  }}
  .notes-body{{
    font-size:12px; line-height:1.5; color:#333;
    white-space:pre-wrap; word-wrap:break-word;
  }}
  .footer{{
    font-size:10px; color:#888; padding-top:10px;
    border-top:1px solid #eee; margin-top:14px;
  }}
</style></head><body>
  <div class="header">
    <span>{full_header}</span>
    <div class="brand-dot"></div>
  </div>
  <div class="frames {layout_class}">
    {"".join(blocks_html)}
  </div>
  <div class="footer">{footer or ""}</div>
</body></html>"""


# ── Existing-PDF page extraction + region detection ──────────────────────────

def _extract_pdf_pages_to_pngs(pdf_path: Path):
    """Return a Pillow Image per page, decoded from each page's largest
    embedded raster image (matches the format of these flattened storyboard
    PDFs). Pages with no images get a white placeholder."""
    from pypdf import PdfReader
    from PIL import Image
    import io
    reader = PdfReader(str(pdf_path))
    pages = []
    for page in reader.pages:
        imgs = list(page.images)
        if not imgs:
            w, h = int(float(page.mediabox.width)), int(float(page.mediabox.height))
            pages.append(Image.new("RGB", (max(w, 100) * 3, max(h, 100) * 3), "white"))
            continue
        imgs.sort(key=lambda i: i.image.size[0] * i.image.size[1], reverse=True)
        biggest = imgs[0]
        try:
            pil = Image.open(io.BytesIO(biggest.data)).convert("RGB")
        except Exception:
            pil = biggest.image.convert("RGB")
        pages.append(pil)
    return pages


def _detect_banner_regions(page_img):
    """Find banner-image rectangles on a storyboard page.

    Strategy: banner ads have **saturated colours** (purple, teal, orange, etc.)
    while frame labels and animation notes are nearly grayscale text. Sample a
    row inside each banner's vertical bounds and look for runs of high-saturation
    columns — those are the banner widths. Then for each strip, walk vertically
    to find the actual top/bottom of the colored region.

    Returns a list of {x, y, w, h, aspect} dicts in reading order (top-left first).
    """
    from PIL import Image

    SAT_THRESH    = 40    # HSV saturation; pixel is "colored" when > this
    MIN_STRIP_FRAC= 0.025  # strip must cover ≥2.5% of the page dimension
    MAX_ASPECT    = 14.0
    MIN_ASPECT    = 0.08
    DARK_GRAY_THRESH = 220  # for the boundary-extension step (anything not-white)

    # Downscale for speed.
    W0, H0 = page_img.size
    if max(W0, H0) > 1200:
        scale = 1200 / max(W0, H0)
        small = page_img.resize((int(W0 * scale), int(H0 * scale)))
    else:
        scale = 1.0
        small = page_img
    sw, sh = small.size

    sat_channel = small.convert("HSV").split()[1]
    sat_px      = sat_channel.load()
    gray_channel = small.convert("L")
    gray_px      = gray_channel.load()

    def _runs(arr, thresh, min_run_len):
        runs = []
        in_run = False
        start  = 0
        for i, v in enumerate(arr):
            if v > thresh and not in_run:
                in_run = True; start = i
            elif v <= thresh and in_run:
                in_run = False
                if i - start >= min_run_len:
                    runs.append((start, i))
        if in_run and len(arr) - start >= min_run_len:
            runs.append((start, len(arr)))
        return runs

    def _strip_bounds_v(x0, x1):
        """For a tall (vertical-strip) banner: return (y0, y1) of the banner
        image. Walks the saturated core downward to capture an ISI panel
        directly below the colored area, but caps extension at 1.5× the core
        height so we don't merge into the next stacked element.
        """
        col_sat = [max(sat_px[x, y] for x in range(x0, x1)) for y in range(sh)]
        strips = _runs(col_sat, SAT_THRESH, max(8, int(sh * 0.04)))
        if not strips:
            return None
        core = max(strips, key=lambda r: r[1] - r[0])
        cy0, cy1 = core
        core_h = cy1 - cy0
        max_extra = int(core_h * 1.5)

        col_gray_min = [min(gray_px[x, y] for x in range(x0, x1)) for y in range(sh)]
        # Walk down — capture trailing non-white rows (ISI panel below banner).
        white_run = 0
        far = min(sh, cy1 + max_extra)
        y = cy1
        while y < far:
            if col_gray_min[y] >= DARK_GRAY_THRESH:
                white_run += 1
                if white_run >= 3:
                    break
            else:
                white_run = 0
                cy1 = y + 1
            y += 1
        return (cy0, cy1)

    def _strip_bounds_h(y0, y1):
        """For a wide (horizontal-strip) banner: return (x0, x1) of the
        saturated core only. We do NOT walk-extend horizontally — animation
        notes typically sit to the right of the banner image and would get
        engulfed.
        """
        row_sat = [max(sat_px[x, y] for y in range(y0, y1)) for x in range(sw)]
        strips = _runs(row_sat, SAT_THRESH, max(8, int(sw * 0.04)))
        if not strips:
            return None
        core = max(strips, key=lambda r: r[1] - r[0])
        return core

    candidates = []

    # ── Strategy A: 4-column tall layout ─────────────────────────────────────
    # Sample a horizontal row near 30% page height — inside the banner image
    # in tall-banner layouts (160×600, 300×600, 300×250).
    sample_y = int(sh * 0.30)
    row_sat  = [max(sat_px[x, y] for y in range(sample_y, sample_y + max(20, int(sh * 0.05))))
                for x in range(sw)]
    col_strips = _runs(row_sat, SAT_THRESH, max(20, int(sw * MIN_STRIP_FRAC)))

    if 2 <= len(col_strips) <= 8:
        for x0, x1 in col_strips:
            v = _strip_bounds_v(x0, x1)
            if not v:
                continue
            y0, y1 = v
            w, h = x1 - x0, y1 - y0
            ar = w / max(h, 1)
            if ar > MAX_ASPECT or ar < MIN_ASPECT:
                continue
            candidates.append({"x": x0, "y": y0, "w": w, "h": h, "aspect": ar})

    # ── Strategy B: 4-row wide layout (e.g., 728×90 stacked vertically) ──────
    if not candidates:
        sample_x = int(sw * 0.30)
        col_sat  = [max(sat_px[x, y] for x in range(sample_x, sample_x + max(20, int(sw * 0.05))))
                    for y in range(sh)]
        row_strips = _runs(col_sat, SAT_THRESH, max(15, int(sh * MIN_STRIP_FRAC)))
        if 2 <= len(row_strips) <= 8:
            for y0, y1 in row_strips:
                h_bounds = _strip_bounds_h(y0, y1)
                if not h_bounds:
                    continue
                x0, x1 = h_bounds
                w, h = x1 - x0, y1 - y0
                ar = w / max(h, 1)
                if ar > MAX_ASPECT or ar < MIN_ASPECT:
                    continue
                candidates.append({"x": x0, "y": y0, "w": w, "h": h, "aspect": ar})

    # Storyboards always have ≤4 banner cells per page. If we found more,
    # keep the 4 largest (filters out tiny noise blobs / dust specks).
    if len(candidates) > 4:
        candidates.sort(key=lambda r: r["w"] * r["h"], reverse=True)
        candidates = candidates[:4]

    # Sort top-to-bottom, then left-to-right (reading order).
    candidates.sort(key=lambda r: (r["y"], r["x"]))

    # Scale back up to original page coordinates.
    if scale != 1.0:
        inv = 1.0 / scale
        for c in candidates:
            c["x"] = int(c["x"] * inv)
            c["y"] = int(c["y"] * inv)
            c["w"] = int(c["w"] * inv)
            c["h"] = int(c["h"] * inv)
            c["kind"] = "banner"
    else:
        for c in candidates:
            c["kind"] = "banner"

    # If there are no banner-shaped colored regions, this is likely an ISI
    # text page. Find the page's content bbox so we can paste a captured
    # isi_full.png into that area when the user assigns a banner to it.
    if not candidates:
        gs   = page_img.convert("L")
        mask = gs.point(lambda v: 255 if v < DARK_GRAY_THRESH else 0, mode="L")
        bbox = mask.getbbox()
        if bbox:
            x0, y0, x1, y1 = bbox
            w, h = x1 - x0, y1 - y0
            # Sanity check — must cover at least 15% of page area to be the
            # "main content block" of an ISI page (rules out tiny page-number
            # text on otherwise-blank pages).
            page_area = page_img.size[0] * page_img.size[1]
            if w * h >= 0.15 * page_area:
                candidates.append({
                    "x": x0, "y": y0, "w": w, "h": h,
                    "aspect": w / max(h, 1), "kind": "isi",
                })

    return candidates


def _suggest_banner_size(regions):
    """Given detected regions, return the most-common aspect-ratio category
    (e.g., '160x600', '300x600', '300x250', '728x90', or None)."""
    if not regions:
        return None
    # Common banner aspect ratios (w/h)
    refs = [
        ("160x600", 160 / 600),
        ("300x600", 300 / 600),
        ("300x250", 300 / 250),
        ("728x90",  728 /  90),
        ("970x250", 970 / 250),
        ("970x90",  970 /  90),
        ("320x50",  320 /  50),
    ]
    votes = {}
    for r in regions:
        a = r["aspect"]
        best, best_d = None, 1e9
        for name, ref in refs:
            d = abs(a - ref) / ref
            if d < best_d:
                best, best_d = name, d
        if best_d < 0.30:   # within 30% of a known ratio
            votes[best] = votes.get(best, 0) + 1
    if not votes:
        return None
    return max(votes.items(), key=lambda kv: kv[1])[0]


def _paste_frame_into_region(page_img, region, frame_path,
                             anchor_top=False, fill=False):
    """Resize a captured frame to fit the region and paste it onto page_img.

    `anchor_top=True` anchors at the top-left of the region (rather than
    centering vertically). Used for banner-frame pages so the new banner
    sits where the old one was, leaving labels/notes below untouched.

    `fill=True` STRETCHES the frame to fill the region exactly (small aspect
    mismatch is OK — invisible at typical banner detail levels). Use when
    the region's aspect ratio is already close to the banner's; the trade-off
    is filling the whole detected area instead of leaving leftover edges of
    the old banner showing through.
    """
    from PIL import Image
    rx, ry, rw, rh = region["x"], region["y"], region["w"], region["h"]
    try:
        frame = Image.open(frame_path).convert("RGB")
    except Exception:
        return
    fw, fh = frame.size
    if fill:
        nw, nh = rw, rh
    else:
        scale = min(rw / fw, rh / fh)
        nw, nh = max(1, int(fw * scale)), max(1, int(fh * scale))
    frame  = frame.resize((nw, nh), Image.LANCZOS)
    ox = rx + (rw - nw) // 2          # always centered horizontally
    oy = ry if anchor_top else ry + (rh - nh) // 2
    page_img.paste(frame, (ox, oy))


def _adjust_region_to_banner_aspect(region, banner_w, banner_h):
    """Re-shape a detected region so its aspect ratio matches the assigned
    banner. Anchors the new region at the top-left of the detected region —
    that way the new banner replaces only the OLD banner's image, leaving any
    surrounding ISI / labels / notes exactly where they are.

    Tolerance: if detection is within 15% of the banner's aspect ratio, the
    detection is considered "close enough" and the region is left unchanged,
    so the captured frame fills the full detected area (any slight stretch is
    invisible). Outside that tolerance, we trim either width or height to
    match the banner's aspect exactly.
    """
    rx, ry, rw, rh = region["x"], region["y"], region["w"], region["h"]
    banner_aspect   = banner_w / max(banner_h, 1)
    detected_aspect = rw / max(rh, 1)
    if detected_aspect > banner_aspect * 1.15:
        new_w = max(1, int(rh * banner_aspect))
        new_w = min(new_w, rw)
        return {**region, "w": new_w, "h": rh, "aspect": banner_aspect}
    if detected_aspect < banner_aspect * 0.85:
        new_h = max(1, int(rw / banner_aspect))
        new_h = min(new_h, rh)
        return {**region, "w": rw, "h": new_h, "aspect": banner_aspect}
    return region


def _page_thumb_b64(page_img, max_w: int = 640):
    """Downscale a PIL page image and return a data-URL base64 PNG."""
    import io, base64
    from PIL import Image
    w, h = page_img.size
    if w > max_w:
        scale = max_w / w
        small = page_img.resize((max_w, int(h * scale)), Image.LANCZOS)
    else:
        small = page_img
    buf = io.BytesIO()
    small.save(buf, format="PNG", optimize=True)
    return ("data:image/png;base64," + base64.b64encode(buf.getvalue()).decode(),
            small.size[0], small.size[1])


def _diff_page_boxes(old_img, new_img, *, block: int = 16, threshold: int = 25,
                     cell_min_brightness: int = 5, min_blocks: int = 2):
    """Compute change-bounding-boxes between two page images for the CPFP tab.

    Approach (no numpy dependency required):
      1. Compute per-pixel max-channel absolute difference via PIL ImageChops.
      2. Threshold to a binary mask (255 where any channel differs by more
         than `threshold`).
      3. Box-filter resize the mask down to a coarse `block`-cell grid; each
         cell's value is then proportional to the fraction of dirty pixels
         it contained. Threshold again at `cell_min_brightness` to mark the
         cell as "dirty".
      4. 4-connected BFS over the dirty cells to merge clusters into
         bounding rectangles. Returns [{x, y, w, h}, …] in old's pixel space
         (new is resized to match if dimensions differ).

    Skips clusters smaller than `min_blocks` cells (filters salt-and-pepper
    noise from anti-aliasing artifacts at edges).
    """
    from PIL import Image, ImageChops
    if old_img.size != new_img.size:
        new_img = new_img.resize(old_img.size, Image.LANCZOS)
    a = old_img.convert("RGB")
    b = new_img.convert("RGB")
    # Per-pixel max-channel diff: ImageChops.difference is per-channel abs.
    # Take the per-pixel max across R/G/B by chaining ImageChops.lighter.
    diff = ImageChops.difference(a, b)
    r, g, bl = diff.split()
    max_chan = ImageChops.lighter(ImageChops.lighter(r, g), bl)
    # Binary mask: 255 above threshold, else 0.
    mask = max_chan.point(lambda v: 255 if v > threshold else 0, mode="L")
    if not mask.getbbox():
        return []

    W, H = mask.size
    bw_cells = (W + block - 1) // block
    bh_cells = (H + block - 1) // block
    # Box-filter downsample: each output pixel is the AVG of its `block`×
    # `block` source pixels — so a cell's brightness is the fraction of
    # dirty pixels × 255 (e.g. 1 dirty pixel in a 16×16 cell ≈ value 1).
    # `cell_min_brightness` sets how dense a cell must be to count as dirty.
    small = mask.resize((bw_cells, bh_cells), Image.BOX)
    pix = list(small.getdata())  # row-major
    dirty = [
        [pix[row * bw_cells + col] > cell_min_brightness for col in range(bw_cells)]
        for row in range(bh_cells)
    ]

    visited = [[False] * bw_cells for _ in range(bh_cells)]
    boxes = []
    for r0 in range(bh_cells):
        for c0 in range(bw_cells):
            if not dirty[r0][c0] or visited[r0][c0]:
                continue
            stack = [(r0, c0)]
            cnt = 0
            min_r = max_r = r0
            min_c = max_c = c0
            while stack:
                rr, cc = stack.pop()
                if visited[rr][cc] or not dirty[rr][cc]:
                    continue
                visited[rr][cc] = True
                cnt += 1
                if rr < min_r: min_r = rr
                if rr > max_r: max_r = rr
                if cc < min_c: min_c = cc
                if cc > max_c: max_c = cc
                if rr > 0:            stack.append((rr - 1, cc))
                if rr < bh_cells - 1: stack.append((rr + 1, cc))
                if cc > 0:            stack.append((rr, cc - 1))
                if cc < bw_cells - 1: stack.append((rr, cc + 1))
            if cnt < min_blocks:
                continue
            x = min_c * block
            y = min_r * block
            boxes.append({
                "x": int(x), "y": int(y),
                "w": int(min((max_c - min_c + 1) * block, W - x)),
                "h": int(min((max_r - min_r + 1) * block, H - y)),
            })
    return boxes


# ── Anthropic API integration for the CPFP Claude analysis ───────────────────
# We talk to api.anthropic.com directly via urllib (no extra dependency on the
# `anthropic` SDK or the `claude` CLI). The API key is stored in a small
# user-local config file with 0600 perms; the user pastes it once via the
# "🔌 Connect Claude" button and it persists across runs.

_CPFP_CONFIG_PATH    = Path.home() / ".banner_app" / "config.json"
_CPFP_DEFAULT_MODEL  = "claude-sonnet-4-5-20250929"
_ANTHROPIC_VERSION   = "2023-06-01"
_ANTHROPIC_ENDPOINT  = "https://api.anthropic.com/v1/messages"


def _cpfp_load_config() -> dict:
    try:
        return json.loads(_CPFP_CONFIG_PATH.read_text())
    except (FileNotFoundError, json.JSONDecodeError):
        return {}


def _cpfp_save_config(cfg: dict):
    _CPFP_CONFIG_PATH.parent.mkdir(parents=True, exist_ok=True)
    _CPFP_CONFIG_PATH.write_text(json.dumps(cfg, indent=2))
    try:
        os.chmod(_CPFP_CONFIG_PATH, 0o600)   # owner-only — key is sensitive.
    except Exception:
        pass


def _cpfp_get_api_key():
    return _cpfp_load_config().get("anthropic_api_key")


def _cpfp_anthropic_call(api_key: str, prompt: str, *,
                         model: str = _CPFP_DEFAULT_MODEL,
                         max_tokens: int = 4096,
                         timeout: int = 180) -> str:
    """Call the Anthropic Messages API with a single user-message prompt.
    Returns the concatenated text content from the response. Raises
    urllib.error.HTTPError on non-2xx.
    """
    payload = json.dumps({
        "model":      model,
        "max_tokens": max_tokens,
        "messages":   [{"role": "user", "content": prompt}],
    }).encode("utf-8")
    req = urllib.request.Request(
        _ANTHROPIC_ENDPOINT,
        data=payload,
        headers={
            "x-api-key":         api_key,
            "anthropic-version": _ANTHROPIC_VERSION,
            "content-type":      "application/json",
        },
        method="POST",
    )
    with urllib.request.urlopen(req, timeout=timeout) as resp:
        data = json.loads(resp.read().decode("utf-8"))
    parts = [b.get("text", "")
             for b in data.get("content", [])
             if b.get("type") == "text"]
    return "\n".join(parts)


def _build_cpfp_claude_prompt(old_name: str, new_name: str, pages_data: list) -> str:
    """Format a CPFP review request for the `claude` CLI. The prompt asks for
    structured JSON so the frontend can render severity badges and inline
    annotations alongside the deterministic diff results.
    """
    lines = []
    for p in pages_data:
        if p.get("added"):
            lines.append(
                f"Page {p['page_num']} — ADDED in new (no counterpart in old).\n"
                f"  New text: {p.get('new_text', '')!r}")
        elif p.get("removed"):
            lines.append(
                f"Page {p['page_num']} — REMOVED from new (was in old).\n"
                f"  Old text: {p.get('old_text', '')!r}")
        else:
            lines.append(
                f"Page {p['page_num']} — visual regions: {p['visual_change_count']}, "
                f"text-word changes: {p['text_change_count']}.\n"
                f"  Old text: {p.get('old_text', '')!r}\n"
                f"  New text: {p.get('new_text', '')!r}")
    body = "\n\n".join(lines)
    return f"""You are reviewing a Change-Proof Full-Proof (CPFP) for pharma marketing storyboards.

Two storyboards were compared:
- OLD: {old_name}
- NEW: {new_name}

A pixel-level + word-level diff was already run by the app — your job is to interpret what changed and explain it to a regulatory reviewer in plain English. Focus on identifying:
1) Intentional content changes (copy edits, new claims, dosing updates).
2) Regulatory-critical changes (safety information, contraindications, dose, brand/generic name, ISI integrity).
3) Cosmetic/incidental changes (formatting, position, font weight, anti-aliasing).

For each page that has changes, identify each distinct change and tag it with:
- "kind":      "text" | "visual" | "both"
- "severity":  "critical"  (regulatory / safety / dose),
               "content"   (intentional copy edit),
               "cosmetic"  (formatting, position, font, anti-aliasing)
- "summary":   1-sentence plain-English description ("Dosage changed from 100mg to 200mg.")
- "confidence": "high" | "medium" | "low"

Also write a 2-3 sentence "overall_summary" describing what changed across the whole storyboard.

Respond with ONLY this JSON shape, no other text or commentary:

```json
{{
  "overall_summary": "...",
  "pages": [
    {{"page_num": 1, "changes": [
      {{"kind": "text", "severity": "critical",
        "summary": "Dosage changed from 100mg to 200mg.", "confidence": "high"}}
    ]}}
  ]
}}
```

Pages with changes:

{body}
"""


def _extract_cpfp_json_payload(raw: str):
    """Try to parse a JSON object out of Claude's raw stdout. Looks for a
    fenced ```json ... ``` block first (which is what we asked for), then
    falls back to the first {…} block. Returns the parsed dict, or None
    if no valid JSON was found."""
    import re, json
    # 1) ```json … ``` fenced block (preferred).
    m = re.search(r"```(?:json)?\s*\n(.*?)```", raw, re.DOTALL | re.IGNORECASE)
    if m:
        try:
            return json.loads(m.group(1))
        except json.JSONDecodeError:
            pass
    # 2) First {...} block. Greedy across lines.
    m = re.search(r"(\{[\s\S]*\})", raw)
    if m:
        try:
            return json.loads(m.group(1))
        except json.JSONDecodeError:
            pass
    return None


def _snap_region_to_content(page_img, region, pad: int = 0,
                            white_threshold: int = 244):
    """Find the bounding box of non-white pixels inside `region` of page_img.

    Treats any pixel where all three channels are ≥ `white_threshold` as
    background (catches near-white scanner artifacts and JPEG-compression noise
    in addition to pure 255s). Returns a tightened {x,y,w,h} in PAGE coords,
    optionally padded outward by `pad` pixels (clamped to page bounds).
    Returns None if the region is entirely white.
    """
    from PIL import Image
    pw, ph = page_img.size
    rx = max(0, int(region.get("x", 0)))
    ry = max(0, int(region.get("y", 0)))
    rw = max(1, int(region.get("w", 0)))
    rh = max(1, int(region.get("h", 0)))
    rx = min(rx, pw - 1)
    ry = min(ry, ph - 1)
    rw = min(rw, pw - rx)
    rh = min(rh, ph - ry)
    crop = page_img.crop((rx, ry, rx + rw, ry + rh)).convert("RGB")

    # Use NumPy if available — much faster than per-pixel Python loops.
    try:
        import numpy as np
        arr = np.asarray(crop)
        # Mask: True where pixel is "content" (any channel below white threshold).
        mask = (arr < white_threshold).any(axis=2)
        if not mask.any():
            return None
        ys = np.where(mask.any(axis=1))[0]
        xs = np.where(mask.any(axis=0))[0]
        y0, y1 = int(ys[0]), int(ys[-1])
        x0, x1 = int(xs[0]), int(xs[-1])
    except ImportError:
        # Pure-PIL fallback: getbbox after converting to a binary mask.
        gray = crop.convert("L").point(
            lambda v: 0 if v >= white_threshold else 255, mode="L")
        bbox = gray.getbbox()
        if not bbox:
            return None
        x0, y0, x1, y1 = bbox
        # getbbox returns half-open coords; convert to inclusive.
        x1 = max(x0, x1 - 1)
        y1 = max(y0, y1 - 1)

    # Convert local crop coords → page coords; add padding; clamp.
    nx = max(0, rx + x0 - pad)
    ny = max(0, ry + y0 - pad)
    nw = min(pw - nx, (x1 - x0 + 1) + 2 * pad)
    nh = min(ph - ny, (y1 - y0 + 1) + 2 * pad)
    return {"x": int(nx), "y": int(ny), "w": int(nw), "h": int(nh)}


def _route_replace_mappings(sess, mappings, return_thumbs: bool):
    """Pick the right backend for either mapping format.

    New (per-region) format: each entry has a `regions` list with `image_path`
    on each region. Old (per-page) format: each entry has `banner_slot`.
    """
    is_new = bool(mappings) and any(
        isinstance(m, dict) and "regions" in m for m in mappings
    )
    if is_new:
        return _apply_per_region_mappings(sess, mappings, return_thumbs)
    return _apply_replacement_mappings(sess, mappings, return_thumbs)


def _apply_per_region_mappings(sess, pages_data, return_thumbs: bool):
    """Apply per-region mappings (new UI: each detected region picks its own
    captured-frame PNG). For each page, walk the user-edited regions; for each
    region with a chosen image_path, wipe that rect on the page with white and
    paste the chosen image scaled-to-fit (centered).

    ISI chain slicing: when the user assigns the *same* `isi_full.png` to two
    or more regions across the document, the ISI image is sliced vertically
    across those regions in document order — region 1 shows the top portion,
    region 2 picks up where region 1 left off, etc. Slice heights are weighted
    by each region's page-px height, so a tall region gets a proportionally
    larger slice. Frame PNGs and one-region ISI assignments are pasted whole
    (no slicing).

    `pages_data` shape (sent from the client):
        [{page_num, page_w, page_h, regions: [{x, y, w, h, kind, image_path}]}]

    Returns thumb-payload dicts (return_thumbs=True) or full-res PIL Images.
    """
    from PIL import Image
    server_pages = {e["page_num"]: e for e in sess.existing_storyboard_pages}
    results, full_pages = [], []

    # ── Pre-compute ISI chain slices ──────────────────────────────────────────
    # Group regions by image_path, but only if the path looks like an ISI
    # capture (filename "isi_full.png"). Frame PNGs are treated as one-shot
    # pastes even when reused across regions.
    def _is_isi(path):
        try:
            return Path(path).name == "isi_full.png"
        except Exception:
            return False

    isi_groups = {}   # path -> list of (page_num, region_idx, page_h)
    for p in pages_data or []:
        pn = p.get("page_num")
        for ri, r in enumerate(p.get("regions") or []):
            path = r.get("image_path")
            if path and _is_isi(path):
                isi_groups.setdefault(path, []).append(
                    (pn, ri, max(1, int(r.get("h", 1) or 1))))

    chain_slice = {}   # (page_num, region_idx) -> (top_frac, bot_frac)
    for path, members in isi_groups.items():
        if len(members) < 2:
            continue
        total = sum(m[2] for m in members) or 1
        cum = 0
        for (pn, ri, h) in members:
            top = cum / total
            cum += h
            bot = cum / total
            chain_slice[(pn, ri)] = (top, bot)

    img_cache = {}
    def _load(path):
        if path not in img_cache:
            try:
                img_cache[path] = Image.open(path).convert("RGB")
            except Exception:
                img_cache[path] = None
        return img_cache[path]

    for p in pages_data or []:
        srv = server_pages.get(p.get("page_num"))
        if not srv:
            continue
        original = srv["full_image"]
        page = original.copy()

        for ri, r in enumerate(p.get("regions") or []):
            ip = r.get("image_path")
            if not ip:
                continue
            full_img = _load(ip)
            if full_img is None:
                continue

            # Slice the ISI vertically when this region is part of a chain or
            # has explicit user-set boundaries. Per-region `isi_top`/`isi_bot`
            # overrides win over the auto-computed (height-weighted) chain.
            slice_info = chain_slice.get((p.get("page_num"), ri))
            user_top, user_bot = r.get("isi_top"), r.get("isi_bot")
            if (user_top is not None and user_bot is not None
                    and _is_isi(ip)):
                try:
                    slice_info = (max(0.0, min(1.0, float(user_top))),
                                  max(0.0, min(1.0, float(user_bot))))
                except (TypeError, ValueError):
                    pass
            if slice_info:
                top_frac, bot_frac = slice_info
                ih = full_img.height
                top_px = max(0, int(ih * top_frac))
                bot_px = min(ih, max(top_px + 1, int(ih * bot_frac)))
                img = full_img.crop((0, top_px, full_img.width, bot_px))
            else:
                img = full_img

            rx = max(0, int(r.get("x", 0)))
            ry = max(0, int(r.get("y", 0)))
            rw = max(1, int(r.get("w", 0)))
            rh = max(1, int(r.get("h", 0)))
            rx = min(rx, page.size[0] - 1)
            ry = min(ry, page.size[1] - 1)
            rw = min(rw, page.size[0] - rx)
            rh = min(rh, page.size[1] - ry)
            iw, ih2 = img.size
            scale = min(rw / iw, rh / ih2) if iw and ih2 else 1.0
            nw = max(1, int(iw * scale))
            nh = max(1, int(ih2 * scale))
            img2 = img.resize((nw, nh), Image.LANCZOS)

            # Wipe the user's edited rect (always). For *frame* regions, also
            # wipe the original auto-detected rect — kills any leftover frame
            # edges when the user moved/resized the region. For ISI regions
            # we deliberately skip the original_rect wipe: the auto-detect
            # often picks up a larger area than the actual ISI panel
            # (including page titles or surrounding text), and wiping that
            # would erase content the user wants to keep.
            if not _is_isi(ip):
                orig = r.get("original_rect") or {}
                if orig:
                    ox_o = max(0, int(orig.get("x", 0)))
                    oy_o = max(0, int(orig.get("y", 0)))
                    ow_o = max(0, int(orig.get("w", 0)))
                    oh_o = max(0, int(orig.get("h", 0)))
                    ox_o = min(ox_o, page.size[0] - 1)
                    oy_o = min(oy_o, page.size[1] - 1)
                    ow_o = min(ow_o, page.size[0] - ox_o)
                    oh_o = min(oh_o, page.size[1] - oy_o)
                    if ow_o > 0 and oh_o > 0:
                        page.paste(Image.new("RGB", (ow_o, oh_o), "white"),
                                   (ox_o, oy_o))
            page.paste(Image.new("RGB", (rw, rh), "white"), (rx, ry))

            # ISI: top-align within the region (oy = ry). The full safety
            # information should always start at the top of its slot, with
            # any extra space falling at the bottom — never centered.
            # Frames stay vertically centered (existing behavior).
            ox = rx + (rw - nw) // 2
            if _is_isi(ip):
                oy = ry
            else:
                oy = ry + (rh - nh) // 2
            page.paste(img2, (ox, oy))

            # Thin gray border around every pasted ISI image — visual cue
            # separating it from surrounding storyboard content. Frame PNGs
            # are not bordered.
            if _is_isi(ip):
                try:
                    from PIL import ImageDraw
                    ImageDraw.Draw(page).rectangle(
                        [ox, oy, ox + nw - 1, oy + nh - 1],
                        outline=(170, 170, 170), width=1)
                except Exception:
                    pass

        if return_thumbs:
            thumb_url, tw, th = _page_thumb_b64(page)
            results.append({
                "page_num": p.get("page_num"),
                "thumb":    thumb_url,
                "thumb_w":  tw, "thumb_h": th,
                "page_w":   original.size[0],
                "page_h":   original.size[1],
                "regions":  p.get("regions") or [],
            })
        else:
            full_pages.append(page)

    return results if return_thumbs else full_pages


def _apply_replacement_mappings(sess, mappings, return_thumbs: bool):
    """Apply page→slot mappings to the uploaded storyboard pages. (Legacy —
    kept for backwards compat; the new UI uses _apply_per_region_mappings.)

    Behaviour by region kind:
    - kind="banner":  paste the i-th captured frame, with the region trimmed
                      to the assigned banner's exact aspect ratio (so a
                      300×250 banner doesn't leak into the ISI text below).
    - kind="isi":     paste the captured isi_full.png. If a run of consecutive
                      ISI pages share the same assigned banner, the captured
                      ISI image is sliced vertically across them — page 1
                      shows the top, page 2 picks up where it left off, etc.

    If `return_thumbs=True`, returns a list of thumbnail-payload dicts ready
    for the UI. If False, returns a list of full-res PIL Images.
    """
    from PIL import Image
    by_page = {m.get("page_num"): m.get("banner_slot") for m in mappings if m}

    # ── Pre-compute ISI sequence groups ─────────────────────────────────────
    # Walk pages in order; group consecutive ISI pages assigned to the same
    # banner so we can slice that banner's isi_full.png across them.
    isi_groups = []   # list of {slot, page_nums:[...], h, w} (one per group)
    current = None
    for entry in sess.existing_storyboard_pages:
        regions = entry["regions"]
        is_isi  = (len(regions) == 1 and regions[0].get("kind") == "isi")
        slot    = by_page.get(entry["page_num"]) if is_isi else None
        if is_isi and slot:
            if current and current["slot"] == slot:
                current["page_nums"].append(entry["page_num"])
            else:
                if current:
                    isi_groups.append(current)
                current = {"slot": slot, "page_nums": [entry["page_num"]]}
        else:
            if current:
                isi_groups.append(current)
                current = None
    if current:
        isi_groups.append(current)
    # page_num → (group_index, position_within_group, group_size)
    isi_index = {}
    for g in isi_groups:
        for i, pn in enumerate(g["page_nums"]):
            isi_index[pn] = (g["slot"], i, len(g["page_nums"]))

    # Cache loaded ISI images per slot (avoid re-decoding the big PNG)
    isi_cache = {}

    def _isi_slice_for(slot, pos, group_size, region_w, region_h):
        """Return the vertical slice of slot's isi_full.png that belongs at
        position `pos` of `group_size` total ISI pages. Slice is sized to
        fit the region (height-fit, preserving aspect) so the user can paste
        it directly with no further scaling."""
        if slot not in isi_cache:
            cap_sub = (sess.last_bundle / sess.multi_slots[slot]["name"]) if sess.last_bundle else None
            if not cap_sub or not (cap_sub / "isi_full.png").exists():
                isi_cache[slot] = None
            else:
                isi_cache[slot] = Image.open(cap_sub / "isi_full.png").convert("RGB")
        img = isi_cache[slot]
        if img is None:
            return None
        H = img.height
        slice_top = int(H * pos / max(group_size, 1))
        slice_bot = int(H * (pos + 1) / max(group_size, 1))
        return img.crop((0, slice_top, img.width, slice_bot))

    results = []
    full_pages = []
    for entry in sess.existing_storyboard_pages:
        page_num = entry["page_num"]
        original = entry["full_image"]
        page = original.copy()
        slot = by_page.get(page_num)

        if slot and slot in sess.multi_slots:
            info    = sess.multi_slots[slot]
            cap_sub = (sess.last_bundle / info["name"]) if sess.last_bundle else None
            frame_paths  = []
            isi_png_path = None
            if cap_sub and cap_sub.exists():
                frame_paths = sorted(
                    p for p in cap_sub.glob("*.png")
                    if p.name not in ("isi_full.png",)
                )
                if (cap_sub / "isi_full.png").exists():
                    isi_png_path = cap_sub / "isi_full.png"

            banner_w = info.get("w") or 160
            banner_h = info.get("h") or 600

            for i, region in enumerate(entry["regions"]):
                kind = region.get("kind", "banner")
                if kind == "isi":
                    # Multi-page ISI continuation
                    grp = isi_index.get(page_num)
                    if not grp:
                        continue
                    _, pos, group_size = grp
                    sl = _isi_slice_for(slot, pos, group_size,
                                        region["w"], region["h"])
                    if sl is None:
                        continue
                    # Wipe the old ISI content, then paste this slice scaled
                    # to fit the region (height-priority since the slice is
                    # tall and narrow).
                    wipe = Image.new("RGB", (region["w"], region["h"]), "white")
                    page.paste(wipe, (region["x"], region["y"]))
                    sw, sh = sl.size
                    scale  = min(region["w"] / sw, region["h"] / sh)
                    nw, nh = max(1, int(sw * scale)), max(1, int(sh * scale))
                    sl_resized = sl.resize((nw, nh), Image.LANCZOS)
                    ox = region["x"] + (region["w"] - nw) // 2
                    oy = region["y"]
                    page.paste(sl_resized, (ox, oy))
                else:
                    # Banner-frame page — re-shape region to match the
                    # assigned banner's aspect ratio so we don't paint into
                    # the ISI/notes area below the actual banner image.
                    if i >= len(frame_paths):
                        break
                    adj = _adjust_region_to_banner_aspect(region, banner_w, banner_h)
                    # If aspect adjustment didn't trim the region (i.e. it was
                    # already close-enough), STRETCH the captured frame to
                    # fully fill the detected area — that prevents leftover
                    # edges of the old banner from showing through.
                    fill = (adj["w"] == region["w"] and adj["h"] == region["h"])
                    _paste_frame_into_region(page, adj, str(frame_paths[i]),
                                             anchor_top=True, fill=fill)

        if return_thumbs:
            thumb_url, tw, th = _page_thumb_b64(page)
            results.append({
                "page_num": page_num,
                "thumb":    thumb_url,
                "thumb_w":  tw, "thumb_h": th,
                "regions":  entry["regions"],
                "page_w":   original.size[0],
                "page_h":   original.size[1],
            })
        else:
            full_pages.append(page)

    return results if return_thumbs else full_pages


async def _generate_storyboard_pdf(spec: dict, out_path: Path) -> Path:
    """Render each banner page as a PNG via Playwright, then assemble into a
    multi-page PDF using Pillow.
    """
    from playwright.async_api import async_playwright
    from PIL import Image
    import io

    DPI = 264   # matches the reference storyboard's effective resolution
    PAGE_W = int(11   * DPI)   # 11" landscape
    PAGE_H = int(8.5  * DPI)   #  8.5" landscape

    header  = spec.get("header", "")
    footer  = spec.get("footer", "")
    banners = spec.get("banners", [])
    if not banners:
        raise ValueError("Storyboard spec contains no banners.")

    pages = []
    async with async_playwright() as pw:
        browser = await pw.chromium.launch(headless=True)
        ctx = await browser.new_context(viewport={"width": PAGE_W, "height": PAGE_H},
                                        device_scale_factor=1)
        page = await ctx.new_page()
        for banner in banners:
            html = _storyboard_page_html(banner, header, footer,
                                         banner.get("notes", []), PAGE_W, PAGE_H)
            await page.set_content(html)
            await page.wait_for_timeout(250)
            png = await page.screenshot(full_page=False)
            pages.append(Image.open(io.BytesIO(png)).convert("RGB"))
        await browser.close()

    out_path.parent.mkdir(parents=True, exist_ok=True)
    pages[0].save(str(out_path), "PDF", resolution=DPI,
                  save_all=True, append_images=pages[1:])
    return out_path


# ═══════════════════════════════════════════════════════════════════════════════
# Async capture
# ═══════════════════════════════════════════════════════════════════════════════

async def extract_banner_info(banner_dir: Path) -> dict:
    """Load a banner headlessly and return its ISI text + animation duration."""
    from playwright.async_api import async_playwright
    bw, bh = _get_banner_size(banner_dir / "index.html")
    async with async_playwright() as pw:
        browser = await pw.chromium.launch(headless=True)
        try:
            ctx  = await browser.new_context(viewport={"width": bw + 40, "height": bh + 200})
            page = await ctx.new_page()
            await page.goto((banner_dir / "index.html").as_uri(),
                            wait_until="domcontentloaded", timeout=20_000)
            await page.wait_for_timeout(2_000)
            await page.evaluate(
                "()=>{const c=document.getElementById('container');if(c)c.style.display='block';}"
            )
            data = await page.evaluate("""() => {
                const cnt = document.querySelector('.content');
                const hdr = document.getElementById('header-cont');
                const dur = (typeof tl !== 'undefined') ? tl.duration() : null;
                return {
                    isi_text:    cnt ? (cnt.innerText || '') : '',
                    header_text: hdr ? (hdr.innerText || '') : '',
                    duration:    dur,
                };
            }""")
            await ctx.close()
        finally:
            await browser.close()

    return {
        "width":       bw,
        "height":      bh,
        "isi_text":    _normalize_text(data.get("isi_text", "")),
        "header_text": _normalize_text(data.get("header_text", "")),
        "duration":    data.get("duration"),
    }


async def screenshot_at(banner_dir: Path, t: float) -> bytes:
    """Capture a PNG screenshot of #container at GSAP time `t` seconds."""
    from playwright.async_api import async_playwright
    bw, bh = _get_banner_size(banner_dir / "index.html")
    async with async_playwright() as pw:
        browser = await pw.chromium.launch(headless=True)
        try:
            ctx  = await browser.new_context(
                viewport={"width": bw, "height": bh}, device_scale_factor=3)
            page = await ctx.new_page()
            await page.goto((banner_dir / "index.html").as_uri(),
                            wait_until="domcontentloaded", timeout=20_000)
            await page.wait_for_timeout(2_000)
            await page.evaluate(
                "()=>{const c=document.getElementById('container');if(c)c.style.display='block';}"
            )
            await page.evaluate(
                f"()=>{{if(typeof tl!='undefined')tl.pause({t});"
                f"if(typeof scrollTL!='undefined')scrollTL.pause();}}"
            )
            await page.wait_for_timeout(280)
            data = await page.locator("#container").screenshot()
            await ctx.close()
            return data
        finally:
            await browser.close()


async def _open_page(browser, url, w=None, h=None, scale=3, record_dir=None):
    # Sane defaults if the caller hasn't passed dimensions. Most callers
    # detect the actual size from the loaded HTML and pass it in explicitly.
    if w is None: w = 160
    if h is None: h = 600
    opts = dict(viewport={"width": w, "height": h}, device_scale_factor=scale)
    if record_dir:
        opts["record_video_dir"] = str(record_dir)
        # IMPORTANT: record_video_size MUST be ≤ viewport. If it's larger,
        # Playwright pads the recording with grey instead of upscaling, so
        # the page content sits in the upper-left corner with grey filling
        # the rest. The ffmpeg pass downstream upscales to ≥3× banner size
        # for output crispness — that's where we get our higher resolution.
        opts["record_video_size"] = {"width": w, "height": h}
    ctx  = await browser.new_context(**opts)
    page = await ctx.new_page()
    await page.goto(url, wait_until="domcontentloaded", timeout=20_000)
    await page.wait_for_timeout(2_000)
    await page.evaluate(
        "()=>{const c=document.getElementById('container');if(c)c.style.display='block';}"
    )
    return ctx, page


async def capture_frames(browser, banner_path, output_dir, frames, log):
    """Capture a PNG for each (label, time) pair in `frames`.

    Adapts to banner structure:
      - Klick/Vertex pattern: uses GSAP `tl.pause(t)` + screenshots `#container`.
      - DoubleClick / Enabler pattern: no global timeline; falls back to
        real-time playback (waits the elapsed time between frames) and
        screenshots whatever wrapper element looks banner-shaped.
    """
    ctx, page = await _open_page(browser, banner_path.as_uri())

    # Pick the screenshot target: prefer #container (Klick), then .banner
    # (DoubleClick), then any sized div with the banner-size class, then body.
    target_sel = await page.evaluate("""() => {
        const tries = ['#container', '.banner', 'div.banner',
                       '[class*="banner"][class*="x"]'];
        for (const s of tries) {
            const el = document.querySelector(s);
            if (el && el.offsetWidth > 0) return s;
        }
        return 'body';
    }""")

    # Detect whether we have a controllable timeline (GSAP `tl` global).
    has_tl = await page.evaluate("() => typeof tl !== 'undefined'")

    if has_tl:
        # Klick path — pause the timeline at each frame's exact t, screenshot.
        for name, t in frames:
            await page.evaluate(
                f"()=>{{if(typeof tl!='undefined')tl.pause({t});"
                f"if(typeof scrollTL!='undefined')scrollTL.pause();}}"
            )
            await page.wait_for_timeout(280)
            out = output_dir / f"{name}.png"
            await page.locator(target_sel).screenshot(path=str(out))
            log(f"  ✓ {name}.png  (t={t}s)")
    else:
        # Real-time path — sort frames by t, let the banner play forward, take
        # screenshots at the right moments. Less precise (animations may have
        # already advanced past the requested time) but works for any banner.
        log(f"  (banner has no GSAP `tl` global — using real-time playback)")
        sorted_frames = sorted(frames, key=lambda f: float(f[1]))
        last_t = 0.0
        for name, t in sorted_frames:
            tt   = float(t)
            wait = max(0, int((tt - last_t) * 1000))
            if wait > 0:
                await page.wait_for_timeout(wait)
            last_t = tt
            out = output_dir / f"{name}.png"
            await page.locator(target_sel).screenshot(path=str(out))
            log(f"  ✓ {name}.png  (t={tt}s, real-time)")

    await ctx.close()


async def capture_isi_png_studio(browser, studio_url, banner_dir, output_dir, log):
    """ISI capture for DoubleClick Studio external-preview URLs.

    The bare creative bundle on s0.2mdn.net doesn't render the ISI panel when
    loaded standalone (Enabler/Studio APIs aren't bootstrapped), so isi_full.png
    comes back blank. Instead, load the Studio preview, find the creative inside
    its nested iframe (google.com → doubleclick/preview/main → s0.2mdn.net/pv2/),
    apply the same ISI-expanding DOM mutations there, serialize the frame's full
    document into a self-contained local file (with <base href> injected so
    relative assets/fonts/CSS still resolve), then defer to capture_isi_png on
    that local file — reusing the proven mutation+screenshot logic.
    """
    ctx, page = await _open_page(browser, studio_url, w=1400, h=900, scale=2)
    creative_frame = None
    try:
        try:
            await page.wait_for_selector(
                'iframe[src*="doubleclick/preview/main"]', timeout=30_000)
        except Exception:
            pass

        # Poll for the creative frame (s0.2mdn.net/pv2/<creativeId>/<ts>/index.html).
        # ~30 × 500ms = 15s total.
        for _ in range(30):
            for f in page.frames:
                u = f.url or ""
                if "s0.2mdn.net" in u and "/pv2/" in u and "index.html" in u:
                    creative_frame = f
                    break
            if creative_frame:
                break
            await page.wait_for_timeout(500)

        # Detect the actual banner width from inside the creative frame.
        # The Studio iframe has its own native creative size; we want the ISI
        # panel to match the actual ISI panel — which on a 728×90 leaderboard
        # is ~240px (a side column), not 728. Default to 160 if detection fails.
        bw = 160
        if creative_frame is not None:
            try:
                # Prefer ISI's natural rendered width (matches text-wrap on the
                # real banner). Fall back to the banner's width if no ISI
                # element is found yet.
                isi_w = await creative_frame.evaluate(DETECT_ISI_NATURAL_WIDTH_JS)
                if isi_w and int(isi_w) > 0:
                    bw = int(isi_w)
                else:
                    sz = await creative_frame.evaluate(DETECT_BANNER_SIZE_JS)
                    if sz and sz.get("w"):
                        bw = int(sz["w"])
            except Exception:
                pass

        if not creative_frame:
            log("  (Studio creative frame didn't load — falling back to local ISI capture)")
            await ctx.close()
            await capture_isi_png(browser, banner_dir / "index.html", output_dir, log)
            return

        # Walk-up ISI isolation: instead of hiding everything but `#isi`/`.isi`
        # at one level (which fails when the ISI is nested inside a hidden
        # wrapper like `#main_link[display:none]`), walk up from the ISI to the
        # body and at each level hide siblings while ensuring ancestors are
        # visible. Handles both flat (.banner > #isi) and nested
        # (.banner > #main_link > .isi) layouts.
        await creative_frame.evaluate(f"""() => {{
            const isi = document.getElementById('isi') || document.querySelector('.isi');
            if (!isi) return;
            let cur = isi;
            while (cur && cur.parentElement) {{
                const parent = cur.parentElement;
                for (const sib of Array.from(parent.children)) {{
                    if (sib !== cur) sib.style.display = 'none';
                }}
                if (parent !== document.body && parent !== document.documentElement) {{
                    parent.style.display = 'block';
                    parent.style.visibility = 'visible';
                    parent.style.opacity = '1';
                    parent.style.height = 'auto';
                    parent.style.maxHeight = 'none';
                    parent.style.overflow = 'visible';
                }}
                if (parent === document.body) break;
                cur = parent;
            }}
            document.body.style.cssText =
                'margin:0;padding:0;background:#fff;width:auto;height:auto;display:block;';
            isi.style.cssText = 'position:relative;top:auto;left:auto;right:auto;bottom:auto;'
                + 'width:{bw}px;height:auto;overflow:visible;background:#fff;display:block;'
                + 'visibility:visible;opacity:1;';
            ['.isiScroll', '.isiBody'].forEach(sel => {{
                document.querySelectorAll(sel).forEach(el => {{
                    el.style.height = 'auto';
                    el.style.maxHeight = 'none';
                    el.style.overflow = 'visible';
                    el.style.display = 'block';
                }});
            }});
            document.querySelectorAll('.isiFix, .isiFixBottom').forEach(el => {{
                el.style.position = 'relative';
                el.style.top = el.style.bottom = 'auto';
                el.style.display = 'block';
            }});
        }}""")
        await page.wait_for_timeout(800)

        # Serialize the frame's document with <base href> injected so relative
        # asset URLs still resolve when we reload the file from disk.
        creative_url = creative_frame.url
        base_href = creative_url.rsplit("/", 1)[0] + "/"
        html = await creative_frame.evaluate("() => document.documentElement.outerHTML")
        if "<base " not in html.lower():
            # Inject right after <head>; if no <head> tag, prepend one.
            if re.search(r'<head[^>]*>', html, re.I):
                html = re.sub(r'(<head[^>]*>)',
                              rf'\1<base href="{base_href}">',
                              html, count=1, flags=re.I)
            else:
                html = f'<head><base href="{base_href}"></head>' + html

        out_html = banner_dir / "_isi_studio.html"
        out_html.write_text(html, encoding="utf-8")
    finally:
        await ctx.close()

    # Reuse the existing local-file capture path (doubleclick branch will fire
    # because the serialized markup contains .isiBody / .isiScroll etc.).
    await capture_isi_png(browser, out_html, output_dir, log)


async def capture_isi_png(browser, banner_path, output_dir, log):
    """Export the full ISI as a single tall PNG.

    The ISI is captured at its *own* on-banner width — we read the natural
    widths of #header-cont and #scroller-container at runtime (so 240px on a
    728×90 leaderboard, 160px on a 160×600 skyscraper, etc.) and constrain the
    container to that ISI panel width. Padding-right and inner widths are left
    untouched so text wraps identically to the live banner.
    """
    # Default to a small skyscraper viewport; the detection block below will
    # resize once we've parsed the banner's actual ad.size. This used to read
    # the module-level _banner_w / _banner_h globals which no longer exist —
    # detection covers all real cases (Klick + DoubleClick both set ad.size
    # or have a .banner element).
    bw, bh = 160, 600
    ctx, page = await _open_page(browser, banner_path.as_uri(), w=bw + 400, h=bh + 400, scale=3)
    await page.wait_for_timeout(300)

    # Detect the actual banner width from the loaded HTML so the ISI panel
    # ends up at the right size for the slot we're capturing.
    detected_bw, detected_bh = None, None
    try:
        sz = await page.evaluate(DETECT_BANNER_SIZE_JS)
        if sz and sz.get("w"):
            detected_bw = int(sz["w"])
        if sz and sz.get("h"):
            detected_bh = int(sz["h"])
    except Exception:
        pass
    if detected_bw and detected_bw != bw:
        bw = detected_bw
        # Resize the viewport too so layout engines can use the full natural
        # banner width without horizontal clipping.
        try:
            await page.set_viewport_size(
                {"width": bw + 400, "height": (detected_bh or bh) + 400})
            await page.wait_for_timeout(150)   # let the layout settle
        except Exception:
            pass
    if detected_bh:
        bh = detected_bh

    # Detect which ISI structure this banner uses.
    isi_kind = await page.evaluate("""() => {
        if (document.getElementById('scroller-container')
            || document.getElementById('header-cont'))   return 'klick';
        if (document.getElementById('isi')
            || document.querySelector('.isiBody, .isiScroll')) return 'doubleclick';
        return null;
    }""")
    if not isi_kind:
        log("  (no ISI panel found — skipping isi_full.png)")
        await ctx.close()
        return

    # ── DoubleClick Studio (Enabler) banners ─────────────────────────────────
    if isi_kind == "doubleclick":
        # Detect the ISI's natural rendered width BEFORE mutating — gives us
        # the panel width the ISI was designed for (e.g. ~240px on a 728×90
        # leaderboard, not 728). Using banner_w here would over-stretch the
        # panel and re-wrap text differently than the real banner.
        try:
            isi_natural = await page.evaluate(DETECT_ISI_NATURAL_WIDTH_JS)
            if isi_natural and int(isi_natural) > 0:
                bw = int(isi_natural)
        except Exception:
            pass

        # Walk-up isolation: starting from the ISI element, hide siblings at
        # every ancestor level and ensure ancestors stay visible. Handles both
        # flat layouts (.banner > #isi) and nested ones where the ISI lives
        # inside a wrapper that's `display:none` by default (e.g.
        # .banner > #main_link[display:none] > .isi — the Vertex 300x250
        # creative on Studio).
        await page.evaluate(f"""() => {{
            const isi = document.getElementById('isi') || document.querySelector('.isi');
            if (!isi) return;
            let cur = isi;
            while (cur && cur.parentElement) {{
                const parent = cur.parentElement;
                for (const sib of Array.from(parent.children)) {{
                    if (sib !== cur) sib.style.display = 'none';
                }}
                if (parent !== document.body && parent !== document.documentElement) {{
                    parent.style.display = 'block';
                    parent.style.visibility = 'visible';
                    parent.style.opacity = '1';
                    parent.style.height = 'auto';
                    parent.style.maxHeight = 'none';
                    parent.style.overflow = 'visible';
                }}
                if (parent === document.body) break;
                cur = parent;
            }}
            document.body.style.cssText =
                'margin:0;padding:0;background:#fff;width:auto;height:auto;display:block;';
            isi.style.cssText = 'position:relative;top:auto;left:auto;right:auto;bottom:auto;'
                + 'width:{bw}px;height:auto;overflow:visible;background:#fff;display:block;'
                + 'visibility:visible;opacity:1;';
            ['.isiScroll', '.isiBody'].forEach(sel => {{
                document.querySelectorAll(sel).forEach(el => {{
                    el.style.height   = 'auto';
                    el.style.maxHeight= 'none';
                    el.style.overflow = 'visible';
                    el.style.display  = 'block';
                }});
            }});
            // .isiFix is normally position:fixed; force it into normal flow
            // so it stacks above the body content.
            document.querySelectorAll('.isiFix, .isiFixBottom').forEach(el => {{
                el.style.position = 'relative';
                el.style.top = el.style.bottom = 'auto';
                el.style.display = 'block';
            }});
        }}""")
        await page.wait_for_timeout(800)

        # Screenshot the ISI element directly — its bounding box covers the full
        # natural height (which can exceed the viewport). Screenshotting body
        # would clip at viewport height when the ISI is taller.
        isi_locator = page.locator("#isi").first
        if await isi_locator.count() == 0:
            isi_locator = page.locator(".isi").first
        await isi_locator.screenshot(path=str(output_dir / "isi_full.png"))
        log("  ✓ isi_full.png  (DoubleClick / Enabler ISI)")
        await ctx.close()
        return

    # ── Klick / Vertex ISI scroller (existing flow) ──────────────────────────

    # Detect the natural width of the ISI panel on the banner (before we mutate styles).
    isi_w = await page.evaluate("""() => {
        const sc  = document.getElementById('scroller-container');
        const hdr = document.getElementById('header-cont');
        const w1  = sc  ? sc.offsetWidth  : 0;
        const w2  = hdr ? hdr.offsetWidth : 0;
        return Math.max(w1, w2) || 160;
    }""")

    await page.evaluate(f"""()=>{{
        {_js_hide(ISI_HIDE)}

        // Constrain #container to the ISI panel's actual on-banner width and let
        // it grow vertically. (For a 728×90 banner where the ISI is a 240px
        // column, this prevents 488px of empty whitespace next to the ISI.)
        const c = document.getElementById('container');
        if (c) {{
            c.style.position   = 'relative';
            c.style.width      = '{isi_w}px';
            c.style.height     = 'auto';
            c.style.overflow   = 'visible';
            c.style.background = '#fff';
            c.style.border     = 'none';
            c.style.boxShadow  = 'none';
            c.style.display    = 'block';
        }}

        // Move #header-cont into normal flow at the top — keep its native width.
        const hdr = document.getElementById('header-cont');
        if (hdr) {{
            hdr.style.position = 'relative';
            hdr.style.top      = 'auto';
            hdr.style.right    = 'auto';
            hdr.style.bottom   = 'auto';
            hdr.style.left     = 'auto';
        }}

        // Move #scroller-container into normal flow below the header — keep its native width.
        const sc = document.getElementById('scroller-container');
        if (sc) {{
            sc.style.position = 'relative';
            sc.style.top      = 'auto';
            sc.style.right    = 'auto';
            sc.style.bottom   = 'auto';
            sc.style.left     = 'auto';
            sc.style.height   = 'auto';
            sc.style.overflow = 'visible';
        }}

        // Expand #scroller vertically. Do NOT change its padding-right —
        // that gutter is what makes text wrap the same as on the banner.
        const s = document.getElementById('scroller');
        if (s) {{
            s.style.height   = 'auto';
            s.style.overflow = 'visible';
        }}

        // Hide the scrollbar visuals (the gutter still occupies space → wrap unchanged).
        const bar = document.getElementById('scrollerbar');
        if (bar) bar.style.display = 'none';
    }}""")
    await page.wait_for_timeout(400)
    await page.locator("#container").screenshot(path=str(output_dir / "isi_full.png"))
    log(f"  ✓ isi_full.png  (ISI panel width: {isi_w}px)")
    await ctx.close()


async def capture_isi_html(browser, banner_path, output_dir, log):
    """Export the full ISI as a standalone HTML file with copyable/selectable text.

    Adapts to either the Klick ISI scroller (#header-cont + .content + #footer-cont)
    or the DoubleClick / Enabler ISI structure (.isiFix + .isiBody + .isiFixBottom).
    """
    ctx, page = await _open_page(browser, banner_path.as_uri(), w=400, h=900, scale=1)
    await page.wait_for_timeout(800)
    await page.evaluate("()=>{const c=document.getElementById('container');if(c)c.style.display='block';}")
    await page.wait_for_timeout(300)

    # Pull markup from both possible structures; whichever one is present wins.
    data = await page.evaluate("""()=>{
        const klick_hdr = document.getElementById('header-cont');
        const klick_ftr = document.getElementById('footer-cont');
        const klick_cnt = document.querySelector('.content');
        if (klick_hdr || klick_cnt) {
            return {
                kind:   'klick',
                header:  klick_hdr ? klick_hdr.innerHTML : '',
                content: klick_cnt ? klick_cnt.innerHTML : '',
                footer:  klick_ftr ? klick_ftr.innerHTML : '',
            };
        }
        const dc_hdr  = document.querySelector('.isiFix');
        const dc_body = document.querySelector('.isiBody');
        const dc_ftr  = document.querySelector('.isiFixBottom');
        if (dc_hdr || dc_body) {
            return {
                kind:   'doubleclick',
                header:  dc_hdr  ? dc_hdr.innerHTML  : '',
                content: dc_body ? dc_body.innerHTML : '',
                footer:  dc_ftr  ? dc_ftr.innerHTML  : '',
            };
        }
        return {kind: null, header: '', content: '', footer: ''};
    }""")
    await ctx.close()

    if not (data.get("content") or data.get("header")):
        log("  (no ISI HTML markup found — skipping isi_text.html)")
        return

    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>Full ISI — ALYFTREK</title>
<link rel="preload" href="https://s0.2mdn.net/creatives/assets/5289881/TTCommons-Regular.woff2"
      as="font" type="font/woff2" crossorigin="anonymous">
<link rel="preload" href="https://s0.2mdn.net/creatives/assets/5289881/TTCommons-Bold.woff2"
      as="font" type="font/woff2" crossorigin="anonymous">
<style>
@font-face {{font-family:TTCommonsReg; src:url(https://s0.2mdn.net/creatives/assets/5289881/TTCommons-Regular.woff2);}}
@font-face {{font-family:TTCommonsBold; src:url(https://s0.2mdn.net/creatives/assets/5289881/TTCommons-Bold.woff2);}}
@font-face {{font-family:TTCommonsItalic; src:url(https://s0.2mdn.net/creatives/assets/5289881/TTCommons-Italic.woff2);}}
* {{box-sizing:border-box;}}
body {{
  max-width: 680px; margin: 40px auto; padding: 0 24px 60px;
  font-family: TTCommonsReg, Arial, sans-serif;
  font-size: 13px; line-height: 1.5; color: #333;
  -webkit-font-smoothing: antialiased;
}}
h1 {{font-size:15px; font-family:TTCommonsBold; color:#411B86; margin-bottom:4px;}}
.header-bar {{
  border-top: 3px solid #00A798; padding: 10px 0 6px;
  margin-bottom: 16px;
}}
.header-bar a {{color:#411B86; font-family:TTCommonsReg; font-size:10px; text-decoration:underline;}}
.header-text {{font-family:TTCommonsBold; font-size:13px; color:#411B86; margin-top:4px;}}
p {{padding-bottom:8px; margin:0;}}
ul {{padding:0 0 4px 0; list-style:none;}}
ul li {{
  margin-left:28px; position:relative; padding-bottom:8px;
  list-style-type:none;
}}
ul li:before {{
  content:"\\2022"; position:absolute; left:-20px;
  color:#FF984A; font-size:14px;
}}
table {{border:2px solid #000; width:100%; margin-bottom:12px;}}
th {{padding:8px; text-align:left;}}
th p {{font-size:10px; line-height:1.4; font-family:TTCommonsReg;}}
.boldTextVio {{font-family:TTCommonsBold; color:#411B86;}}
.boldTextTeal {{font-family:TTCommonsBold; color:#00A798;}}
a {{color:#411B86;}}
#isiLogo {{width:120px; height:auto; margin:8px 0 16px 24px;}}
.footer-bar {{
  border-top:1px solid #DBDBDB; margin-top:20px; padding-top:10px;
  font-family:TTCommonsBold; font-size:12px; color:#411B86;
}}
</style>
</head>
<body>
<div class="header-bar">
  {data['header']}
</div>
<div class="content">
  {data['content']}
</div>
<div class="footer-bar">
  {data['footer']}
</div>
</body>
</html>"""

    out = output_dir / "isi_text.html"
    out.write_text(html, encoding="utf-8")
    log("  ✓ isi_text.html  (copyable text — open in browser)")


async def record_video(browser, banner_path, output_dir, log):
    """Record an MP4 of the banner playing in its NATURAL layout — banner
    animation on the left, ISI panel scrolling in its built-in right column,
    both visible simultaneously. Matches how the ad actually looks in-place
    on a webpage (e.g., a 728×90 leaderboard with a 240px ISI side column).

    Single phase: banner auto-plays / loops while the ISI scrolls once at
    238 WPM. Total duration tracks the ISI scroll length so the whole panel
    is read by the time the video ends.
    """
    vtmp = output_dir / "_vtmp"
    vtmp.mkdir(exist_ok=True)

    # Detect the banner's natural size BEFORE we open the recording context
    # — Playwright fixes the recording viewport at context creation, and any
    # later page.set_viewport_size() doesn't resize the recording. So we read
    # the dimensions out of the HTML's <meta name="ad.size"> first (cheap,
    # no browser needed). Falls back to a 160×600 skyscraper if detection
    # fails (it almost never does for real banners).
    try:
        det_w, det_h = _get_banner_size(banner_path)
        bw, bh = (det_w or 160), (det_h or 600)
    except Exception:
        bw, bh = 160, 600

    # Record at scale=1: device_scale_factor>1 makes Playwright render at
    # higher device-pixel density, but record_video_size is interpreted as
    # device pixels too — combining the two crops the recording to the top-
    # left of the rendered viewport. Stick to dsf=1 for the recording; the
    # ffmpeg pass below upscales to 3× the natural banner size for crispness.
    ctx, page = await _open_page(
        browser, banner_path.as_uri(),
        w=bw, h=bh, scale=1, record_dir=vtmp
    )

    # Let the banner mount + start its timeline.
    await page.wait_for_timeout(500)

    # Detect ISI kind so we know which scroller to drive.
    isi_kind = await page.evaluate("""() => {
        if (document.getElementById('scroller')
            && document.getElementById('container')) return 'klick';
        if (document.querySelector('.isiScroll'))    return 'doubleclick';
        return null;
    }""")

    # ── Start ISI scroll WITHOUT hiding the banner ───────────────────────
    # Both panels stay visible together. The banner auto-plays / loops in
    # its built-in column on the left; the ISI scrolls in its built-in
    # column on the right (or wherever the design places it). Total recording
    # duration = ISI scroll time (or 22 s if there's no ISI / very little
    # text), matching how a reviewer would naturally watch the ad.
    if isi_kind == "klick":
        words = await page.evaluate("""() => {
            const c = document.querySelector('.content');
            return c ? c.innerText.trim().split(/\\s+/).length : 500;
        }""")
        scroll_ms = max(22_000, int(words / (238 / 60) * 1000))
        log(f"  Banner + Klick ISI side-by-side · {words} words · "
            f"~{scroll_ms/1000:.0f}s at 238 WPM")
        await page.evaluate(f"""() => {{
            const el = document.getElementById('scroller');
            if (!el) return;
            // Hide the scrollbar visuals (the gutter still occupies space →
            // wrap unchanged) but keep the rest of the layout intact.
            const bar = document.getElementById('scrollerbar');
            if (bar) bar.style.display = 'none';
            el.scrollTop = 0;
            const t0 = performance.now(), dur = {scroll_ms};
            const max = el.scrollHeight - el.clientHeight;
            function step(n) {{
                const p = Math.min((n - t0) / dur, 1);
                el.scrollTop = max * p;
                if (p < 1) requestAnimationFrame(step);
            }}
            requestAnimationFrame(step);
        }}""")
        await page.wait_for_timeout(scroll_ms + 2_000)

    elif isi_kind == "doubleclick":
        words = await page.evaluate("""() => {
            const body = document.querySelector('.isiBody');
            const text = body ? body.innerText.trim() : '';
            return text ? text.split(/\\s+/).length : 500;
        }""")
        scroll_ms = max(22_000, int(words / (238 / 60) * 1000))
        log(f"  Banner + DoubleClick ISI side-by-side · {words} words · "
            f"~{scroll_ms/1000:.0f}s at 238 WPM")
        await page.evaluate(f"""() => {{
            const el = document.querySelector('.isiScroll');
            if (!el) return;
            el.scrollTop = 0;
            const t0 = performance.now(), dur = {scroll_ms};
            const max = el.scrollHeight - el.clientHeight;
            function step(n) {{
                const p = Math.min((n - t0) / dur, 1);
                el.scrollTop = max * p;
                if (p < 1) requestAnimationFrame(step);
            }}
            requestAnimationFrame(step);
        }}""")
        await page.wait_for_timeout(scroll_ms + 2_000)

    else:
        log("  (no ISI scroller found — recording banner animation only, ~22s)")
        await page.wait_for_timeout(22_000)

    webm_path = str(await page.video.path())
    await ctx.close()

    # High-quality encode: video is now recorded at 2× viewport via Playwright,
    # ffmpeg keeps that resolution and upscales to 3× native banner size with
    # the lanczos filter (sharpest standard scaler) and CRF 8 (visually lossless).
    mp4_out = output_dir / "banner_animation.mp4"
    ffmpeg  = imageio_ffmpeg.get_ffmpeg_exe()
    out_w, out_h = bw * 3, bh * 3
    result  = subprocess.run(
        [ffmpeg, "-y", "-i", webm_path,
         "-vf", f"scale={out_w}:{out_h}:flags=lanczos,setsar=1",
         "-c:v", "libx264", "-pix_fmt", "yuv420p",
         "-preset", "slow", "-crf", "8", "-r", "30", str(mp4_out)],
        capture_output=True, text=True
    )
    log("  ✓ banner_animation.mp4" if result.returncode == 0
        else f"  ERROR (ffmpeg):\n{result.stderr[-300:]}")

    try: Path(webm_path).unlink()
    except: pass
    try: vtmp.rmdir()
    except: pass


async def run_all(banner_path, output_dir, frames,
                  do_frames, do_video, log,
                  do_isi_png=None, do_isi_html=None,
                  studio_url=None):
    """Run the capture pipeline. The four booleans are independent.

    For backward-compat, if `do_isi_png` / `do_isi_html` are left as None
    they fall back to `do_frames` (which is how page 1's "PNG frames + ISI"
    checkbox was historically wired — frames + both ISI exports together).

    `studio_url`, when supplied, routes ISI PNG capture through the live
    DoubleClick Studio preview instead of the local creative bundle (the bare
    s0.2mdn.net creative doesn't render its ISI standalone).
    """
    if do_isi_png  is None: do_isi_png  = do_frames
    if do_isi_html is None: do_isi_html = do_frames
    async with async_playwright() as pw:
        browser = await pw.chromium.launch(headless=True)
        try:
            if do_frames and frames:
                log("── PNG Frames ────────────────────────")
                await capture_frames(browser, banner_path, output_dir, frames, log)
            if do_isi_png:
                log("── ISI Image ─────────────────────────")
                if studio_url:
                    await capture_isi_png_studio(
                        browser, studio_url, banner_path.parent, output_dir, log)
                else:
                    await capture_isi_png(browser, banner_path, output_dir, log)
            if do_isi_html:
                log("── ISI HTML ──────────────────────────")
                await capture_isi_html(browser, banner_path, output_dir, log)
            if do_video:
                log("── MP4 Video ─────────────────────────")
                await record_video(browser, banner_path, output_dir, log)
        finally:
            await browser.close()


# ═══════════════════════════════════════════════════════════════════════════════
# Script injected into the banner iframe for live preview controls
# ═══════════════════════════════════════════════════════════════════════════════

BANNER_INJECT = """<script>
(function(){
  function init(){
    if(typeof tl==='undefined'){setTimeout(init,200);return;}
    // Tell parent the total duration so it can size the scrubber
    try{ window.parent.postMessage({bc:'ready',duration:tl.duration()},'*'); }catch(e){}
    window.addEventListener('message',function(e){
      if(!e.data||!e.data.bc)return;
      var d=e.data;
      if(d.bc==='play')   tl.play();
      if(d.bc==='pause')  {tl.pause();if(typeof scrollTL!=='undefined')scrollTL.pause();}
      if(d.bc==='restart'){tl.restart();if(typeof scrollTL!=='undefined'){scrollTL.kill();scrollTL=gsap.timeline({paused:true});}}
      if(d.bc==='seek')   {tl.pause(d.time);if(typeof scrollTL!=='undefined')scrollTL.pause();}
    });
    setInterval(function(){
      try{
        window.parent.postMessage({
          bc:'state',
          time:    tl.time(),
          duration:tl.duration(),
          paused:  tl.paused()
        },'*');
      }catch(e){}
    },80);
  }
  init();
})();
</script>"""


# ═══════════════════════════════════════════════════════════════════════════════
# HTML UI
# ═══════════════════════════════════════════════════════════════════════════════

HTML = r"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>Banner Tool</title>
<style>
/* ─── Theme tokens ───────────────────────────────────────────────────────
 * Each combo of {scheme, mode} declares the same set of variables, so the
 * rest of the CSS just reads `var(--bg-page)` etc. The JS ThemeManager
 * resolves "auto" mode to a concrete dark/light value at runtime via
 * matchMedia, so CSS only needs concrete (scheme, mode) blocks below.
 *
 * Schemes: catppuccin (default) | tokyo | nord | solarized
 * Modes:   dark | light
 *
 * To toggle:
 *   document.documentElement.dataset.scheme = 'tokyo';
 *   document.documentElement.dataset.theme  = 'light';
 */

/* ── Catppuccin (Mocha / Latte) — DEFAULT ────────────────────────────── */
:root,
:root[data-scheme="catppuccin"][data-theme="dark"]{
  --bg-page:        #1e1e2e;
  --bg-topbar:      #181825;
  --bg-card:        #313244;
  --bg-elevated:    #45475a;
  --bg-input:       #181825;
  --bg-hover:       #45475a;
  --border-default: #45475a;
  --border-strong:  #6c7086;
  --border-soft:    #313244;
  --text-primary:   #cdd6f4;
  --text-secondary: #a6adc8;
  --text-muted:     #6c7086;
  --text-inverse:   #1e1e2e;
  --text-on-accent: #1e1e2e;
  --accent-primary: #cba6f7;   /* mauve */
  --accent-success: #a6e3a1;   /* green */
  --accent-warn:    #f9e2af;   /* yellow */
  --accent-error:   #f38ba8;   /* red */
  --accent-info:    #89dceb;   /* sky */
  --accent-pink:    #f5c2e7;
  --shadow-card:    0 4px 12px rgba(0,0,0,.4);
  --overlay-dim:    rgba(17,17,27,.72);
  color-scheme: dark;
}
:root[data-scheme="catppuccin"][data-theme="light"]{
  --bg-page:        #eff1f5;
  --bg-topbar:      #e6e9ef;
  --bg-card:        #ffffff;
  --bg-elevated:    #ccd0da;
  --bg-input:       #ffffff;
  --bg-hover:       #dce0e8;
  --border-default: #bcc0cc;
  --border-strong:  #9ca0b0;
  --border-soft:    #ccd0da;
  --text-primary:   #4c4f69;
  --text-secondary: #6c6f85;
  --text-muted:     #9ca0b0;
  --text-inverse:   #ffffff;
  --text-on-accent: #ffffff;
  --accent-primary: #8839ef;
  --accent-success: #40a02b;
  --accent-warn:    #df8e1d;
  --accent-error:   #d20f39;
  --accent-info:    #04a5e5;
  --accent-pink:    #ea76cb;
  --shadow-card:    0 4px 12px rgba(76,79,105,.12);
  --overlay-dim:    rgba(76,79,105,.45);
  color-scheme: light;
}

/* ── Tokyo Night ─────────────────────────────────────────────────────── */
:root[data-scheme="tokyo"][data-theme="dark"]{
  --bg-page:        #1a1b26;
  --bg-topbar:      #16161e;
  --bg-card:        #24283b;
  --bg-elevated:    #2f334d;
  --bg-input:       #16161e;
  --bg-hover:       #292e42;
  --border-default: #292e42;
  --border-strong:  #3b4261;
  --border-soft:    #1f2335;
  --text-primary:   #c0caf5;
  --text-secondary: #9aa5ce;
  --text-muted:     #565f89;
  --text-inverse:   #1a1b26;
  --text-on-accent: #1a1b26;
  --accent-primary: #bb9af7;
  --accent-success: #9ece6a;
  --accent-warn:    #e0af68;
  --accent-error:   #f7768e;
  --accent-info:    #7aa2f7;
  --accent-pink:    #ff7eb6;
  --shadow-card:    0 4px 12px rgba(0,0,0,.5);
  --overlay-dim:    rgba(15,16,28,.78);
  color-scheme: dark;
}
:root[data-scheme="tokyo"][data-theme="light"]{
  --bg-page:        #e1e2e7;
  --bg-topbar:      #d5d6db;
  --bg-card:        #ffffff;
  --bg-elevated:    #cbcdda;
  --bg-input:       #ffffff;
  --bg-hover:       #c4c8da;
  --border-default: #b6b9c8;
  --border-strong:  #8990b3;
  --border-soft:    #cbcdda;
  --text-primary:   #3760bf;
  --text-secondary: #6172b0;
  --text-muted:     #8990b3;
  --text-inverse:   #ffffff;
  --text-on-accent: #ffffff;
  --accent-primary: #9854f1;
  --accent-success: #587539;
  --accent-warn:    #8c6c3e;
  --accent-error:   #f52a65;
  --accent-info:    #2e7de9;
  --accent-pink:    #d20065;
  --shadow-card:    0 4px 12px rgba(55,96,191,.12);
  --overlay-dim:    rgba(55,96,191,.45);
  color-scheme: light;
}

/* ── Nord ────────────────────────────────────────────────────────────── */
:root[data-scheme="nord"][data-theme="dark"]{
  --bg-page:        #2e3440;
  --bg-topbar:      #242933;
  --bg-card:        #3b4252;
  --bg-elevated:    #434c5e;
  --bg-input:       #2e3440;
  --bg-hover:       #4c566a;
  --border-default: #4c566a;
  --border-strong:  #6c7080;
  --border-soft:    #3b4252;
  --text-primary:   #eceff4;
  --text-secondary: #d8dee9;
  --text-muted:     #6c7080;
  --text-inverse:   #2e3440;
  --text-on-accent: #2e3440;
  --accent-primary: #88c0d0;   /* frost cyan */
  --accent-success: #a3be8c;   /* aurora green */
  --accent-warn:    #ebcb8b;
  --accent-error:   #bf616a;
  --accent-info:    #81a1c1;
  --accent-pink:    #b48ead;
  --shadow-card:    0 4px 12px rgba(0,0,0,.4);
  --overlay-dim:    rgba(46,52,64,.78);
  color-scheme: dark;
}
:root[data-scheme="nord"][data-theme="light"]{
  --bg-page:        #eceff4;
  --bg-topbar:      #ffffff;
  --bg-card:        #ffffff;
  --bg-elevated:    #e5e9f0;
  --bg-input:       #ffffff;
  --bg-hover:       #d8dee9;
  --border-default: #d8dee9;
  --border-strong:  #b6bdca;
  --border-soft:    #e5e9f0;
  --text-primary:   #2e3440;
  --text-secondary: #4c566a;
  --text-muted:     #88909d;
  --text-inverse:   #ffffff;
  --text-on-accent: #ffffff;
  --accent-primary: #5e81ac;   /* frost deep blue */
  --accent-success: #587539;
  --accent-warn:    #b58900;
  --accent-error:   #bf616a;
  --accent-info:    #81a1c1;
  --accent-pink:    #b48ead;
  --shadow-card:    0 4px 12px rgba(46,52,64,.12);
  --overlay-dim:    rgba(46,52,64,.45);
  color-scheme: light;
}

/* ── Solarized (dark = base03; light = base3) ────────────────────────── */
:root[data-scheme="solarized"][data-theme="dark"]{
  --bg-page:        #002b36;
  --bg-topbar:      #00212b;
  --bg-card:        #073642;
  --bg-elevated:    #0d4858;
  --bg-input:       #073642;
  --bg-hover:       #0d4858;
  --border-default: #0d4858;
  --border-strong:  #586e75;
  --border-soft:    #073642;
  --text-primary:   #eee8d5;
  --text-secondary: #93a1a1;
  --text-muted:     #586e75;
  --text-inverse:   #002b36;
  --text-on-accent: #002b36;
  --accent-primary: #6c71c4;   /* violet */
  --accent-success: #859900;   /* green */
  --accent-warn:    #b58900;   /* yellow */
  --accent-error:   #dc322f;   /* red */
  --accent-info:    #268bd2;   /* blue */
  --accent-pink:    #d33682;   /* magenta */
  --shadow-card:    0 4px 12px rgba(0,0,0,.45);
  --overlay-dim:    rgba(0,21,28,.78);
  color-scheme: dark;
}
:root[data-scheme="solarized"][data-theme="light"]{
  --bg-page:        #fdf6e3;
  --bg-topbar:      #f5efd9;
  --bg-card:        #eee8d5;
  --bg-elevated:    #e7e0c5;
  --bg-input:       #fdf6e3;
  --bg-hover:       #e7e0c5;
  --border-default: #d9d2bd;
  --border-strong:  #93a1a1;
  --border-soft:    #eee8d5;
  --text-primary:   #586e75;
  --text-secondary: #657b83;
  --text-muted:     #93a1a1;
  --text-inverse:   #fdf6e3;
  --text-on-accent: #fdf6e3;
  --accent-primary: #6c71c4;
  --accent-success: #859900;
  --accent-warn:    #b58900;
  --accent-error:   #dc322f;
  --accent-info:    #268bd2;
  --accent-pink:    #d33682;
  --shadow-card:    0 4px 12px rgba(88,110,117,.12);
  --overlay-dim:    rgba(88,110,117,.45);
  color-scheme: light;
}
*{box-sizing:border-box;margin:0;padding:0;}

/* ── Page shell ── */
body{
  font-family:-apple-system,"Helvetica Neue",sans-serif;
  background:var(--bg-page);color:var(--text-primary);
  min-height:100vh;padding:0 0 60px;
  transition:background-color .2s ease, color .2s ease;
}

/* Top bar */
.topbar{
  background:var(--bg-topbar);border-bottom:1px solid var(--border-soft);
  padding:14px 24px;display:flex;align-items:center;gap:18px;
  position:sticky;top:0;z-index:10;
}
.topbar h1{font-size:17px;font-weight:700;letter-spacing:-.01em;flex-shrink:0;}
.topbar .sub{color:var(--text-muted);font-size:12px;margin-left:auto;}

/* Theme toggle (3-state: Light / Auto / Dark) */
.theme-toggle{
  display:inline-flex;background:var(--bg-input);
  border:1px solid var(--border-default);border-radius:6px;
  overflow:hidden;flex-shrink:0;
}
.theme-toggle button{
  background:transparent;border:none;color:var(--text-secondary);
  cursor:pointer;padding:5px 9px;font-size:13px;font-family:inherit;
  line-height:1;transition:background .15s ease, color .15s ease;
}
.theme-toggle button:hover{background:var(--bg-hover);color:var(--text-primary);}
.theme-toggle button.active{
  background:var(--accent-primary);color:var(--text-on-accent);
}

/* Color-scheme picker (button + dropdown menu) */
.scheme-picker{position:relative;flex-shrink:0;}
.scheme-picker > button{
  background:var(--bg-input);border:1px solid var(--border-default);
  color:var(--text-primary);cursor:pointer;
  padding:5px 10px;font-size:12px;font-family:inherit;border-radius:6px;
  display:inline-flex;align-items:center;gap:6px;line-height:1;
  transition:background .15s ease, border-color .15s ease;
}
.scheme-picker > button:hover{
  background:var(--bg-hover);border-color:var(--border-strong);
}
.scheme-picker > button .chev{font-size:10px;color:var(--text-secondary);}
.scheme-menu{
  position:absolute;top:calc(100% + 4px);right:0;
  background:var(--bg-card);border:1px solid var(--border-default);
  border-radius:8px;box-shadow:var(--shadow-card);
  min-width:180px;padding:4px;z-index:1000;display:none;
}
.scheme-menu.open{display:block;}
.scheme-menu button{
  display:flex;align-items:center;gap:10px;width:100%;
  background:transparent;border:none;cursor:pointer;
  padding:8px 10px;font-size:12px;font-family:inherit;color:var(--text-primary);
  text-align:left;border-radius:5px;transition:background .12s ease;
}
.scheme-menu button:hover{background:var(--bg-hover);}
.scheme-menu button .swatch{
  width:14px;height:14px;border-radius:3px;flex-shrink:0;
  border:1px solid var(--border-default);
  /* Three diagonal stripes show off the scheme's palette at a glance */
  background:linear-gradient(135deg,
    var(--swatch-1) 0 33%,
    var(--swatch-2) 33% 66%,
    var(--swatch-3) 66% 100%);
}
.scheme-menu button.active{background:var(--bg-elevated);font-weight:600;}
.scheme-menu button.active::after{content:"✓";margin-left:auto;color:var(--accent-success);}

/* Tab nav */
.tabs{display:flex;gap:4px;}
.tab-btn{
  background:transparent;border:none;color:var(--text-secondary);
  padding:8px 14px;border-radius:7px;font-size:13px;font-family:inherit;
  cursor:pointer;transition:background .15s,color .15s;
}
.tab-btn:hover{background:var(--bg-hover);color:var(--text-primary);}
.tab-btn.active{background:var(--border-default);color:#fff;font-weight:600;}

/* Tab panes — only one visible at a time */
.tab-pane{display:none;}
.tab-pane.active{display:block;}

/* Main two-column layout */
.layout{
  display:grid;
  grid-template-columns:340px 1fr;
  grid-template-rows:auto;
  gap:0;
  min-height:calc(100vh - 57px);
}

/* Left sidebar */
.sidebar{
  background:var(--bg-hover);border-right:1px solid var(--border-soft);
  padding:20px;display:flex;flex-direction:column;gap:16px;
  overflow-y:auto;
}

/* Right main area */
.main{
  padding:20px;display:flex;flex-direction:column;gap:16px;
  overflow-y:auto;
}

/* Cards */
.card{
  background:var(--bg-elevated);border:1px solid var(--border-default);
  border-radius:10px;padding:18px;
}
.card-title{
  font-size:10px;font-weight:700;letter-spacing:.12em;text-transform:uppercase;
  color:var(--text-muted);margin-bottom:12px;
}

/* Form elements */
.row{display:flex;gap:8px;align-items:center;flex-wrap:wrap;}
.btn{
  flex-shrink:0;padding:8px 14px;border-radius:7px;border:none;cursor:pointer;
  font-size:13px;font-family:inherit;white-space:nowrap;transition:background .15s,opacity .15s;
}
.btn:disabled{opacity:.3;cursor:default;}
.btn-p{background:var(--accent-primary);color:#fff;font-weight:700;}
.btn-p:not(:disabled):hover{background:#6a4dad;}
.btn-g{background:var(--border-default);color:var(--text-primary);}
.btn-g:not(:disabled):hover{background:var(--bg-hover);}
.btn-sm{font-size:11px;padding:5px 10px;}
.btn-full{width:100%;text-align:center;padding:10px;}
input[type=text],input[type=number],input[type=password]{
  flex:1;min-width:0;padding:8px 10px;border-radius:7px;
  border:1.5px solid var(--border-default);background:var(--bg-input);
  color:var(--text-primary);font-family:"Menlo",monospace;font-size:12px;outline:none;
  transition:border-color .2s;
}
input:focus{border-color:var(--accent-primary);}
input.ok{border-color:var(--accent-success);color:var(--accent-success);}
.hint{font-size:11px;color:var(--text-muted);margin-top:8px;line-height:1.6;}
.src-label{font-size:11px;color:var(--text-secondary);margin:6px 0 4px;letter-spacing:.04em;}
.src-divider{font-size:10px;color:var(--border-default);text-align:center;margin:10px 0 4px;letter-spacing:.15em;}
#pstatus{font-size:12px;min-height:16px;margin-top:8px;}
#pstatus.ok{color:var(--accent-success);} #pstatus.err{color:var(--accent-error);}
input[type=checkbox]{width:14px;height:14px;accent-color:var(--accent-primary);cursor:pointer;flex-shrink:0;}
label{display:flex;align-items:center;gap:7px;cursor:pointer;font-size:13px;}

/* ── Preview area ── */
#preview-panel{display:none;flex-direction:column;gap:14px;}
.preview-wrap{
  display:flex;gap:16px;align-items:flex-start;flex-wrap:wrap;
}
.banner-shell{
  flex-shrink:0;background:#000;border-radius:6px;
  border:1px solid var(--bg-hover);overflow:hidden;
}
.banner-shell iframe{border:none;display:block;pointer-events:none;}

/* Playback controls */
.controls{flex:1;min-width:220px;display:flex;flex-direction:column;gap:10px;}
.time-display{
  font-family:"Menlo",monospace;font-size:22px;font-weight:700;
  color:var(--accent-success);letter-spacing:.04em;
  /* Steady-width fix: tabular-nums forces equal-width digit columns; min-width
   * reserves space for the longest reasonable value ("t = 999.99 s"); left-
   * align anchors the digits so subsequent buttons don't shift as the counter
   * widens from "9.99" to "10.00" to "100.00". white-space:nowrap prevents
   * an aggressive flex-shrink from breaking the label across lines. */
  font-variant-numeric:tabular-nums;
  min-width:9ch;
  text-align:left;
  white-space:nowrap;
}
.pb-row{display:flex;gap:6px;flex-wrap:wrap;}
#scrubber{
  width:100%;height:5px;cursor:pointer;accent-color:var(--accent-primary);border-radius:3px;
  margin-top:2px;
}

/* Frame list */
.frame-list{display:flex;flex-direction:column;gap:5px;}
.frame-item{
  display:flex;align-items:center;gap:6px;
  background:var(--bg-page);border-radius:6px;padding:5px 8px;
  font-family:"Menlo",monospace;font-size:11px;
}
.frame-item input[type=text]{font-size:11px;padding:4px 7px;}
.frame-item input[type=number]{width:64px;flex:none;font-size:11px;padding:4px 7px;}
.frame-del{
  margin-left:auto;flex-shrink:0;background:none;border:none;
  color:var(--text-muted);cursor:pointer;font-size:14px;padding:0 3px;line-height:1;
}
.frame-del:hover{color:var(--accent-error);}
.add-frame-btn{
  align-self:flex-start;background:#1a2e22;color:var(--accent-success);
  border:1px solid #2a5a3a;border-radius:6px;padding:5px 11px;
  font-size:12px;cursor:pointer;transition:background .15s;
}
.add-frame-btn:hover{background:#1e3a2a;}

/* ── Export & log (sidebar) ── */
.export-opts{display:flex;flex-direction:column;gap:10px;}
.check-row{display:flex;flex-direction:column;gap:8px;}
#log{
  background:var(--bg-input);border-radius:7px;padding:12px;
  font-family:"Menlo",monospace;font-size:11px;color:var(--accent-success);
  height:180px;overflow-y:auto;white-space:pre-wrap;word-break:break-all;
  margin-top:4px;line-height:1.5;
}
#runstatus{margin-top:6px;font-size:12px;color:var(--text-secondary);text-align:center;}

/* ── Responsive: collapse to single column ── */
@media(max-width:720px){
  .layout{grid-template-columns:1fr;}
  .sidebar{border-right:none;border-bottom:1px solid var(--border-soft);}
  .main{padding:16px;}
}

/* ═══════ ISI Compare page ═══════ */
.compare-wrap{padding:24px;max-width:1400px;margin:0 auto;}
.compare-toolbar{
  display:flex;gap:10px;align-items:center;flex-wrap:wrap;
  background:var(--bg-elevated);border:1px solid var(--border-default);border-radius:10px;
  padding:14px 18px;margin-bottom:16px;
}
.compare-pill{
  font-size:11px;color:var(--text-secondary);padding:5px 10px;background:var(--bg-hover);
  border-radius:6px;font-family:"Menlo",monospace;
}
.compare-pill.ok{color:var(--accent-success);}
.compare-pill.warn{color:#f9c97c;}
.compare-grid{
  display:grid;grid-template-columns:1fr 1fr;gap:14px;
}
.compare-col{
  background:var(--bg-elevated);border:1px solid var(--border-default);border-radius:10px;
  padding:16px;min-height:200px;
}
.compare-col h3{
  font-size:11px;color:var(--text-secondary);text-transform:uppercase;letter-spacing:.12em;
  margin-bottom:10px;font-weight:600;
}
.diff-text{
  font-family:"Menlo",monospace;font-size:11px;line-height:1.6;
  white-space:pre-wrap;word-break:break-word;color:var(--text-primary);
  max-height:60vh;overflow-y:auto;padding:8px;
  background:var(--bg-input);border-radius:6px;
}
.diff-source{
  width:100%;min-height:220px;max-height:400px;
  font-family:"Menlo",monospace;font-size:11px;line-height:1.6;
  color:var(--text-primary);background:var(--bg-input);
  border:1.5px solid var(--border-default);border-radius:6px;padding:8px;
  resize:vertical;outline:none;
}
.diff-source:focus{border-color:var(--accent-primary);}
.isi-sel-info{
  font-size:10px;color:var(--accent-success);font-family:"Menlo",monospace;
  margin-left:8px;font-weight:400;text-transform:none;letter-spacing:0;
}
.diff-text del{
  background:#43222a;color:var(--accent-error);text-decoration:line-through;
  padding:0 2px;border-radius:2px;
}
.diff-text ins{
  background:#1e3d2e;color:var(--accent-success);text-decoration:none;
  padding:0 2px;border-radius:2px;
}
.diff-summary{
  margin-bottom:14px;padding:12px 16px;border-radius:8px;
  font-size:13px;
}
.diff-summary.match{background:#1e3a2a;color:var(--accent-success);border:1px solid #2a5a3a;}
.diff-summary.diff{background:#3a1e1e;color:var(--accent-error);border:1px solid #6a3a3a;}

/* ═══════ Banner Compare page ═══════ */
.bc-toolbar{
  display:grid;grid-template-columns:1fr 1fr;gap:14px;margin-bottom:16px;
}
.bc-loader{
  background:var(--bg-elevated);border:1px solid var(--border-default);border-radius:10px;padding:14px;
}
.bc-loader h3{
  font-size:11px;color:var(--text-secondary);text-transform:uppercase;letter-spacing:.12em;
  margin-bottom:10px;font-weight:600;
}
.bc-loader .row{margin-bottom:8px;}
.bc-status{font-size:12px;min-height:16px;color:var(--text-secondary);}
.bc-status.ok{color:var(--accent-success);}
.bc-status.err{color:var(--accent-error);}

.bc-master{
  background:var(--bg-elevated);border:1px solid var(--border-default);border-radius:10px;
  padding:14px 18px;display:flex;align-items:center;gap:14px;flex-wrap:wrap;
  margin-bottom:16px;
}
.bc-master .time-display{font-size:18px;flex-shrink:0;}
.bc-master input[type=range]{flex:1;min-width:200px;height:5px;accent-color:var(--accent-primary);}
.bc-link{
  display:flex;align-items:center;gap:6px;font-size:12px;color:var(--accent-success);
  background:#1e3a2a;padding:6px 12px;border-radius:7px;border:1px solid #2a5a3a;
  cursor:pointer;
}
.bc-link.unlinked{color:#f9c97c;background:#3a2e1e;border-color:#5a4a2a;}

.bc-views{
  display:grid;grid-template-columns:1fr 1fr;gap:16px;margin-bottom:16px;
}
.bc-view{
  background:var(--bg-elevated);border:1px solid var(--border-default);border-radius:10px;
  padding:14px;display:flex;flex-direction:column;gap:10px;align-items:center;
}
.bc-view-title{
  align-self:stretch;font-size:11px;color:var(--text-secondary);text-transform:uppercase;
  letter-spacing:.12em;font-weight:600;
}
.bc-shell{
  flex-shrink:0;background:#000;border-radius:6px;
  border:1px solid var(--bg-hover);overflow:hidden;
}
.bc-shell iframe{border:none;display:block;pointer-events:none;}
.bc-indep-controls{display:none;width:100%;flex-direction:column;gap:6px;}
.bc-indep-controls.show{display:flex;}
.bc-indep-controls input[type=range]{width:100%;height:4px;accent-color:var(--accent-primary);}

.bc-diff-section{
  background:var(--bg-elevated);border:1px solid var(--border-default);border-radius:10px;
  padding:18px;
}
.bc-diff-section h3{font-size:13px;font-weight:700;margin-bottom:12px;}
.bc-meta{
  display:grid;grid-template-columns:1fr 1fr;gap:8px;font-size:12px;
  font-family:"Menlo",monospace;color:var(--text-primary);margin-bottom:14px;
}
.bc-meta-row{
  display:grid;grid-template-columns:auto 1fr 1fr;gap:14px;
  padding:6px 10px;background:var(--bg-hover);border-radius:6px;align-items:center;
}
.bc-meta-row.different{background:#3a1e1e;color:var(--accent-error);}
.bc-meta-label{color:var(--text-secondary);font-size:11px;}

.bc-vis{
  display:grid;grid-template-columns:1fr 1fr 1fr;gap:10px;margin-top:12px;
}
.bc-vis figure{margin:0;text-align:center;}
.bc-vis figcaption{font-size:10px;color:var(--text-secondary);margin-top:4px;}
.bc-vis img{
  max-width:100%;border:1px solid var(--border-default);border-radius:4px;background:#fff;
}

/* ═══════ Multi-Banner Viewer page ═══════ */
.mb-loader{
  background:var(--bg-elevated);border:1px solid var(--border-default);border-radius:10px;
  padding:14px 18px;margin-bottom:16px;
}
.mb-loader .row{align-items:center;gap:10px;flex-wrap:wrap;}
.mb-loader label{
  display:flex;flex-direction:column;font-size:11px;color:var(--text-secondary);
  letter-spacing:.04em;gap:4px;flex:1;min-width:140px;
}
.mb-loader label input{font-family:"Menlo",monospace;}
.mb-loader .hint{margin-top:6px;}

.mb-views{
  display:flex;flex-wrap:wrap;gap:18px;justify-content:flex-start;
  margin-top:6px;
}
.mb-tile{
  background:var(--bg-elevated);border:1px solid var(--border-default);border-radius:10px;
  padding:14px;display:flex;flex-direction:column;gap:6px;align-items:center;
}
.mb-tile h4{
  font-size:11px;color:var(--text-secondary);text-transform:uppercase;letter-spacing:.12em;
  margin:0;align-self:stretch;font-weight:600;
}
.mb-shell{
  background:#000;border:1px solid var(--bg-hover);border-radius:6px;overflow:hidden;
  flex-shrink:0;
}
.mb-shell iframe{border:none;display:block;pointer-events:none;}
.mb-tile .mb-t{font-size:11px;color:var(--accent-success);font-family:"Menlo",monospace;}

.mb-master{
  background:var(--bg-elevated);border:1px solid var(--border-default);border-radius:10px;
  padding:14px 18px;display:flex;align-items:center;gap:10px;
  margin-bottom:16px;
  /* Was flex-wrap:wrap with gap:14 + scrubber min-width:240 — long labels
   * spilled onto a second row. Now: stay on one line; if the viewport is
   * too narrow to fit, scroll horizontally inside the bar instead of
   * wrapping (avoids the controls jumping rows as the scrubber updates). */
  flex-wrap:nowrap;
  overflow-x:auto;
}
.mb-master > *{flex-shrink:0;}            /* buttons keep their natural width */
.mb-master input[type=range]{
  flex:1 1 auto;min-width:140px;          /* scrubber shrinks as needed */
  height:5px;accent-color:var(--accent-primary);
}

/* ═══════ Storyboard page ═══════ */
.story-toolbar{
  background:var(--bg-elevated);border:1px solid var(--border-default);border-radius:10px;
  padding:14px 18px;margin-bottom:16px;
  display:flex;flex-wrap:wrap;gap:10px;align-items:end;
}
.story-toolbar label{
  display:flex;flex-direction:column;font-size:11px;color:var(--text-secondary);
  letter-spacing:.04em;gap:4px;flex:1;min-width:240px;
}
.story-toolbar label input{font-family:"Menlo",monospace;}

.story-banner{
  background:var(--bg-elevated);border:1px solid var(--border-default);border-radius:10px;
  padding:16px 18px;margin-bottom:14px;
}
.story-banner h3{
  font-size:14px;font-weight:700;color:var(--text-primary);margin:0 0 12px;
  display:flex;align-items:center;gap:10px;
}
.story-banner h3 .mute{font-size:11px;color:var(--text-secondary);font-weight:400;}
.story-frames-grid{
  display:grid;grid-template-columns:repeat(4,1fr);gap:14px;
}
.story-frame{display:flex;flex-direction:column;gap:6px;}
.story-frame .label{
  font-size:11px;color:var(--text-secondary);letter-spacing:.06em;text-transform:uppercase;
  font-weight:600;
}
.story-frame img{
  background:#fff;border:1px solid var(--bg-hover);border-radius:4px;
  max-width:100%;max-height:280px;object-fit:contain;align-self:flex-start;
}
.story-frame textarea{
  width:100%;min-height:90px;
  font-family:-apple-system,sans-serif;font-size:12px;line-height:1.5;
  background:var(--bg-input);color:var(--text-primary);
  border:1.5px solid var(--border-default);border-radius:6px;padding:8px;
  resize:vertical;outline:none;
}
.story-frame textarea:focus{border-color:var(--accent-primary);}
.story-actions{
  background:var(--bg-elevated);border:1px solid var(--border-default);border-radius:10px;
  padding:14px 18px;display:flex;gap:10px;align-items:center;flex-wrap:wrap;
  margin-top:14px;
}
.story-status{font-size:12px;color:var(--text-secondary);margin-left:auto;}
.story-status.ok{color:var(--accent-success);}
.story-status.err{color:var(--accent-error);}

@media(max-width:1100px){.story-frames-grid{grid-template-columns:repeat(2,1fr);}}
@media(max-width:600px){.story-frames-grid{grid-template-columns:1fr;}}

/* Mode toggle (Generate vs Update existing) */
.story-mode{
  display:inline-flex;background:var(--bg-hover);border:1px solid var(--border-default);
  border-radius:8px;padding:3px;margin-bottom:14px;
}
.story-mode .mode-btn{
  background:transparent;border:none;color:var(--text-secondary);
  padding:7px 16px;border-radius:6px;cursor:pointer;
  font-family:inherit;font-size:13px;
}
.story-mode .mode-btn.active{background:var(--border-default);color:#fff;font-weight:600;}

/* Replace-mode page rows */
.replace-page{
  display:flex;gap:18px;background:var(--bg-elevated);border:1px solid var(--border-default);
  border-radius:10px;padding:14px;margin-bottom:12px;align-items:flex-start;
}
.replace-page-meta{
  font-size:11px;color:var(--text-secondary);letter-spacing:.04em;text-transform:uppercase;
  font-weight:600;margin-bottom:6px;
}
.replace-thumb-wrap{
  position:relative;flex-shrink:0;
  background:#fff;border:1px solid var(--bg-hover);border-radius:6px;overflow:hidden;
}
.replace-thumb-wrap img{display:block;user-select:none;-webkit-user-drag:none;}
.replace-thumb-wrap canvas{
  position:absolute;top:0;left:0;
}
.replace-thumb-wrap canvas.region-canvas{
  pointer-events:auto;cursor:default;touch-action:none;
}
.replace-controls{flex:1;min-width:0;display:flex;flex-direction:column;gap:8px;}
.replace-controls select{
  background:var(--bg-input);color:var(--text-primary);border:1.5px solid var(--border-default);
  border-radius:6px;padding:7px 10px;font-family:"Menlo",monospace;
  font-size:12px;outline:none;
}
.replace-controls select:focus{border-color:var(--accent-primary);}
.replace-detected{font-size:11px;color:var(--accent-success);font-family:"Menlo",monospace;}
.replace-detected.warn{color:#f9c97c;}

/* Per-region pickers (one row per detected region within a page) */
.region-pickers{display:flex;flex-direction:column;gap:6px;}
.region-picker{
  display:flex;align-items:center;gap:8px;padding:6px 8px;
  background:var(--bg-card);border:1px solid var(--border-default);border-radius:6px;
}
.region-picker .swatch{
  flex-shrink:0;width:14px;height:14px;border-radius:3px;border:2px solid #fff;
}
.region-picker .lbl{
  flex-shrink:0;font-size:11px;color:var(--text-secondary);letter-spacing:.04em;
  text-transform:uppercase;min-width:64px;font-weight:600;
}
.region-picker select{flex:1;min-width:0;}
.region-picker button.region-snap-btn,
.region-picker button.region-del-btn{
  flex-shrink:0;background:var(--bg-input);color:var(--text-primary);border:1px solid var(--border-default);
  border-radius:5px;padding:4px 8px;font-size:11px;font-family:inherit;
  cursor:pointer;line-height:1;
}
.region-picker button.region-snap-btn:hover{border-color:var(--accent-success);color:var(--accent-success);}
.region-picker button.region-del-btn{padding:4px 9px;font-weight:700;}
.region-picker button.region-del-btn:hover{border-color:var(--accent-error);color:var(--accent-error);}
.region-add-btn{align-self:flex-start;margin-top:4px;}

/* ── ISI chain slicing editor ──────────────────────────────────────────── */
#isi-chain-editors{display:flex;flex-direction:column;gap:14px;margin:14px 0;}
.isi-chain-editor{
  background:var(--bg-card);border:1px solid var(--border-default);border-radius:10px;
  padding:14px 16px;display:flex;flex-direction:column;gap:10px;
}
.isi-chain-head{display:flex;align-items:center;gap:12px;flex-wrap:wrap;}
.isi-chain-head strong{color:var(--text-primary);font-size:13px;}
.isi-chain-meta{color:var(--text-secondary);font-size:11px;letter-spacing:.04em;text-transform:uppercase;}
.isi-chain-hint{color:var(--text-secondary);font-size:11px;line-height:1.5;}
.isi-chain-strip{
  position:relative;display:inline-block;align-self:flex-start;
  border:1px solid var(--border-default);border-radius:6px;overflow:hidden;
  user-select:none;-webkit-user-select:none;
}
.isi-strip-img{
  display:block;max-width:320px;max-height:520px;width:auto;height:auto;
}
.isi-band{
  position:absolute;left:0;right:0;pointer-events:none;
  display:flex;align-items:flex-start;padding:4px 6px;box-sizing:border-box;
}
.isi-band-label{
  font-size:10px;font-weight:700;text-shadow:0 0 4px #000,0 0 4px #000;
  letter-spacing:.04em;text-transform:uppercase;
}
.isi-divider{
  position:absolute;left:0;right:0;height:8px;margin-top:-4px;
  cursor:ns-resize;background:var(--text-primary);opacity:.0;
  border-top:2px solid var(--accent-pink);border-bottom:2px solid var(--accent-pink);
}
.isi-divider:hover{opacity:.35;}
.isi-divider:active{opacity:.55;}

/* ── Open-modal buttons (one per ISI chain) ──────────────────────────── */
.isi-chain-open{
  display:inline-flex;align-items:center;gap:8px;
  margin:6px 8px 6px 0;font-size:12px;
}
.isi-chain-open .muted{color:var(--text-secondary);font-weight:400;font-size:11px;}

/* ── ISI slicing modal ───────────────────────────────────────────────── */
.isi-modal-overlay{
  position:fixed;inset:0;background:rgba(0,0,0,.7);z-index:2000;
  display:flex;align-items:center;justify-content:center;padding:20px;
}
.isi-modal-overlay[hidden]{display:none;}
.isi-modal{
  background:var(--bg-card);border:1px solid var(--border-default);border-radius:10px;
  box-shadow:0 20px 60px rgba(0,0,0,.5);
  width:80vw;max-width:80vw;max-height:90vh;
  display:flex;flex-direction:column;color:var(--text-primary);
}
.isi-modal-head{
  display:flex;align-items:center;gap:14px;padding:12px 16px;
  border-bottom:1px solid var(--border-default);flex-shrink:0;
}
.isi-modal-head strong{font-size:13px;flex-shrink:0;}
.isi-modal-tabs{
  display:flex;gap:4px;flex-wrap:wrap;flex:1;min-width:0;
}
.isi-modal-tab{
  background:var(--bg-input);border:1px solid var(--border-default);color:var(--text-primary);
  padding:4px 10px;font-size:11px;font-family:inherit;border-radius:5px;
  cursor:pointer;letter-spacing:.04em;text-transform:uppercase;font-weight:600;
}
.isi-modal-tab:hover{border-color:var(--accent-success);}
.isi-modal-tab.active{background:var(--accent-success);color:var(--text-inverse);border-color:var(--accent-success);}
.isi-modal-tools{
  display:flex;align-items:center;gap:6px;flex-shrink:0;
  background:var(--bg-input);border:1px solid var(--border-default);border-radius:6px;padding:3px 6px;
}
.isi-modal-tools button{
  background:none;border:none;color:var(--text-primary);cursor:pointer;
  padding:2px 8px;font-size:13px;font-family:inherit;border-radius:3px;
}
.isi-modal-tools button:hover{background:var(--border-default);}
.isi-modal-tools span{font-size:11px;color:var(--text-secondary);min-width:34px;text-align:center;}
.isi-modal-close{
  background:none;border:none;color:var(--text-primary);cursor:pointer;
  font-size:22px;line-height:1;padding:2px 8px;border-radius:4px;
  font-family:inherit;flex-shrink:0;
}
.isi-modal-close:hover{background:var(--border-default);color:var(--accent-error);}

/* Body is a flex row: [left pane] [splitter] [right pane]. Each pane is
 * resizable via the splitter; widths persist for the modal session. */
.isi-modal-body{
  flex:1;display:flex;background:var(--text-inverse);min-height:0;
}
.isi-modal-pane{
  flex:1 1 50%;min-width:120px;overflow:auto;padding:14px;
  display:flex;flex-direction:column;gap:10px;
}
.isi-modal-pane-label{
  font-size:10px;color:var(--text-secondary);letter-spacing:.06em;text-transform:uppercase;
  font-weight:700;flex-shrink:0;
}
.isi-modal-splitter{
  width:6px;background:var(--border-default);cursor:col-resize;flex-shrink:0;
  transition:background .15s ease;
}
.isi-modal-splitter:hover,.isi-modal-splitter.dragging{background:var(--accent-success);}

/* Live preview canvas */
.isi-modal-preview-wrap{
  flex:1;display:flex;align-items:flex-start;justify-content:center;
  overflow:auto;background:var(--bg-input);border:1px solid var(--border-default);border-radius:6px;
  padding:8px;
}
#isi-modal-preview-canvas{
  background:#fff;display:block;max-width:100%;
  box-shadow:0 4px 12px rgba(0,0,0,.4);
}

/* ISI slicer strip */
.isi-modal-strip-wrap{
  flex:1;overflow:auto;display:flex;align-items:flex-start;justify-content:center;
  background:var(--bg-input);border:1px solid var(--border-default);border-radius:6px;padding:8px;
}
.isi-modal-strip{
  position:relative;display:inline-block;
  border:1px solid var(--border-default);border-radius:6px;overflow:hidden;
  user-select:none;-webkit-user-select:none;background:#fff;
}
.isi-modal-strip .isi-strip-img{
  display:block;max-width:none;max-height:70vh;width:auto;height:auto;
}
.isi-modal-strip .isi-band{
  position:absolute;left:0;right:0;pointer-events:none;
  display:flex;align-items:flex-start;padding:4px 6px;box-sizing:border-box;
}
.isi-modal-strip .isi-band-label{
  font-size:10px;font-weight:700;text-shadow:0 0 4px #000,0 0 4px #000;
  letter-spacing:.04em;text-transform:uppercase;
}
/* Independent top + bot handles per region. Higher z-index than bands so
 * clicks always land on the handle even when bands overlap visually. */
.isi-handle{
  position:absolute;left:0;right:0;height:14px;margin-top:-7px;
  cursor:ns-resize;z-index:10;display:flex;align-items:center;
  border-top:2px solid currentColor;border-bottom:2px solid currentColor;
  opacity:.55;
}
.isi-handle:hover,.isi-handle:active{opacity:.85;}
.isi-handle-top{border-bottom:none;}
.isi-handle-bot{border-top:none;}
.isi-handle-label{
  font-size:9px;font-weight:700;letter-spacing:.04em;text-transform:uppercase;
  background:var(--text-inverse);padding:1px 5px;border-radius:3px;color:#fff;
  margin-left:6px;pointer-events:none;
}

.isi-modal-foot{
  display:flex;align-items:center;gap:10px;padding:10px 16px;
  border-top:1px solid var(--border-default);flex-shrink:0;
}
.isi-modal-hint{font-size:11px;color:var(--text-secondary);line-height:1.4;max-width:50%;}

/* Bundle thumbnail strip — visual confirmation of loaded captures */
.bundle-strip{
  display:flex;flex-wrap:wrap;gap:10px;align-items:flex-start;
  margin:0 0 14px 0;padding:10px 12px;background:var(--bg-card);
  border:1px dashed var(--border-default);border-radius:8px;
}
.bundle-strip .bundle-empty{font-size:12px;color:var(--text-secondary);font-style:italic;}
.bundle-strip .bundle-pill{
  display:flex;flex-direction:column;align-items:center;gap:4px;
  padding:6px 8px;background:var(--bg-input);border:1px solid var(--border-default);
  border-radius:6px;min-width:74px;
}
.bundle-strip .bundle-pill img{
  width:64px;height:64px;object-fit:contain;background:#fff;
  border-radius:3px;
}
.bundle-strip .bundle-pill .nm{
  font-size:11px;color:var(--text-primary);font-family:"Menlo",monospace;
}
.bundle-strip .bundle-pill .ct{font-size:10px;color:var(--text-secondary);}
.bundle-strip .bundle-title{
  font-size:11px;color:var(--accent-success);letter-spacing:.04em;text-transform:uppercase;
  font-weight:700;width:100%;margin:0 0 4px 0;
}

/* Global busy overlay (shown during async backend calls) */
#busy-overlay{
  position:fixed;inset:0;background:var(--overlay-dim);
  display:none;align-items:center;justify-content:center;
  z-index:9999;backdrop-filter:blur(2px);
  flex-direction:column;gap:14px;
}
#busy-overlay.show{display:flex;}
#busy-overlay .spinner{
  width:46px;height:46px;border-radius:50%;
  border:4px solid var(--border-default);border-top-color:var(--accent-primary);
  animation:busy-spin 0.9s linear infinite;
}
#busy-overlay .msg{
  font-size:13px;color:var(--text-primary);font-family:"Menlo",monospace;
  letter-spacing:.04em;max-width:520px;text-align:center;
}
@keyframes busy-spin{to{transform:rotate(360deg);}}

/* Inline busy hint shown in toolbars while a fetch is in flight */
.busy-inline{
  display:inline-flex;align-items:center;gap:6px;font-size:12px;color:var(--text-secondary);
}
.busy-inline::before{
  content:"";width:12px;height:12px;border-radius:50%;
  border:2px solid var(--border-default);border-top-color:var(--accent-primary);
  animation:busy-spin 0.7s linear infinite;
}

/* ═══════ CPFP page (tab 6) ═══════ */
.lbl-small{
  font-size:10px;color:var(--text-secondary);letter-spacing:.06em;
  text-transform:uppercase;font-weight:700;
}
.cpfp-pages{display:flex;flex-direction:column;gap:14px;margin-top:14px;}
.cpfp-row{
  background:var(--bg-card);border:1px solid var(--border-default);
  border-radius:10px;padding:14px 16px;
  display:flex;flex-direction:column;gap:10px;
}
.cpfp-row.match{opacity:.65;}        /* identical pages dim slightly */
.cpfp-row.added,
.cpfp-row.removed{border-color:var(--accent-warn);}
.cpfp-row-head{
  display:flex;align-items:center;gap:10px;flex-wrap:wrap;
  font-size:13px;
}
.cpfp-row-head .pill{
  background:var(--bg-elevated);color:var(--text-primary);
  padding:3px 9px;border-radius:11px;font-size:11px;font-weight:700;
  letter-spacing:.04em;text-transform:uppercase;
}
.cpfp-row.match    .pill{background:var(--accent-success);color:var(--text-on-accent);}
.cpfp-row.diff     .pill{background:var(--accent-error);color:var(--text-on-accent);}
.cpfp-row.added    .pill{background:var(--accent-info);color:var(--text-on-accent);}
.cpfp-row.removed  .pill{background:var(--accent-warn);color:var(--text-inverse);}
.cpfp-row-meta{color:var(--text-secondary);font-size:11px;}
.cpfp-pair{
  display:grid;grid-template-columns:1fr 1fr;gap:14px;align-items:flex-start;
}
.cpfp-thumb-wrap{
  position:relative;border:1px solid var(--border-default);border-radius:6px;
  overflow:hidden;background:var(--bg-input);
  display:flex;flex-direction:column;
}
.cpfp-thumb-wrap img{display:block;width:100%;height:auto;}
.cpfp-thumb-label{
  background:var(--bg-elevated);color:var(--text-secondary);
  font-size:10px;letter-spacing:.06em;text-transform:uppercase;font-weight:700;
  padding:5px 10px;border-bottom:1px solid var(--border-default);
}
/* Box overlay over the new-side thumbnail. boxes are positioned/sized in
 * percentages so they scale with the rendered thumbnail width. */
.cpfp-box-layer{position:absolute;top:auto;left:0;right:0;bottom:0;pointer-events:none;}
.cpfp-box{
  position:absolute;border:2px solid var(--accent-error);
  background:rgba(243,139,168,.12);
  box-shadow:0 0 0 1px rgba(0,0,0,.35) inset;
  pointer-events:auto;cursor:zoom-in;
  transition:background .15s ease, border-color .15s ease;
}
.cpfp-box:hover{background:rgba(243,139,168,.28);}
.cpfp-row.match .cpfp-box{display:none;}
.cpfp-stats-line{
  display:flex;flex-wrap:wrap;gap:10px 18px;font-size:11px;
  color:var(--text-secondary);
}
.cpfp-stats-line strong{color:var(--text-primary);}

/* Text-diff panes (CPFP "Text" / "Both" modes) */
.cpfp-text-pair{
  display:grid;grid-template-columns:1fr 1fr;gap:14px;margin-top:10px;
}
.cpfp-text-pane{
  background:var(--bg-input);border:1px solid var(--border-default);
  border-radius:6px;overflow:hidden;display:flex;flex-direction:column;
}
.cpfp-text-body{
  padding:12px 14px;font-size:12px;line-height:1.55;
  white-space:pre-wrap;word-break:break-word;color:var(--text-primary);
  font-family:-apple-system, "Helvetica Neue", sans-serif;
  max-height:340px;overflow-y:auto;
}
.cpfp-text-body em{color:var(--text-muted);font-style:italic;}
.cpfp-text-body ins{
  background:rgba(166,227,161,.18);color:var(--accent-success);
  text-decoration:none;font-weight:700;border-radius:3px;padding:0 2px;
}
.cpfp-text-body del{
  background:rgba(243,139,168,.18);color:var(--accent-error);
  text-decoration:line-through;border-radius:3px;padding:0 2px;
}

/* Claude analysis: summary card + inline per-page annotations */
.cpfp-claude-card .card-title{font-size:13px;}
.cpfp-claude-overall{
  font-size:13px;line-height:1.6;color:var(--text-primary);
  background:var(--bg-elevated);border-radius:8px;padding:14px 16px;
}
.cpfp-claude-raw{
  font-size:12px;line-height:1.5;color:var(--text-secondary);
  background:var(--bg-input);border:1px solid var(--border-default);
  border-radius:6px;padding:12px 14px;white-space:pre-wrap;
  font-family:"Menlo",monospace;max-height:300px;overflow:auto;
}
.cpfp-claude-list{
  display:flex;flex-direction:column;gap:6px;margin-top:8px;
  border-top:1px solid var(--border-default);padding-top:10px;
}
.cpfp-claude-row{
  display:flex;align-items:flex-start;gap:8px;font-size:12px;
  line-height:1.5;color:var(--text-primary);padding:4px 0;
}
.cpfp-claude-sev{
  flex-shrink:0;font-size:11px;font-weight:700;
  letter-spacing:.04em;text-transform:uppercase;border-radius:4px;
  padding:2px 8px;line-height:1.4;
}
.cpfp-claude-sev.critical{background:var(--accent-error);  color:var(--text-on-accent);}
.cpfp-claude-sev.content {background:var(--accent-warn);   color:var(--text-inverse);}
.cpfp-claude-sev.cosmetic{background:var(--accent-info);   color:var(--text-on-accent);}
.cpfp-claude-sev.unknown {background:var(--bg-elevated);   color:var(--text-secondary);}
.cpfp-claude-kind{
  flex-shrink:0;font-size:10px;color:var(--text-secondary);
  letter-spacing:.04em;text-transform:uppercase;font-weight:600;
  align-self:center;
}
.cpfp-claude-conf{
  flex-shrink:0;font-size:10px;color:var(--text-muted);align-self:center;
}
</style>
<script>
// Theme + color-scheme manager. Two independent dimensions:
//   data-theme  ("light" | "dark")        — set on <html>; resolved from
//                                            user choice "auto" via matchMedia
//   data-scheme ("catppuccin" | "tokyo" | "nord" | "solarized")
//
// User stores preference (mode, scheme) in localStorage. JS resolves "auto"
// to a concrete light/dark value at runtime — so the CSS only needs concrete
// (scheme, mode) blocks, no @media queries.
const ThemeManager = {
  KEY_MODE:   'banner_app_theme',
  KEY_SCHEME: 'banner_app_scheme',
  SCHEMES: ['catppuccin', 'tokyo', 'nord', 'solarized'],
  SCHEME_LABELS: {
    catppuccin: 'Catppuccin',
    tokyo:      'Tokyo Night',
    nord:       'Nord',
    solarized:  'Solarized',
  },

  init(){
    const mode   = localStorage.getItem(this.KEY_MODE)   || 'auto';
    const scheme = localStorage.getItem(this.KEY_SCHEME) || 'catppuccin';
    this._applyScheme(scheme);
    this._applyMode(mode);
    if (window.matchMedia){
      window.matchMedia('(prefers-color-scheme: dark)')
            .addEventListener('change', () => {
              if (this.get() === 'auto') this._applyMode('auto');
            });
    }
  },

  // ── Mode (light / dark / auto) ────────────────────────────────────
  get(){ return localStorage.getItem(this.KEY_MODE) || 'auto'; },
  set(mode){
    if (mode !== 'light' && mode !== 'dark' && mode !== 'auto') mode = 'auto';
    localStorage.setItem(this.KEY_MODE, mode);
    this._applyMode(mode);
  },
  _applyMode(mode){
    // Resolve "auto" to a concrete light/dark via the OS preference now,
    // and re-resolve whenever the OS preference changes.
    let concrete = mode;
    if (mode === 'auto'){
      concrete = (window.matchMedia &&
                  window.matchMedia('(prefers-color-scheme: dark)').matches)
        ? 'dark' : 'light';
    }
    document.documentElement.dataset.theme = concrete;
    this._refreshThemeButtons();
  },
  _refreshThemeButtons(){
    const cur = this.get();
    document.querySelectorAll('.theme-toggle button[data-theme]').forEach(b => {
      b.classList.toggle('active', b.dataset.theme === cur);
    });
  },

  // ── Scheme (catppuccin / tokyo / nord / solarized) ────────────────
  getScheme(){ return localStorage.getItem(this.KEY_SCHEME) || 'catppuccin'; },
  setScheme(name){
    if (!this.SCHEMES.includes(name)) name = 'catppuccin';
    localStorage.setItem(this.KEY_SCHEME, name);
    this._applyScheme(name);
    this._closeSchemeMenu();
  },
  _applyScheme(name){
    document.documentElement.dataset.scheme = name;
    this._refreshSchemeUI();
  },
  _refreshSchemeUI(){
    const cur = this.getScheme();
    const lbl = document.getElementById('scheme-name');
    if (lbl) lbl.textContent = this.SCHEME_LABELS[cur] || cur;
    document.querySelectorAll('.scheme-menu button[data-scheme]').forEach(b => {
      b.classList.toggle('active', b.dataset.scheme === cur);
    });
  },
  toggleSchemeMenu(ev){
    if (ev) ev.stopPropagation();
    const menu = document.getElementById('scheme-menu');
    if (!menu) return;
    const isOpen = menu.classList.toggle('open');
    const btn = document.querySelector('#scheme-picker > button');
    if (btn) btn.setAttribute('aria-expanded', isOpen ? 'true' : 'false');
  },
  _closeSchemeMenu(){
    const menu = document.getElementById('scheme-menu');
    if (menu) menu.classList.remove('open');
    const btn = document.querySelector('#scheme-picker > button');
    if (btn) btn.setAttribute('aria-expanded', 'false');
  },
};
ThemeManager.init();
document.addEventListener('DOMContentLoaded', () => {
  ThemeManager._refreshThemeButtons();
  ThemeManager._refreshSchemeUI();
  // Click outside the picker closes it.
  document.addEventListener('click', e => {
    const picker = document.getElementById('scheme-picker');
    if (picker && !picker.contains(e.target)) ThemeManager._closeSchemeMenu();
  });
  // Escape closes the picker.
  document.addEventListener('keydown', e => {
    if (e.key === 'Escape') ThemeManager._closeSchemeMenu();
  });

  // ── Bookmarklet integration ──────────────────────────────────────────────
  // The Banner Tool ships a bookmarklet (see Multi-Banner Viewer tab) that
  // collects every banner URL from a Ziflow proof (running in the user's
  // logged-in browser context, so it works for private Klick proofs without
  // the user having to paste cookies). The bookmarklet sends users back here
  // with `?ml=<url-list>` and `#tab-multi`. Honour both: switch tab, fill
  // the URL field, and auto-fire mbLoad().
  try {
    const params = new URLSearchParams(window.location.search);
    const ml = params.get('ml');
    if (ml) {
      // Pre-fill BOTH tabs so the user can choose which workflow they want:
      //   - Multi-Banner Viewer: bulk capture (auto-loads here, default)
      //   - Frame Capture:       inspect / scrub / capture frames per banner
      // Frame Capture's URL textarea accepts the multi-line list and shows
      // a size dropdown after the user clicks Load on that tab.
      const fcUrl = document.getElementById('url-input');
      if (fcUrl) fcUrl.value = ml;
      // Switch to the Multi-Banner Viewer tab + auto-fire its load.
      try { switchTab('multi'); } catch (e) {}
      const mbUrl = document.getElementById('mb-url');
      if (mbUrl) {
        mbUrl.value = ml;
        setTimeout(() => { try { mbLoad(); } catch (e) {} }, 100);
      }
      // Clean the URL so a refresh doesn't re-trigger.
      const cleaned = window.location.pathname + (window.location.hash || '');
      window.history.replaceState({}, '', cleaned);
    } else if (window.location.hash && window.location.hash.startsWith('#tab-')) {
      // Allow deep-linking to a tab without other params.
      const tab = window.location.hash.slice('#tab-'.length);
      try { switchTab(tab); } catch (e) {}
    }
  } catch (e) { /* never let init noise break the page */ }
});
</script>
</head>
<body>

<div id="busy-overlay">
  <div class="spinner"></div>
  <div class="msg" id="busy-msg">Working…</div>
</div>

<div class="topbar">
  <h1>Banner Tool</h1>
  <nav class="tabs">
    <button class="tab-btn active" data-tab="capture" onclick="switchTab('capture')">Frame Capture</button>
    <button class="tab-btn"        data-tab="isi"     onclick="switchTab('isi')">ISI Compare</button>
    <button class="tab-btn"        data-tab="bc"      onclick="switchTab('bc')">Banner Compare</button>
    <button class="tab-btn"        data-tab="multi"   onclick="switchTab('multi')">Multi-Banner Viewer</button>
    <button class="tab-btn"        data-tab="story"   onclick="switchTab('story')">Storyboard</button>
    <button class="tab-btn"        data-tab="cpfp"    onclick="switchTab('cpfp')">CPFP</button>
  </nav>
  <span class="sub" id="topbar-sub">Folder or Ziflow link · Preview · Export PNG + MP4</span>
  <div class="scheme-picker" id="scheme-picker">
    <button type="button" onclick="ThemeManager.toggleSchemeMenu(event)"
            title="Choose color scheme" aria-haspopup="menu" aria-expanded="false">
      🎨 <span id="scheme-name">Catppuccin</span> <span class="chev">▾</span>
    </button>
    <div class="scheme-menu" id="scheme-menu" role="menu">
      <button type="button" data-scheme="catppuccin"
              onclick="ThemeManager.setScheme('catppuccin')"
              style="--swatch-1:#cba6f7;--swatch-2:#a6e3a1;--swatch-3:#f38ba8;">
        <span class="swatch"></span>Catppuccin
      </button>
      <button type="button" data-scheme="tokyo"
              onclick="ThemeManager.setScheme('tokyo')"
              style="--swatch-1:#bb9af7;--swatch-2:#7aa2f7;--swatch-3:#7dcfff;">
        <span class="swatch"></span>Tokyo Night
      </button>
      <button type="button" data-scheme="nord"
              onclick="ThemeManager.setScheme('nord')"
              style="--swatch-1:#88c0d0;--swatch-2:#a3be8c;--swatch-3:#b48ead;">
        <span class="swatch"></span>Nord
      </button>
      <button type="button" data-scheme="solarized"
              onclick="ThemeManager.setScheme('solarized')"
              style="--swatch-1:#b58900;--swatch-2:#268bd2;--swatch-3:#dc322f;">
        <span class="swatch"></span>Solarized
      </button>
    </div>
  </div>
  <div class="theme-toggle" role="group" aria-label="Color theme">
    <button type="button" data-theme="light" onclick="ThemeManager.set('light')"
            title="Light mode" aria-label="Light mode">☀</button>
    <button type="button" data-theme="auto"  onclick="ThemeManager.set('auto')"
            title="Follow system" aria-label="Follow system">🖥</button>
    <button type="button" data-theme="dark"  onclick="ThemeManager.set('dark')"
            title="Dark mode" aria-label="Dark mode">🌙</button>
  </div>
</div>

<!-- ═══════════════════════ TAB 1: FRAME CAPTURE ═══════════════════════ -->
<section class="tab-pane active" id="tab-capture">
<div class="layout">

  <!-- ═══════════ LEFT SIDEBAR ═══════════ -->
  <aside class="sidebar">

    <!-- Load banner -->
    <div class="card">
      <p class="card-title">Banner Source</p>

      <!-- Local folder upload -->
      <p class="src-label">Upload a banner folder</p>
      <div class="row">
        <button class="btn btn-p btn-sm" id="browse-btn" onclick="document.getElementById('banner-folder-input').click()">Choose folder…</button>
        <input id="banner-folder-input" type="file" webkitdirectory directory multiple
               style="display:none" onchange="uploadBannerFolder(this)">
        <input id="folder-input" type="text"
               placeholder="Folder name (auto-fills)"
               readonly
               style="background:#f7f8fa;color:#666">
      </div>

      <p class="src-divider">— or —</p>

      <!-- Ziflow / DoubleClick URL — textarea so the bookmarklet's multi-line
           asset-URL paste works here too (private Klick Ziflow proofs that
           the server can't discover anonymously). -->
      <p class="src-label">Ziflow proof URL, DoubleClick preview URL, or paste multiple asset URLs</p>
      <div class="row">
        <textarea id="url-input" rows="1"
               placeholder="https://….ziflow.io/proof/… &#10;OR https://www.google.com/doubleclick/studio/externalpreview/… &#10;OR https://proof-assets.ziflow.io/Proofs/.../richMedia/.../index.html (one per line)"
               style="resize:none;min-height:34px;font-family:Menlo,monospace;font-size:12px;padding:8px 10px;border-radius:7px;border:1.5px solid var(--border-default);background:var(--bg-input);color:var(--text-primary);outline:none;"
               onkeydown="if(event.key==='Enter' && !event.shiftKey){event.preventDefault();commitPath(this.value);}"></textarea>
        <button class="btn btn-p btn-sm" id="url-btn"
                onclick="commitPath(document.getElementById('url-input').value)">Load</button>
      </div>

      <!-- Banner-size picker — appears only when the loaded URL has multiple
           creatives (DoubleClick Studio with N sizes). Switching reloads the
           preview iframe with the chosen creative. -->
      <div id="dc-picker-wrap" class="row" style="display:none;margin-top:8px;align-items:center;gap:8px;">
        <label style="font-size:11px;color:var(--text-secondary);letter-spacing:.04em;">SIZE</label>
        <select id="dc-picker"
                style="flex:1;padding:6px 8px;border-radius:7px;border:1.5px solid var(--border-default);background:var(--bg-input);color:var(--text-primary);font-family:Menlo,monospace;font-size:12px;"
                onchange="onDcPickerChange()"></select>
      </div>

      <p id="pstatus"></p>
      <p class="hint">
        For local: hold <strong>Option</strong> + right-click folder in Finder →
        <em>Copy "…" as Pathname</em>, then paste above.
      </p>
    </div>

    <!-- Export options — match the Multi-Banner Viewer's 4-checkbox layout. -->
    <div class="card export-opts">
      <p class="card-title">Export Options</p>
      <div class="check-row" style="flex-direction:column;align-items:flex-start;gap:6px;">
        <label><input type="checkbox" id="opt-frames"   checked> Frames (PNG stills at the times below)</label>
        <label><input type="checkbox" id="opt-isi-png"  checked> ISI image</label>
        <label><input type="checkbox" id="opt-isi-html" checked> ISI HTML</label>
        <label><input type="checkbox" id="opt-video"    checked> MP4 video</label>
      </div>
      <button class="btn btn-p btn-full" id="run-btn" disabled onclick="runCapture()">
        Run Capture
      </button>
      <button class="btn btn-g btn-full" id="open-btn" disabled onclick="openFolder()">
        ⬇ Download zip
      </button>
      <p class="card-title" style="margin-top:10px;margin-bottom:4px;">Progress</p>
      <div id="log">Ready — load a banner folder to begin.
</div>
      <div id="runstatus"></div>
    </div>

  </aside>

  <!-- ═══════════ MAIN AREA ═══════════ -->
  <main class="main">

    <!-- Empty state -->
    <div id="empty-state" style="display:flex;flex-direction:column;align-items:center;
         justify-content:center;flex:1;gap:12px;color:var(--border-default);padding:60px 0;">
      <svg width="56" height="56" viewBox="0 0 24 24" fill="none" stroke="currentColor"
           stroke-width="1.2" stroke-linecap="round" stroke-linejoin="round">
        <rect x="2" y="3" width="20" height="14" rx="2"/>
        <path d="M8 21h8M12 17v4"/>
      </svg>
      <p style="font-size:14px;">Load a banner folder to preview it here</p>
    </div>

    <!-- Preview panel -->
    <div id="preview-panel">
      <div class="card">
        <p class="card-title">Live Preview</p>
        <div class="preview-wrap">

          <div class="banner-shell">
            <iframe id="banner-frame" src="about:blank"></iframe>
          </div>

          <div class="controls">
            <div class="time-display" id="time-disp">t = 0.00 s</div>
            <div class="pb-row">
              <button class="btn btn-p btn-sm" onclick="iCmd('play')">▶ Play</button>
              <button class="btn btn-g btn-sm" onclick="iCmd('pause')">⏸ Pause</button>
              <button class="btn btn-g btn-sm" onclick="iCmd('restart')">↺ Restart</button>
            </div>
            <input type="range" id="scrubber" min="0" max="1000" value="0"
                   onmousedown="scrubbing=true" ontouchstart="scrubbing=true"
                   onmouseup="scrubEnd()"      ontouchend="scrubEnd()"
                   oninput="onScrub(this.value)">
            <button class="add-frame-btn" onclick="addCurrentFrame()">
              + Capture this frame
            </button>
          </div>

        </div>
      </div>

      <!-- Frame list -->
      <div class="card">
        <p class="card-title">Frames to Export</p>
        <div class="frame-list" id="frame-list"></div>
        <button class="btn btn-g btn-sm" style="margin-top:10px"
                onclick="addBlankFrame()">+ Add frame</button>
      </div>
    </div>

  </main>
</div>
</section>

<!-- ═══════════════════════ TAB 2: ISI COMPARE ═══════════════════════ -->
<section class="tab-pane" id="tab-isi">
  <div class="compare-wrap">

    <div class="compare-toolbar">
      <span style="font-size:13px;color:var(--text-primary);">Compare a banner's ISI to a PDF/Word ISI document.</span>
      <span class="compare-pill" id="isi-banner-pill">No banner loaded yet — load one on Frame Capture</span>
      <button class="btn btn-g btn-sm" onclick="isiLoadBanner()">⟳ Refresh from page 1</button>
      <button class="btn btn-p btn-sm" onclick="document.getElementById('isi-doc-input').click()">📄 Upload PDF / Word</button>
      <input id="isi-doc-input" type="file" accept=".pdf,.doc,.docx"
             style="display:none" onchange="isiUploadDoc(this)">
      <span class="compare-pill" id="isi-doc-pill">No document loaded</span>
    </div>

    <p class="hint" style="margin:0 0 12px;">
      <strong>Tip:</strong> edit either box to trim away unrelated text,
      or highlight a portion and click <em>Run diff</em> — only the selected
      text in each box will be compared.
    </p>

    <!-- Editable source panes -->
    <div class="compare-grid">
      <div class="compare-col">
        <h3>Banner ISI <span class="isi-sel-info" id="isi-left-info"></span></h3>
        <textarea class="diff-source" id="isi-left-src"
                  placeholder="Load a banner on the Frame Capture page, then click ‘Refresh from page 1’."
                  oninput="isiInputChanged()"
                  onmouseup="isiSelChanged()" onkeyup="isiSelChanged()"></textarea>
      </div>
      <div class="compare-col">
        <h3>Document ISI <span class="isi-sel-info" id="isi-right-info"></span></h3>
        <textarea class="diff-source" id="isi-right-src"
                  placeholder="Upload a PDF or Word file with the reference ISI text."
                  oninput="isiInputChanged()"
                  onmouseup="isiSelChanged()" onkeyup="isiSelChanged()"></textarea>
      </div>
    </div>

    <div class="row" style="margin:14px 0 6px;align-items:center;gap:10px;">
      <button class="btn btn-p btn-sm" onclick="isiRunDiff()">Run diff</button>
      <button class="btn btn-g btn-sm" onclick="isiResetSources()">Reset to extracted text</button>
      <span class="compare-pill" id="isi-scope" style="margin-left:auto;">Will diff: full text on both sides</span>
    </div>

    <!-- Diff result -->
    <div class="diff-summary" id="isi-summary" style="display:none;"></div>
    <div class="compare-grid" id="isi-result" style="display:none;">
      <div class="compare-col">
        <h3>Banner ISI · diff</h3>
        <div class="diff-text" id="isi-left"></div>
      </div>
      <div class="compare-col">
        <h3>Document ISI · diff</h3>
        <div class="diff-text" id="isi-right"></div>
      </div>
    </div>

  </div>
</section>

<!-- ═══════════════════════ TAB 3: BANNER COMPARE ═══════════════════════ -->
<section class="tab-pane" id="tab-bc">
  <div class="compare-wrap">

    <!-- Two loaders side by side -->
    <div class="bc-toolbar">
      <div class="bc-loader">
        <h3>Banner A</h3>
        <div class="row">
          <button class="btn btn-p btn-sm" onclick="document.getElementById('bcA-folder-input').click()">Choose folder…</button>
          <input id="bcA-folder-input" type="file" webkitdirectory directory multiple
                 style="display:none" onchange="bcUpload('A',this)">
          <input type="text" id="bcA-folder" placeholder="Folder name (auto-fills)"
                 readonly style="background:#f7f8fa;color:#666">
        </div>
        <div class="row">
          <input type="text" id="bcA-url" placeholder="Or Ziflow URL…"
                 onkeydown="if(event.key==='Enter')bcLoad('A',this.value)">
          <button class="btn btn-g btn-sm" onclick="bcLoad('A',document.getElementById('bcA-url').value)">Load</button>
        </div>
        <p class="bc-status" id="bcA-status"></p>
      </div>
      <div class="bc-loader">
        <h3>Banner B</h3>
        <div class="row">
          <button class="btn btn-p btn-sm" onclick="document.getElementById('bcB-folder-input').click()">Choose folder…</button>
          <input id="bcB-folder-input" type="file" webkitdirectory directory multiple
                 style="display:none" onchange="bcUpload('B',this)">
          <input type="text" id="bcB-folder" placeholder="Folder name (auto-fills)"
                 readonly style="background:#f7f8fa;color:#666">
        </div>
        <div class="row">
          <input type="text" id="bcB-url" placeholder="Or Ziflow URL…"
                 onkeydown="if(event.key==='Enter')bcLoad('B',this.value)">
          <button class="btn btn-g btn-sm" onclick="bcLoad('B',document.getElementById('bcB-url').value)">Load</button>
        </div>
        <p class="bc-status" id="bcB-status"></p>
      </div>
    </div>

    <!-- Master controls (synced) -->
    <div class="bc-master" id="bc-master" style="display:none;">
      <div class="time-display" id="bc-time">t = 0.00 s</div>
      <button class="btn btn-p btn-sm" onclick="bcSyncCmd('play')">▶ Play</button>
      <button class="btn btn-g btn-sm" onclick="bcSyncCmd('pause')">⏸ Pause</button>
      <button class="btn btn-g btn-sm" onclick="bcSyncCmd('restart')">↺ Restart</button>
      <input type="range" id="bc-scrubber" min="0" max="1000" value="0"
             onmousedown="bcScrubbing=true" ontouchstart="bcScrubbing=true"
             onmouseup="bcScrubbing=false" ontouchend="bcScrubbing=false"
             oninput="bcOnScrub(this.value)">
      <button class="bc-link" id="bc-link-toggle" onclick="bcToggleLink()">🔗 Linked</button>
      <button class="btn btn-g btn-sm" onclick="bcDoVisualDiff()">🔍 Visual diff at this time</button>
    </div>

    <!-- Side-by-side previews -->
    <div class="bc-views" id="bc-views" style="display:none;">
      <div class="bc-view">
        <div class="bc-view-title" id="bc-A-title">Banner A</div>
        <div class="bc-shell"><iframe id="bcA-frame" src="about:blank"></iframe></div>
        <div class="bc-indep-controls" id="bcA-indep">
          <div style="display:flex;gap:6px;">
            <button class="btn btn-p btn-sm" onclick="bcIndepCmd('A','play')">▶</button>
            <button class="btn btn-g btn-sm" onclick="bcIndepCmd('A','pause')">⏸</button>
            <button class="btn btn-g btn-sm" onclick="bcIndepCmd('A','restart')">↺</button>
            <span style="font-family:Menlo;font-size:12px;flex:1;text-align:right;color:var(--accent-success);" id="bcA-t">0.00 s</span>
          </div>
          <input type="range" min="0" max="1000" value="0" id="bcA-scrub"
                 oninput="bcIndepScrub('A',this.value)">
        </div>
      </div>
      <div class="bc-view">
        <div class="bc-view-title" id="bc-B-title">Banner B</div>
        <div class="bc-shell"><iframe id="bcB-frame" src="about:blank"></iframe></div>
        <div class="bc-indep-controls" id="bcB-indep">
          <div style="display:flex;gap:6px;">
            <button class="btn btn-p btn-sm" onclick="bcIndepCmd('B','play')">▶</button>
            <button class="btn btn-g btn-sm" onclick="bcIndepCmd('B','pause')">⏸</button>
            <button class="btn btn-g btn-sm" onclick="bcIndepCmd('B','restart')">↺</button>
            <span style="font-family:Menlo;font-size:12px;flex:1;text-align:right;color:var(--accent-success);" id="bcB-t">0.00 s</span>
          </div>
          <input type="range" min="0" max="1000" value="0" id="bcB-scrub"
                 oninput="bcIndepScrub('B',this.value)">
        </div>
      </div>
    </div>

    <!-- Differences panel -->
    <div class="bc-diff-section" id="bc-diff-section" style="display:none;">
      <h3>Differences</h3>
      <div id="bc-meta-block"></div>

      <h3 style="margin-top:18px;">ISI text diff</h3>
      <div class="diff-summary" id="bc-isi-summary" style="display:none;"></div>
      <div class="compare-grid">
        <div class="compare-col">
          <h3>Banner A · ISI</h3>
          <div class="diff-text" id="bc-isi-left"></div>
        </div>
        <div class="compare-col">
          <h3>Banner B · ISI</h3>
          <div class="diff-text" id="bc-isi-right"></div>
        </div>
      </div>

      <div id="bc-vis-result" style="display:none;margin-top:18px;">
        <h3>Visual diff at <span id="bc-vis-t"></span></h3>
        <div class="bc-vis">
          <figure><img id="bc-vis-a"   alt="Banner A"><figcaption>Banner A</figcaption></figure>
          <figure><img id="bc-vis-b"   alt="Banner B"><figcaption>Banner B</figcaption></figure>
          <figure><img id="bc-vis-diff" alt="Diff"><figcaption>Differences (red overlay)</figcaption></figure>
        </div>
      </div>
    </div>

  </div>
</section>

<!-- ═══════════════════════ TAB 4: MULTI-BANNER VIEWER ═══════════════════════ -->
<section class="tab-pane" id="tab-multi">
  <div class="compare-wrap">

    <!-- Loader card -->
    <div class="mb-loader">
      <!-- ═══ Bookmarklet for private Klick Ziflow proofs ═══════════════════
           One-time setup: drag the purple button to your bookmarks bar.
           Then on any Ziflow proof page (where you're already logged in)
           click the bookmark — it cycles through every banner, collects
           each iframe src, and re-opens the Banner Tool with the URLs
           auto-filled. No DevTools, no manual paste. -->
      <div style="padding:10px 12px;margin-bottom:10px;border:1.5px dashed var(--border-default);border-radius:8px;background:var(--bg-elevated);">
        <p style="margin:0 0 6px;font-size:11px;color:var(--text-secondary);letter-spacing:.04em;">
          🚀 ONE-CLICK ZIFLOW LOAD (Klick proofs)
        </p>
        <div class="row" style="align-items:center;gap:10px;">
          <a id="mb-bookmarklet" class="btn btn-p btn-sm"
             href="javascript:void(0)"
             title="Drag this to your bookmarks bar (one-time). Then visit any Ziflow proof while logged in and click the bookmark — banners auto-load here."
             onclick="event.preventDefault(); alert('Drag this button to your bookmarks bar instead of clicking it.\n\n(Right-click it and choose Bookmark Link to add it manually.)');">
            📌 Banner Tool: Get from Ziflow
          </a>
          <button class="btn btn-g btn-sm" onclick="mbCopyBookmarklet()">📋 Copy bookmarklet code</button>
          <span style="font-size:11px;color:var(--text-muted);">
            Drag the purple button → bookmarks bar (one-time). Then click it on any Ziflow proof.
          </span>
        </div>
      </div>

      <div class="row" style="margin-bottom:8px;">
        <label style="flex:1;">Staging URL <span style="color:var(--text-muted);">(or paste many asset URLs, one per line)</span>
          <textarea id="mb-url" rows="1"
                    placeholder="https://klick.link-staging.com/... &#10;OR&#10;https://proof-assets.ziflow.io/Proofs/.../richMedia/.../index.html (one per line)"
                    style="resize:none;min-height:34px;font-family:Menlo,monospace;font-size:12px;padding:8px 10px;border-radius:7px;border:1.5px solid var(--border-default);background:var(--bg-input);color:var(--text-primary);outline:none;"></textarea>
        </label>
        <label style="flex:1;">Username <span style="color:var(--text-muted);">(optional)</span>
          <input type="text" id="mb-user" placeholder="leave blank for Klick links">
        </label>
        <label style="flex:1;">Password
          <input type="password" id="mb-pass" placeholder="password">
        </label>
        <button class="btn btn-p btn-sm" onclick="mbLoad()" style="align-self:flex-end;">Load all banners</button>
      </div>
      <p class="bc-status" id="mb-status"></p>

      <p class="hint" style="margin:8px 0 0;">
        <strong>Klick / AdPiler:</strong> use the password from the staging email — leave Username blank.<br>
        <strong>DoubleClick Studio</strong> (URLs starting with
        <code>www.google.com/doubleclick/studio/externalpreview</code>): no
        login needed; just paste the URL.<br>
        <strong>Ziflow proof</strong>: public proofs work directly. For
        <em>private</em> Klick proofs
        (<code>klickhealth.ziflow.io/proof/&lt;id&gt;</code>) use the
        <em>📌 Get from Ziflow</em> bookmarklet above, or paste each banner's
        <code>proof-assets.ziflow.io/.../richMedia/&lt;guid&gt;/index.html</code>
        URL on a separate line, or use <em>Storyboard → Load folder</em>.<br>
        <strong>HTTP Basic Auth</strong>: fill both Username + Password fields.
      </p>
    </div>

    <!-- Master controls (hidden until banners load) -->
    <div class="mb-master" id="mb-master" style="display:none;">
      <div class="time-display" id="mb-time">t = 0.00 s</div>
      <button class="btn btn-p btn-sm" onclick="mbCmd('play')">▶ Play</button>
      <button class="btn btn-g btn-sm" onclick="mbCmd('pause')">⏸ Pause</button>
      <button class="btn btn-g btn-sm" onclick="mbCmd('restart')">↺ Restart</button>
      <input type="range" id="mb-scrubber" min="0" max="1000" value="0"
             onmousedown="mbScrubbing=true" ontouchstart="mbScrubbing=true"
             onmouseup="mbScrubbing=false" ontouchend="mbScrubbing=false"
             oninput="mbOnScrub(this.value)">
      <button class="btn btn-g btn-sm" onclick="mbAddTime()"
              title="Add this time to the capture list">+ Add time</button>
      <button class="btn btn-g btn-sm" onclick="mbSelectAll(true)">Select all</button>
      <button class="btn btn-g btn-sm" onclick="mbSelectAll(false)">Select none</button>
      <button class="btn btn-p" onclick="mbRunAll()"
              title="Capture selected banners and download as a zip">⬇ Capture &amp; download</button>
    </div>

    <!-- Frame list shared across all banners -->
    <div class="card" id="mb-frame-list-wrap" style="display:none;margin-bottom:16px;">
      <p class="card-title">Frames to capture (applied to every banner)</p>
      <div class="frame-list" id="mb-frame-list"></div>
      <p class="hint" style="margin-top:6px;">
        Use the master scrubber + <em>+ Add time</em> above to populate this list.
        Empty list → no per-frame stills (ISI/MP4 still export if checked).
      </p>
      <div class="row" style="margin-top:8px;gap:14px;flex-wrap:wrap;">
        <label><input type="checkbox" id="mb-do-frames"   checked> Frames (PNG stills at the times above)</label>
        <label><input type="checkbox" id="mb-do-isi-png"  checked> ISI image</label>
        <label><input type="checkbox" id="mb-do-isi-html" checked> ISI HTML</label>
        <label><input type="checkbox" id="mb-do-video"    checked> MP4 video</label>
      </div>
    </div>

    <!-- Per-banner tiles -->
    <div class="mb-views" id="mb-views"></div>

    <!-- Run progress + log -->
    <div class="card" id="mb-runlog-wrap" style="display:none;margin-top:16px;">
      <p class="card-title">Capture progress</p>
      <div id="mb-log" style="background:var(--bg-input);border-radius:7px;padding:12px;
                              font-family:Menlo;font-size:11px;color:var(--accent-success);
                              height:200px;overflow-y:auto;white-space:pre-wrap;
                              word-break:break-all;line-height:1.5;"></div>
      <div id="mb-runstatus" style="margin-top:8px;font-size:12px;color:var(--text-secondary);text-align:center;"></div>
    </div>

  </div>
</section>

<!-- ═══════════════════════ TAB 5: STORYBOARD ═══════════════════════ -->
<section class="tab-pane" id="tab-story">
  <div class="compare-wrap">

    <!-- Mode toggle: Generate new vs. Update existing PDF -->
    <div class="story-mode">
      <button class="mode-btn active" id="story-mode-gen"     onclick="storySetMode('generate')">Generate new</button>
      <button class="mode-btn"        id="story-mode-replace" onclick="storySetMode('replace')">Update existing PDF</button>
    </div>

    <!-- ════════ MODE 1: Generate new storyboard ════════ -->
    <div id="story-generate-mode">
      <div class="story-toolbar">
        <label>Storyboard header
          <input type="text" id="story-header" placeholder="ALYFTREK HCP 2026 — Q1 Media Display Banners — RM Announcement">
        </label>
        <label>Footer text
          <input type="text" id="story-footer" placeholder="© 2026 Vertex Pharmaceuticals Incorporated | VXR-US-34-2500240 (v7.0) | 04/2026">
        </label>
        <button class="btn btn-g btn-sm" onclick="storyAutoLoad()">⟳ Refresh from latest capture</button>
        <button class="btn btn-g btn-sm" onclick="document.getElementById('story-bundle-zip-input').click()">📦 Load .zip…</button>
        <button class="btn btn-g btn-sm" onclick="document.getElementById('story-bundle-folder-input').click()">📁 Load folder…</button>
        <input id="story-bundle-zip-input" type="file" accept=".zip"
               style="display:none" onchange="storyUploadBundle(this)">
        <input id="story-bundle-folder-input" type="file" webkitdirectory directory multiple
               style="display:none" onchange="storyUploadBundle(this)">
      </div>

      <div class="bundle-strip" id="bundle-strip-gen">
        <span class="bundle-empty">No capture bundle loaded yet.</span>
      </div>

      <div id="story-banners">
        <p class="hint" style="padding:24px 8px;color:var(--text-secondary);">
          Run a capture on the <strong>Multi-Banner Viewer</strong> tab first.
          When it finishes, come back here — the captured frames will appear automatically.
          Or click <em>Load bundle (.zip)</em> to upload a previously-downloaded capture zip.
        </p>
      </div>

      <div class="story-actions" id="story-actions" style="display:none;">
        <button class="btn btn-p" onclick="storyBuild()">⬇ Generate Storyboard PDF</button>
        <button class="btn btn-g btn-sm" id="story-download-link" style="display:none;"
                onclick="storyDownloadLastPdf()">Download last PDF</button>
        <span class="story-status" id="story-status"></span>
      </div>
    </div>

    <!-- ════════ MODE 2: Update existing PDF ════════ -->
    <div id="story-replace-mode" style="display:none;">

      <div class="story-toolbar">
        <button class="btn btn-p btn-sm" onclick="document.getElementById('story-replace-input').click()">📄 Upload existing storyboard PDF…</button>
        <input id="story-replace-input" type="file" accept=".pdf"
               style="display:none" onchange="storyReplaceUpload(this)">
        <span class="compare-pill" id="story-replace-name">No PDF loaded</span>
        <button class="btn btn-g btn-sm" onclick="storyAutoLoad()">⟳ Refresh capture bundle</button>
      </div>

      <div class="bundle-strip" id="bundle-strip-rep">
        <span class="bundle-empty">No capture bundle loaded yet.</span>
      </div>

      <p class="hint" style="margin:0 0 12px;">
        Upload the existing storyboard PDF, then for each page pick which captured banner's frames should replace
        the images on it. The animation notes, headers and footers stay exactly as they were.
        Click <em>Refresh preview</em> to see what each page will look like, then <em>Export</em> when satisfied.
      </p>

      <div id="story-replace-pages"></div>

      <div class="story-actions" id="story-replace-actions" style="display:none;">
        <button class="btn btn-g" onclick="storyReplacePreview()">↻ Refresh preview</button>
        <button class="btn btn-p" onclick="storyReplaceExport()">⬇ Export updated PDF</button>
        <button class="btn btn-g btn-sm" id="story-replace-download-link"
                style="display:none;" onclick="storyDownloadReplacedPdf()">Download last updated PDF</button>
        <span class="story-status" id="story-replace-status"></span>
      </div>

      <!-- ISI slicing modal — opened per chain by openIsiModal(path).
           Layout: 80vw wide, split into a left live-preview pane and a
           right slicing pane separated by a draggable splitter. Page tabs
           let the user switch which page is being previewed when the chain
           spans multiple pages. -->
      <div id="isi-modal" class="isi-modal-overlay" hidden>
        <div class="isi-modal" role="dialog" aria-modal="true">
          <div class="isi-modal-head">
            <strong>ISI Slicing — <span id="isi-modal-title"></span></strong>
            <div id="isi-modal-tabs" class="isi-modal-tabs"></div>
            <div class="isi-modal-tools">
              <button type="button" onclick="isiModalZoom('out')" title="Zoom out 10%">−</button>
              <span id="isi-modal-zoom-label">Fit</span>
              <button type="button" onclick="isiModalZoom('in')"  title="Zoom in 10%">+</button>
              <button type="button" onclick="isiModalZoom('fit')" title="Fit to modal">Fit</button>
            </div>
            <button type="button" class="isi-modal-close"
                    onclick="closeIsiModal()" aria-label="Close">×</button>
          </div>
          <div class="isi-modal-body">
            <div id="isi-modal-left" class="isi-modal-pane">
              <div class="isi-modal-pane-label">Live preview</div>
              <div id="isi-modal-preview-wrap" class="isi-modal-preview-wrap">
                <canvas id="isi-modal-preview-canvas"></canvas>
              </div>
            </div>
            <div id="isi-modal-splitter" class="isi-modal-splitter" title="Drag to resize"></div>
            <div id="isi-modal-right" class="isi-modal-pane">
              <div class="isi-modal-pane-label">ISI slicer</div>
              <div id="isi-modal-strip-wrap" class="isi-modal-strip-wrap"></div>
            </div>
          </div>
          <div class="isi-modal-foot">
            <span class="isi-modal-hint">
              Drag colored top/bottom handles to set each region's slice
              bounds. Slices may overlap or leave gaps. Preview updates live.
            </span>
            <div style="flex:1"></div>
            <button type="button" class="btn btn-g btn-sm" onclick="isiModalReset()">↺ Reset to auto</button>
            <button type="button" class="btn btn-p btn-sm" onclick="isiModalApply()">✓ Apply this page</button>
          </div>
        </div>
      </div>

    </div>

  </div>
</section>

<!-- ═══════════════════════ TAB 6: CPFP ═══════════════════════ -->
<section class="tab-pane" id="tab-cpfp">
  <div class="compare-wrap">
    <div class="card">
      <p class="card-title">Change-proof full-proof (CPFP)</p>
      <p class="hint" style="margin-bottom:10px;">
        Upload the <strong>old</strong> storyboard and the <strong>new</strong>
        storyboard. The app aligns them strictly by page index, computes a
        pixel-level diff for each pair, and draws red boxes around every
        region that differs — so you can spot both requested changes and
        unintended ones at a glance.
      </p>
      <div class="row" style="gap:14px;flex-wrap:wrap;align-items:end;">
        <div style="display:flex;flex-direction:column;gap:4px;">
          <span class="lbl-small">Old version</span>
          <button class="btn btn-g" onclick="document.getElementById('cpfp-old-input').click()"
                  title="Choose the original / pre-revision storyboard PDF">
            📂 Choose old PDF…
          </button>
          <input id="cpfp-old-input" type="file" accept=".pdf"
                 style="display:none" onchange="cpfpUpload('old', this)">
          <span class="compare-pill" id="cpfp-old-pill">No file loaded</span>
        </div>
        <div style="font-size:24px;color:var(--text-muted);align-self:center;">⇄</div>
        <div style="display:flex;flex-direction:column;gap:4px;">
          <span class="lbl-small">New version</span>
          <button class="btn btn-g" onclick="document.getElementById('cpfp-new-input').click()"
                  title="Choose the revised storyboard PDF">
            📂 Choose new PDF…
          </button>
          <input id="cpfp-new-input" type="file" accept=".pdf"
                 style="display:none" onchange="cpfpUpload('new', this)">
          <span class="compare-pill" id="cpfp-new-pill">No file loaded</span>
        </div>
        <div style="margin-left:auto;display:flex;gap:8px;align-items:center;">
          <div class="theme-toggle" role="group" aria-label="Diff mode" id="cpfp-mode">
            <button type="button" data-mode="visual" class="active"
                    onclick="cpfpSetMode('visual')"
                    title="Pixel diff with red boxes around regions that differ">📐 Visual</button>
            <button type="button" data-mode="text"
                    onclick="cpfpSetMode('text')"
                    title="Word-level text diff — fast, focused on copy edits">📝 Text</button>
            <button type="button" data-mode="both"
                    onclick="cpfpSetMode('both')"
                    title="Show both visual boxes and text diff">⚡ Both</button>
          </div>
          <button class="btn btn-p" id="cpfp-run-btn" onclick="cpfpRun()" disabled>
            ⚖ Compare
          </button>
          <button class="btn btn-g" id="cpfp-claude-btn" onclick="cpfpAnalyzeWithClaude()"
                  disabled title="Layer Claude's plain-English analysis on top of the deterministic diff">
            🤖 Analyze with Claude
          </button>
          <button class="btn btn-g btn-sm" id="cpfp-connect-btn"
                  onclick="cpfpOpenConnectModal()"
                  title="Connect your Anthropic API key so Claude can analyze the diff">
            🔌 Connect Claude
          </button>
          <button class="btn btn-g btn-sm" onclick="cpfpClear()">↺ Reset</button>
        </div>
      </div>
      <p class="bc-status" id="cpfp-status" style="margin-top:8px;"></p>
    </div>

    <!-- Per-page diff results -->
    <div id="cpfp-summary" class="card" style="display:none;margin-top:16px;">
      <p class="card-title" id="cpfp-summary-title">Diff results</p>
      <div id="cpfp-summary-stats" class="hint" style="margin-bottom:6px;"></div>
    </div>

    <!-- Claude analysis card (populated by cpfpRenderClaude when present) -->
    <div id="cpfp-claude-card" class="card cpfp-claude-card" style="display:none;margin-top:14px;">
      <p class="card-title">🤖 Claude's review</p>
      <div id="cpfp-claude-summary"></div>
      <div id="cpfp-claude-stats" class="cpfp-stats-line" style="margin-top:8px;"></div>
    </div>

    <div id="cpfp-pages" class="cpfp-pages"></div>

    <!-- API-key modal — opened by 🔌 Connect Claude / closed on ✓ Connect or ✕ -->
    <div id="cpfp-connect-modal" class="isi-modal-overlay" hidden>
      <div class="isi-modal" role="dialog" aria-modal="true" style="width:520px;max-width:90vw;">
        <div class="isi-modal-head">
          <strong>🔌 Connect Claude</strong>
          <button type="button" class="isi-modal-close" style="margin-left:auto;"
                  onclick="cpfpCloseConnectModal()" aria-label="Close">×</button>
        </div>
        <div class="isi-modal-body" style="display:block;padding:16px 20px;">
          <p style="font-size:13px;line-height:1.55;color:var(--text-primary);margin-bottom:12px;">
            Paste your Anthropic API key. The key is saved locally to
            <code style="background:var(--bg-input);padding:1px 5px;border-radius:3px;font-size:12px;">~/.banner_app/config.json</code>
            with owner-only permissions and used only for the CPFP "Analyze with Claude" button.
          </p>
          <p style="font-size:12px;color:var(--text-secondary);margin-bottom:14px;">
            Don't have a key? Get one at
            <a href="https://console.anthropic.com/settings/keys" target="_blank"
               style="color:var(--accent-primary);">console.anthropic.com</a>.
          </p>
          <input type="password" id="cpfp-api-key-input"
                 placeholder="sk-ant-…"
                 style="width:100%;padding:9px 12px;border-radius:6px;
                        border:1px solid var(--border-default);background:var(--bg-input);
                        color:var(--text-primary);font-family:Menlo,monospace;font-size:12px;
                        outline:none;"
                 onkeydown="if(event.key==='Enter') cpfpDoConnect();">
          <p class="bc-status" id="cpfp-connect-status" style="margin-top:10px;font-size:12px;"></p>
        </div>
        <div class="isi-modal-foot">
          <span class="isi-modal-hint">
            Your key never leaves this Mac except to call api.anthropic.com.
          </span>
          <div style="flex:1;"></div>
          <button class="btn btn-g btn-sm" id="cpfp-disconnect-btn"
                  style="display:none;" onclick="cpfpDoDisconnect()">Disconnect</button>
          <button class="btn btn-p btn-sm" onclick="cpfpDoConnect()">✓ Connect</button>
        </div>
      </div>
    </div>
  </div>
</section>

<script>
// ── Global busy indicator ────────────────────────────────────────────────────
// Stack-based so concurrent fetches don't dismiss each other early.
let _busyDepth = 0;
function setBusy(msg){
  _busyDepth += 1;
  const ov = document.getElementById('busy-overlay');
  const m  = document.getElementById('busy-msg');
  if(m)  m.textContent = msg || 'Working…';
  if(ov) ov.classList.add('show');
}
function clearBusy(){
  _busyDepth = Math.max(0, _busyDepth - 1);
  if(_busyDepth === 0){
    const ov = document.getElementById('busy-overlay');
    if(ov) ov.classList.remove('show');
  }
}
// busyFetch() wraps fetch() so callers don't have to remember to clear.
async function busyFetch(url, opts, msg){
  setBusy(msg);
  try { return await fetch(url, opts); }
  finally { clearBusy(); }
}

// ── Toast notifications ──────────────────────────────────────────────────────
// Briefly pops up at the bottom-right with success / info / error. Auto-fades.
// `kind` is 'ok' | 'info' | 'err'. Survives multiple stacked toasts.
function showToast(message, kind){
  let host = document.getElementById('toast-host');
  if (!host) {
    host = document.createElement('div');
    host.id = 'toast-host';
    host.style.cssText = 'position:fixed;right:16px;bottom:16px;z-index:9999;'
                       + 'display:flex;flex-direction:column;gap:8px;'
                       + 'pointer-events:none;';
    document.body.appendChild(host);
  }
  const t = document.createElement('div');
  const palette = (kind === 'err')
    ? 'background:#fee2e2;color:#991b1b;border:1px solid #fca5a5;'
    : (kind === 'ok')
      ? 'background:#dcfce7;color:#14532d;border:1px solid #86efac;'
      : 'background:#e0e7ff;color:#1e3a8a;border:1px solid #a5b4fc;';
  t.style.cssText = palette
    + 'padding:10px 14px;border-radius:8px;font-size:13px;'
    + 'box-shadow:0 4px 12px rgba(0,0,0,0.15);max-width:380px;'
    + 'opacity:0;transform:translateY(8px);transition:opacity .2s, transform .2s;'
    + 'pointer-events:auto;font-family:inherit;';
  t.textContent = message;
  host.appendChild(t);
  // Trigger transition into view.
  requestAnimationFrame(() => { t.style.opacity = '1'; t.style.transform = 'translateY(0)'; });
  // Auto-dismiss after 5s. Click also dismisses early.
  const dismiss = () => {
    t.style.opacity = '0'; t.style.transform = 'translateY(8px)';
    setTimeout(() => t.remove(), 250);
  };
  t.addEventListener('click', dismiss);
  setTimeout(dismiss, 5000);
}

// ── Save-As download helper ───────────────────────────────────────────────────
// Two-tier strategy:
//
//   1. showSaveFilePicker() — true Save As dialog (Chromium browsers, top-level
//      pages only). Lets the user pick folder + filename.
//
//   2. Blob URL + <a download> — Safari, Firefox, embedded contexts. Fetch the
//      bytes ourselves, wrap in a blob, click an <a download href="blob:...">.
//      Works in all browsers; download attribute on a blob URL is honoured.
//
// (No third "navigate to URL" fallback — it caused the tab to land on a "not
// found" page when same-origin clicks raced with the download.) On failure,
// `downloadFile` shows an error toast and returns false.
//
// MUST be called from a user-gesture handler (button onclick).
async function downloadFile(url, suggestedName, mimeType, extension){
  mimeType = mimeType || 'application/octet-stream';
  const _toplevel = (window.self === window.top);

  // Path 1: Save As dialog (Chromium top-level only).
  if (_toplevel && typeof window.showSaveFilePicker === 'function') {
    try {
      const opts = { suggestedName: suggestedName };
      if (extension) {
        opts.types = [{
          description: extension.replace(/^\./,'').toUpperCase() + ' file',
          accept: { [mimeType]: [extension] }
        }];
      }
      const handle = await window.showSaveFilePicker(opts);
      setBusy('Downloading ' + suggestedName + '…');
      const resp = await fetch(url, { credentials: 'same-origin' });
      if (!resp.ok) throw new Error('Server returned HTTP ' + resp.status);
      const blob = await resp.blob();
      const writable = await handle.createWritable();
      await writable.write(blob);
      await writable.close();
      clearBusy();
      showToast('✓ Saved ' + suggestedName, 'ok');
      return true;
    } catch (e) {
      clearBusy();
      if (e && e.name === 'AbortError') return false;   // user cancelled — silent
      console.warn('showSaveFilePicker failed; falling back to blob link:', e);
      // fall through to path 2
    }
  }

  // Path 2: Fetch as blob, then click <a download> with a blob URL.
  try {
    setBusy('Downloading ' + suggestedName + '…');
    const resp = await fetch(url, { credentials: 'same-origin' });
    if (!resp.ok) throw new Error('Server returned HTTP ' + resp.status);
    const blob = await resp.blob();
    const blobUrl = URL.createObjectURL(blob);
    const a = document.createElement('a');
    a.href = blobUrl;
    a.download = suggestedName || '';
    a.style.display = 'none';
    document.body.appendChild(a);
    try { a.click(); } finally {
      // Cleanup, but don't let removal errors mask the click.
      try { document.body.removeChild(a); } catch(_) {}
      setTimeout(() => { try { URL.revokeObjectURL(blobUrl); } catch(_) {} }, 60_000);
    }
    clearBusy();
    showToast('✓ Downloaded ' + suggestedName + ' — check your Downloads folder', 'ok');
    return true;
  } catch (e) {
    clearBusy();
    console.error('Download failed:', e);
    showToast('⚠ Download failed: ' + (e && e.message || e), 'err');
    return false;
  }
}

// ── Tab switching ──────────────────────────────────────────────────────────────
function switchTab(name){
  document.querySelectorAll('.tab-pane').forEach(el=>el.classList.remove('active'));
  document.querySelectorAll('.tab-btn').forEach(el=>el.classList.remove('active'));
  document.getElementById('tab-'+name).classList.add('active');
  document.querySelector('.tab-btn[data-tab="'+name+'"]').classList.add('active');
  // Update subtitle
  const subs = {
    capture: 'Folder or Ziflow link · Preview · Export PNG + MP4',
    isi:     'Compare banner ISI text against a PDF or Word document',
    bc:      'Load two banners · Side-by-side preview · Spot the differences',
    multi:   'Klick staging URL · Load every size at once · Bulk capture',
    story:   'Build a storyboard PDF from captured frames + your animation notes',
  };
  if(name === 'story') storyAutoLoad();
  document.getElementById('topbar-sub').textContent = subs[name] || '';
}

// ── State ──────────────────────────────────────────────────────────────────────
let selectedPath = null;
let currentTime  = 0;
let frameCounter = 0;
let isRemote     = false;   // true when banner came from a Ziflow URL

// Frame list starts empty — user adds frames via "📸 Capture this frame" or "+ Add frame".

// ── Browse / Upload ────────────────────────────────────────────────────────────
async function uploadBannerFolder(input){
  // User picked a folder via webkitdirectory. Upload every file as
  // multipart/form-data, preserving the relative path so the server can
  // find index.html.
  const btn = document.getElementById('browse-btn');
  const txt = document.getElementById('folder-input');
  if(!input.files || !input.files.length){ return; }
  btn.disabled = true; const old = btn.textContent; btn.textContent='Uploading…';
  setStatus('Uploading ' + input.files.length + ' files…', '');
  const fd = new FormData();
  for(const f of input.files){
    // webkitRelativePath preserves the folder structure (e.g. "myBanner/index.html")
    fd.append('files', f, f.webkitRelativePath || f.name);
  }
  try{
    const r = await busyFetch('/upload/banner', {method:'POST', body: fd},
                              'Uploading banner folder…');
    const data = await r.json();
    if(data.ok){
      txt.value = data.name + '  (' + data.files + ' files)';
      selectedPath = data.name;
      isRemote = true;
      _bothInputs().forEach(i => i.classList.remove('ok'));
      txt.classList.add('ok');
      // Mirror commitPath's success branch so the iframe loads and the
      // run button enables.
      setStatus('✓  '+data.name+' ('+data.width+'×'+data.height+')', 'ok');
      document.getElementById('run-btn').disabled = false;
      document.getElementById('empty-state').style.display = 'none';
      document.getElementById('preview-panel').style.display = 'flex';
      const bw = (data.width||160)+'px', bh = (data.height||600)+'px';
      const shell = document.querySelector('.banner-shell');
      const frame = document.getElementById('banner-frame');
      shell.style.width=bw; shell.style.height=bh;
      frame.style.width=bw; frame.style.height=bh;
      // Cache-bust the iframe so the new banner is loaded.
      frame.src = '/banner/index.html?t=' + Date.now();
      appendLog('Loaded: '+data.name+' ('+data.width+'×'+data.height+')\n');
    } else {
      setStatus('Upload error: ' + (data.error || 'unknown'), 'err');
    }
  }catch(e){ setStatus('Upload error: '+e.message, 'err'); }
  finally{ btn.textContent = old; btn.disabled = false; }
}
function _activeInput(p){
  // Highlight whichever input matches the value being committed.
  const url = /^https?:\/\//i.test(p);
  return document.getElementById(url ? 'url-input' : 'folder-input');
}
function _bothInputs(){
  return [document.getElementById('folder-input'),
          document.getElementById('url-input')];
}
async function commitPath(p, creativeIndex){
  // creativeIndex is optional — passed when re-validating a DoubleClick URL
  // to pick a different size from the dropdown.
  p = p.trim().replace(/\/+$/,'');
  if(!p)return;
  const isUrl = /^https?:\/\//i.test(p);
  const isDoubleClick = /google\.com\/doubleclick\/studio\/externalpreview/i.test(p);
  if(isUrl){
    setStatus(isDoubleClick
      ? (creativeIndex != null
          ? 'Switching to ' + (window._dcCreatives && window._dcCreatives[creativeIndex] ? window._dcCreatives[creativeIndex].name : 'creative #' + creativeIndex) + '…'
          : 'Loading DoubleClick preview (cycling through creatives — ~15 s)…')
      : 'Downloading from Ziflow…', '');
  }
  const payload = {path: p};
  if(creativeIndex != null) payload.creative_index = creativeIndex;
  const data = await (await busyFetch('/validate', {
    method:'POST', headers:{'Content-Type':'application/json'},
    body: JSON.stringify(payload)
  }, isUrl ? 'Downloading banner from URL…' : 'Validating banner folder…')).json();
  if(data.ok){
    selectedPath = p;
    isRemote     = !!data.is_remote;
    _bothInputs().forEach(i => i.classList.remove('ok'));
    _activeInput(p).classList.add('ok');
    let badge = '';
    // The size dropdown handles three sources uniformly: DoubleClick Studio
    // (`data.doubleclick`), multi-asset paste from the bookmarklet for
    // private Klick Ziflow proofs (`data.multi_banners`), and single-banner
    // URLs (no dropdown).
    const picker = data.doubleclick || data.multi_banners;
    if(picker){
      const label = data.doubleclick ? 'DoubleClick' : 'Ziflow';
      badge = '  — ' + label + ' · loaded ' + picker.picked_name
            + (picker.total_count > 1
                ? ' (' + ((picker.picked_index || 0) + 1) + ' of ' + picker.total_count + ')'
                : '');
      _renderDcPicker(picker);
    } else {
      const wrap = document.getElementById('dc-picker-wrap');
      if(wrap) wrap.style.display = 'none';
      window._dcCreatives = null;
      if(isRemote){
        badge = '  — Ziflow';
        // Klick proofs are private and only the og:image's banner shows up.
        // Hint at the bookmarklet so the user can pull in the others.
        if(/klickhealth\.ziflow\.io\/proof\//i.test(p)){
          badge += '  · multi-banner? Click the 📌 bookmarklet on this proof to load all sizes.';
        }
      }
    }
    setStatus('✓  '+data.name+' ('+data.width+'×'+data.height+')' + badge, 'ok');
    document.getElementById('run-btn').disabled=false;
    document.getElementById('empty-state').style.display='none';
    document.getElementById('preview-panel').style.display='flex';
    // Size the shell + iframe to the detected banner dimensions
    const bw = (data.width||160)+'px', bh = (data.height||600)+'px';
    const shell = document.querySelector('.banner-shell');
    const frame = document.getElementById('banner-frame');
    shell.style.width=bw; shell.style.height=bh;
    frame.style.width=bw; frame.style.height=bh;
    // Cache-bust so the iframe reloads when we switch creatives.
    frame.src = '/banner/index.html?t=' + Date.now();
    appendLog('Loaded: '+p+' ('+data.width+'×'+data.height+')' + badge + '\n');
  }else{
    _bothInputs().forEach(i => i.classList.remove('ok'));
    setStatus(data.error||'No index.html found.','err');
    document.getElementById('run-btn').disabled=true;
    document.getElementById('preview-panel').style.display='none';
    document.getElementById('empty-state').style.display='flex';
  }
}

// Build the size dropdown from DoubleClick discovery info. `dc.all_sizes`
// is an array of names (e.g. ['160x600', '300x250', ...]); `dc.picked_index`
// is the currently-loaded creative's position in that list.
function _renderDcPicker(dc){
  const wrap = document.getElementById('dc-picker-wrap');
  const sel  = document.getElementById('dc-picker');
  if(!wrap || !sel) return;
  if(!dc.total_count || dc.total_count < 2){
    wrap.style.display = 'none';
    window._dcCreatives = null;
    return;
  }
  // Cache the creative names so onDcPickerChange can show user-readable
  // status while the swap is in flight.
  window._dcCreatives = (dc.all_sizes || []).map((name, i) => ({ name, index: i }));
  sel.innerHTML = '';
  (dc.all_sizes || []).forEach((name, i) => {
    const opt = document.createElement('option');
    opt.value = String(i);
    opt.textContent = name;
    if(i === (dc.picked_index || 0)) opt.selected = true;
    sel.appendChild(opt);
  });
  wrap.style.display = '';
}

// Dropdown onchange — re-validate the same URL with a different creative_index.
async function onDcPickerChange(){
  const sel = document.getElementById('dc-picker');
  if(!sel || !selectedPath) return;
  const idx = parseInt(sel.value, 10);
  if(isNaN(idx)) return;
  await commitPath(selectedPath, idx);
}
function setStatus(msg,cls){
  const el=document.getElementById('pstatus');
  el.textContent=msg; el.className=cls||'';
}

// ── iframe postMessage bridge ──────────────────────────────────────────────────
let scrubbing    = false;
let animDuration = 20;   // updated when iframe sends 'ready'

function iCmd(cmd, extra){
  const f=document.getElementById('banner-frame');
  try{ f.contentWindow.postMessage(Object.assign({bc:cmd},extra),'*'); }catch(e){}
}

// Scrubber interaction
function onScrub(val){
  const t = (val/1000)*animDuration;
  currentTime = t;
  document.getElementById('time-disp').textContent='t = '+t.toFixed(2)+' s';
  iCmd('seek',{time:t});
}
function scrubEnd(){ scrubbing=false; }

// Messages from iframe
window.addEventListener('message',function(e){
  if(!e.data||!e.data.bc) return;
  if(e.data.bc==='ready'){
    animDuration = e.data.duration||20;
  }
  if(e.data.bc==='state'){
    currentTime = e.data.time||0;
    document.getElementById('time-disp').textContent='t = '+currentTime.toFixed(2)+' s';
    if(!scrubbing && animDuration>0){
      document.getElementById('scrubber').value=
        Math.round((currentTime/animDuration)*1000);
    }
  }
});

// ── Frame list ─────────────────────────────────────────────────────────────────
function addCurrentFrame(){
  const t = parseFloat(currentTime.toFixed(2));
  frameCounter++;
  const name = 'custom_'+String(frameCounter).padStart(2,'0')+'_t'+t+'s';
  addFrameItem(name, t);
}
function addBlankFrame(){
  frameCounter++;
  addFrameItem('frame_'+String(frameCounter).padStart(2,'0'), 0);
}
function addFrameItem(name, t){
  const id = 'fr-'+Date.now()+'-'+Math.random().toString(36).slice(2);
  const div = document.createElement('div');
  div.className='frame-item'; div.id=id;
  div.innerHTML=`
    <input type="text"   value="${name}" placeholder="label">
    <input type="number" value="${t}" step="0.1" min="0" style="width:70px;flex:none;" placeholder="t(s)">
    <button class="frame-del" onclick="document.getElementById('${id}').remove()" title="Remove">×</button>`;
  document.getElementById('frame-list').appendChild(div);
}
function getFrames(){
  return [...document.querySelectorAll('.frame-item')].map(el=>{
    const inputs = el.querySelectorAll('input');
    return {name: inputs[0].value.trim()||'frame', t: parseFloat(inputs[1].value)||0};
  });
}

// ── Run capture ────────────────────────────────────────────────────────────────
async function runCapture(){
  if(!selectedPath)return;
  const frames      = getFrames();
  const wantFrames  = document.getElementById('opt-frames').checked;
  const wantIsiPng  = document.getElementById('opt-isi-png').checked;
  const wantIsiHtml = document.getElementById('opt-isi-html').checked;
  const wantVideo   = document.getElementById('opt-video').checked;
  const rs          = document.getElementById('runstatus');

  // Block: Frames is checked but no frame entries.
  if(wantFrames && frames.length === 0){
    rs.textContent = '⚠ Add at least one frame to export — use "📸 Capture this frame" or "+ Add frame", or uncheck Frames.';
    rs.style.color = 'var(--accent-error)';
    return;
  }
  // Block: nothing at all is selected.
  if(!wantFrames && !wantIsiPng && !wantIsiHtml && !wantVideo){
    rs.textContent = '⚠ Pick at least one export option.';
    rs.style.color = 'var(--accent-error)';
    return;
  }
  // Reset the warning color when we're actually proceeding.
  rs.style.color = '';

  document.getElementById('run-btn').disabled=true;
  document.getElementById('open-btn').disabled=true;
  document.getElementById('runstatus').textContent='Running…';

  await busyFetch('/run', {
    method:'POST', headers:{'Content-Type':'application/json'},
    body: JSON.stringify({
      path:        selectedPath,
      frames:      frames,
      do_frames:   wantFrames,
      do_isi_png:  wantIsiPng,
      do_isi_html: wantIsiHtml,
      do_video:    wantVideo,
    })
  }, 'Starting capture…');

  // Track the zip URL the server tells us about so the "Download" button
  // can stream it. Reset on each run.
  window._lastZipUrl = null;

  const es=new EventSource('/progress');
  es.onmessage=(e)=>{
    const msg=JSON.parse(e.data);
    if(msg.type==='log'){
      appendLog(msg.text+'\n');
    }else if(msg.type==='done'){
      appendLog('\n✓ Done — capture complete.\n');
      document.getElementById('runstatus').textContent='Complete ✓ — click Download to save the zip';
      document.getElementById('run-btn').disabled=false;
      if(msg.zip_url){
        window._lastZipUrl = msg.zip_url;
        const ob = document.getElementById('open-btn');
        ob.textContent = '⬇ Download zip'; ob.disabled = false;
      }
      es.close();
    }else if(msg.type==='error'){
      appendLog('ERROR: '+msg.text+'\n');
      document.getElementById('runstatus').textContent='Error — see log';
      document.getElementById('run-btn').disabled=false;
      es.close();
    }
  };
}
async function openFolder(){
  // Hosted: download the zip the server prepared. The "Open" button is
  // repurposed as "Download zip" — opens a Save As dialog where supported.
  if(window._lastZipUrl){
    const name = (selectedPath || 'banner').split(/[\\\/]/).pop() + '_frames.zip';
    await downloadFile(window._lastZipUrl, name, 'application/zip', '.zip');
  }
}
function appendLog(t){
  const el=document.getElementById('log');
  el.textContent+=t; el.scrollTop=el.scrollHeight;
}

// ════════════════════════════════════════════════════════════════════════════
// TAB 2 — ISI Compare
// ════════════════════════════════════════════════════════════════════════════
// Originals as extracted; user can edit the sources freely.
let isiBannerText = '';
let isiDocText    = '';

function _isiSel(el){
  // Returns the currently-selected substring of a textarea, or null if none.
  const s = el.selectionStart, e = el.selectionEnd;
  if(s == null || e == null || s === e) return null;
  return el.value.substring(s, e);
}
function _wordCount(s){
  return (s||'').trim().split(/\s+/).filter(Boolean).length;
}

function isiInputChanged(){ isiSelChanged(); }

function isiSelChanged(){
  // Update selection-info badges + the "Will diff" pill so the user knows what
  // is going to be compared on the next Run-diff click.
  const left  = document.getElementById('isi-left-src');
  const right = document.getElementById('isi-right-src');
  const lSel  = _isiSel(left);
  const rSel  = _isiSel(right);

  const lInfo = document.getElementById('isi-left-info');
  const rInfo = document.getElementById('isi-right-info');
  lInfo.textContent = lSel ? ('— ' + _wordCount(lSel) + ' word selection') : '';
  rInfo.textContent = rSel ? ('— ' + _wordCount(rSel) + ' word selection') : '';

  const scope = document.getElementById('isi-scope');
  if(lSel && rSel){
    scope.textContent = 'Will diff: selections on both sides';
    scope.className   = 'compare-pill ok';
  } else if(lSel){
    scope.textContent = 'Will diff: selection on left vs full text on right';
    scope.className   = 'compare-pill ok';
  } else if(rSel){
    scope.textContent = 'Will diff: full text on left vs selection on right';
    scope.className   = 'compare-pill ok';
  } else {
    scope.textContent = 'Will diff: full text on both sides';
    scope.className   = 'compare-pill';
  }
}

async function isiLoadBanner(){
  const pill = document.getElementById('isi-banner-pill');
  pill.textContent = 'Reading banner ISI…'; pill.className = 'compare-pill';
  try{
    const r = await busyFetch('/isi/banner',
        {method:'POST', headers:{'Content-Type':'application/json'}, body:'{}'},
        'Reading banner ISI…');
    const d = await r.json();
    if(!d.ok){
      pill.textContent = '⚠ ' + (d.error || 'Failed');
      pill.className = 'compare-pill warn';
      return;
    }
    isiBannerText = d.isi || '';
    pill.textContent = '✓ ' + d.name;
    pill.className   = 'compare-pill ok';
    document.getElementById('isi-left-src').value = isiBannerText;
    isiSelChanged();
    if(isiDocText) isiRunDiff();
  }catch(e){
    pill.textContent = '⚠ ' + e.message;
    pill.className = 'compare-pill warn';
  }
}

async function isiUploadDoc(input){
  const pill = document.getElementById('isi-doc-pill');
  if(!input.files || !input.files.length){ return; }
  const file = input.files[0];
  pill.textContent = 'Uploading ' + file.name + '…';
  pill.className = 'compare-pill';
  try{
    // Step 1: upload the binary to /upload/isi-doc
    const fd = new FormData(); fd.append('file', file, file.name);
    const u = await busyFetch('/upload/isi-doc', {method:'POST', body: fd},
                              'Uploading document…');
    const ud = await u.json();
    if(!ud.ok){
      pill.textContent = '⚠ ' + (ud.error || 'Upload failed');
      pill.className = 'compare-pill warn';
      return;
    }
    // Step 2: ask /isi/upload to extract the text
    const r = await busyFetch('/isi/upload', {
      method:'POST', headers:{'Content-Type':'application/json'},
      body: JSON.stringify({filename: ud.filename})
    }, 'Reading uploaded document…');
    const d = await r.json();
    if(!d.ok){
      pill.textContent = '⚠ ' + (d.error || 'Failed to read');
      pill.className = 'compare-pill warn';
      return;
    }
    isiDocText = d.text || '';
    pill.textContent = '✓ ' + d.name;
    pill.className   = 'compare-pill ok';
    document.getElementById('isi-right-src').value = isiDocText;
    isiSelChanged();
    if(isiBannerText) isiRunDiff();
  }catch(e){
    pill.textContent = '⚠ ' + e.message;
    pill.className = 'compare-pill warn';
  }
}

function isiResetSources(){
  document.getElementById('isi-left-src').value  = isiBannerText;
  document.getElementById('isi-right-src').value = isiDocText;
  isiSelChanged();
}

async function isiRunDiff(){
  const left  = document.getElementById('isi-left-src');
  const right = document.getElementById('isi-right-src');
  // Use selection if present in a given side, otherwise the full textarea.
  const a = _isiSel(left)  ?? left.value;
  const b = _isiSel(right) ?? right.value;

  if(!a.trim() || !b.trim()){
    const sum = document.getElementById('isi-summary');
    sum.style.display = 'block';
    sum.className     = 'diff-summary diff';
    sum.textContent   = '⚠ Both sides need text. Load a banner + upload a document, or paste text directly.';
    document.getElementById('isi-result').style.display = 'none';
    return;
  }

  const r = await busyFetch('/isi/diff', {
    method:'POST', headers:{'Content-Type':'application/json'},
    body: JSON.stringify({a: a, b: b}),
  }, 'Comparing ISI text…');
  const d = await r.json();
  if(!d.ok) return;
  document.getElementById('isi-result').style.display = 'grid';
  document.getElementById('isi-left').innerHTML  = d.left;
  document.getElementById('isi-right').innerHTML = d.right;
  const sum = document.getElementById('isi-summary');
  sum.style.display = 'block';
  if(d.identical){
    sum.className   = 'diff-summary match';
    sum.textContent = '✓ Identical (text only — styling ignored).';
  }else{
    sum.className   = 'diff-summary diff';
    sum.textContent = '⚠ ' + d.diffs + ' word' + (d.diffs===1?'':'s')
                    + ' differ. Red strikethrough = only on left. Green = only on right.';
  }
}

// ════════════════════════════════════════════════════════════════════════════
// TAB 3 — Banner Compare
// ════════════════════════════════════════════════════════════════════════════
let bcA_loaded = false, bcB_loaded = false;
let bcLinked   = true;
let bcScrubbing= false;
let bcA_dur    = 20, bcB_dur = 20;
let bcA_time   = 0,  bcB_time = 0;

async function bcUpload(slot, input){
  if(!input.files || !input.files.length) return;
  const status = document.getElementById('bc'+slot+'-status');
  status.textContent = 'Uploading ' + input.files.length + ' files…';
  status.className   = 'bc-status';
  const fd = new FormData();
  fd.append('slot', slot);
  for(const f of input.files){
    fd.append('files', f, f.webkitRelativePath || f.name);
  }
  try{
    const r = await busyFetch('/upload/banner', {method:'POST', body: fd},
                              'Uploading banner folder ' + slot + '…');
    const d = await r.json();
    if(!d.ok){
      status.textContent = '⚠ ' + (d.error || 'Upload failed');
      status.className   = 'bc-status err';
      return;
    }
    document.getElementById('bc'+slot+'-folder').value = d.name + '  (' + d.files + ' files)';
    status.textContent = '✓ ' + d.name + ' (' + d.width + '×' + d.height + ')';
    status.className   = 'bc-status ok';
    if(slot==='A') bcA_loaded = true; else bcB_loaded = true;
    // Size + load the iframe.
    const shell = document.querySelector('#bc'+slot+'-frame').parentElement;
    const frame = document.getElementById('bc'+slot+'-frame');
    shell.style.width = d.width+'px';  shell.style.height = d.height+'px';
    frame.style.width = d.width+'px';  frame.style.height = d.height+'px';
    frame.src = '/banner'+slot+'/index.html?t=' + Date.now();
    if(bcA_loaded && bcB_loaded){
      document.getElementById('bc-master').style.display = 'flex';
      document.getElementById('bc-views').style.display  = 'grid';
      document.getElementById('bc-diff-section').style.display = 'block';
      bcLoadDiffInfo();
    } else {
      document.getElementById('bc-views').style.display = 'grid';
    }
  }catch(e){
    status.textContent = '⚠ ' + e.message;
    status.className   = 'bc-status err';
  }
}

async function bcLoad(slot, raw){
  raw = (raw||'').trim();
  if(!raw) return;
  const status = document.getElementById('bc'+slot+'-status');
  const isUrl  = /^https?:\/\//i.test(raw);
  status.textContent = isUrl ? 'Downloading from Ziflow…' : 'Validating…';
  status.className   = 'bc-status';
  try{
    const r = await busyFetch('/p3/validate', {
      method:'POST', headers:{'Content-Type':'application/json'},
      body: JSON.stringify({slot:slot, path:raw}),
    }, isUrl ? 'Downloading banner from URL…' : 'Validating banner folder…');
    const d = await r.json();
    if(!d.ok){
      status.textContent = '⚠ ' + (d.error || 'Failed');
      status.className   = 'bc-status err';
      return;
    }
    status.textContent = '✓ ' + d.name + ' (' + d.width + '×' + d.height + ')'
                         + (d.is_remote ? ' — Ziflow' : '');
    status.className   = 'bc-status ok';
    if(slot==='A') bcA_loaded = true; else bcB_loaded = true;

    // Size the iframe and shell to the detected dims
    const shell = document.querySelector('#bc'+slot+'-frame').parentElement;
    const frame = document.getElementById('bc'+slot+'-frame');
    shell.style.width = d.width+'px';  shell.style.height = d.height+'px';
    frame.style.width = d.width+'px';  frame.style.height = d.height+'px';
    frame.src = '/banner'+slot+'/index.html';

    if(bcA_loaded && bcB_loaded){
      document.getElementById('bc-master').style.display = 'flex';
      document.getElementById('bc-views').style.display  = 'grid';
      document.getElementById('bc-diff-section').style.display = 'block';
      bcLoadDiffInfo();
    } else {
      document.getElementById('bc-views').style.display = 'grid';
    }
  }catch(e){
    status.textContent = '⚠ ' + e.message;
    status.className   = 'bc-status err';
  }
}

function bcSyncCmd(cmd){
  bcSendIframe('A', cmd);
  bcSendIframe('B', cmd);
}
function bcIndepCmd(slot, cmd){ bcSendIframe(slot, cmd); }
function bcSendIframe(slot, cmd, extra){
  const f = document.getElementById('bc'+slot+'-frame');
  try{ f.contentWindow.postMessage(Object.assign({bc:cmd}, extra||{}), '*'); }catch(e){}
}
function bcOnScrub(val){
  const tA = (val/1000)*bcA_dur;
  const tB = (val/1000)*bcB_dur;
  bcSendIframe('A', 'seek', {time: tA});
  bcSendIframe('B', 'seek', {time: tB});
  document.getElementById('bc-time').textContent = 't = ' + tA.toFixed(2) + ' s';
}
function bcIndepScrub(slot, val){
  const dur = (slot==='A') ? bcA_dur : bcB_dur;
  const t   = (val/1000) * dur;
  bcSendIframe(slot, 'seek', {time: t});
  document.getElementById('bc'+slot+'-t').textContent = t.toFixed(2) + ' s';
}
function bcToggleLink(){
  bcLinked = !bcLinked;
  const btn = document.getElementById('bc-link-toggle');
  if(bcLinked){
    btn.textContent = '🔗 Linked';
    btn.classList.remove('unlinked');
    document.getElementById('bcA-indep').classList.remove('show');
    document.getElementById('bcB-indep').classList.remove('show');
  } else {
    btn.textContent = '🔓 Unlinked';
    btn.classList.add('unlinked');
    document.getElementById('bcA-indep').classList.add('show');
    document.getElementById('bcB-indep').classList.add('show');
  }
}

// Listen for state messages from the two iframes
window.addEventListener('message', function(e){
  if(!e.data || !e.data.bc) return;
  // Figure out which slot
  let slot = null;
  const fA = document.getElementById('bcA-frame');
  const fB = document.getElementById('bcB-frame');
  if(fA && e.source === fA.contentWindow) slot = 'A';
  else if(fB && e.source === fB.contentWindow) slot = 'B';
  if(!slot) return;
  if(e.data.bc === 'ready'){
    if(slot==='A') bcA_dur = e.data.duration || 20;
    else           bcB_dur = e.data.duration || 20;
  }
  if(e.data.bc === 'state'){
    const t = e.data.time || 0;
    if(slot==='A'){ bcA_time = t; document.getElementById('bcA-t').textContent = t.toFixed(2)+' s'; }
    else          { bcB_time = t; document.getElementById('bcB-t').textContent = t.toFixed(2)+' s'; }
    if(bcLinked && !bcScrubbing && bcA_dur > 0){
      const pct = bcA_time / bcA_dur;
      document.getElementById('bc-scrubber').value = Math.round(pct * 1000);
      document.getElementById('bc-time').textContent = 't = ' + bcA_time.toFixed(2) + ' s';
    }
    if(!bcLinked){
      const dur = (slot==='A') ? bcA_dur : bcB_dur;
      if(dur>0) document.getElementById('bc'+slot+'-scrub').value = Math.round((t/dur)*1000);
    }
  }
});

async function bcLoadDiffInfo(){
  const block = document.getElementById('bc-meta-block');
  block.innerHTML = '<p style="color:var(--text-secondary);font-size:12px;">Reading both banners…</p>';
  const r = await busyFetch('/p3/info',
      {method:'POST', headers:{'Content-Type':'application/json'}, body:'{}'},
      'Reading banner metadata…');
  const d = await r.json();
  if(!d.ok){
    block.innerHTML = '<p style="color:var(--accent-error);">' + (d.error || 'Failed') + '</p>';
    return;
  }
  // Metadata diff
  const rows = [
    ['Name',           d.a.name,                                  d.b.name],
    ['Dimensions',     d.a.width + '×' + d.a.height,              d.b.width + '×' + d.b.height],
    ['Animation',      (d.a.duration||'?') + ' s',                (d.b.duration||'?') + ' s'],
    ['ISI word count', String((d.a.isi_text||'').split(/\s+/).filter(Boolean).length),
                       String((d.b.isi_text||'').split(/\s+/).filter(Boolean).length)],
  ];
  block.innerHTML = '<div class="bc-meta">' + rows.map(([k,v1,v2])=>{
    const same = (v1 === v2);
    return '<div class="bc-meta-row' + (same ? '' : ' different') + '">'
         + '<span class="bc-meta-label">' + k + '</span>'
         + '<span>' + v1 + '</span>'
         + '<span>' + v2 + '</span>'
         + '</div>';
  }).join('') + '</div>';

  // ISI diff
  const sum = document.getElementById('bc-isi-summary');
  sum.style.display = 'block';
  if(d.isi_diff.identical){
    sum.className   = 'diff-summary match';
    sum.textContent = '✓ ISI text is identical between the two banners.';
  } else {
    sum.className   = 'diff-summary diff';
    sum.textContent = '⚠ ' + d.isi_diff.diffs + ' word' + (d.isi_diff.diffs===1?'':'s')
                    + ' differ in the ISI text.';
  }
  document.getElementById('bc-isi-left').innerHTML  = d.isi_diff.left;
  document.getElementById('bc-isi-right').innerHTML = d.isi_diff.right;
}

async function bcDoVisualDiff(){
  const t = bcLinked ? bcA_time : bcA_time;  // use A's time as the reference
  document.getElementById('bc-vis-t').textContent = t.toFixed(2) + ' s';
  const result = document.getElementById('bc-vis-result');
  result.style.display = 'block';
  ['bc-vis-a','bc-vis-b','bc-vis-diff'].forEach(id => document.getElementById(id).removeAttribute('src'));

  const r = await busyFetch('/p3/visdiff', {
    method:'POST', headers:{'Content-Type':'application/json'},
    body: JSON.stringify({t: t}),
  }, 'Computing visual diff…');
  const d = await r.json();
  if(!d.ok){ alert(d.error || 'Visual diff failed'); return; }
  document.getElementById('bc-vis-a').src    = d.a;
  document.getElementById('bc-vis-b').src    = d.b;
  document.getElementById('bc-vis-diff').src = d.diff;
}

// ════════════════════════════════════════════════════════════════════════════
// TAB 4 — Multi-Banner Viewer
// ════════════════════════════════════════════════════════════════════════════
let mbSlots     = [];   // [{slot, name, w, h, dur, time}]
let mbDur       = 20;
let mbScrubbing = false;
let mbFrameCounter = 0;
let mbBundleName= 'klick_capture';

function mbStatus(msg, kind){
  const el = document.getElementById('mb-status');
  el.textContent = msg;
  el.className   = 'bc-status ' + (kind || '');
}
function mbAppendLog(t){
  const el = document.getElementById('mb-log');
  el.textContent += t;
  el.scrollTop = el.scrollHeight;
}

// ── Ziflow bookmarklet ────────────────────────────────────────────────────
// Build the bookmarklet href once on page load. The bookmarklet, when
// triggered on a Ziflow proof page, cycles through every banner using the
// same selectors the server-side Playwright code uses, collects the iframe
// srcs, and opens the Banner Tool with `?ml=<urls>#tab-multi`. Auto-loads
// without further user action.
function _bannerToolBookmarkletSource() {
  // Hard-code the deployed origin so the bookmarklet works no matter which
  // page initially provided it. (Falls back to current origin during local
  // dev so a developer hitting localhost gets a localhost-pointing bookmarklet.)
  const APP_HOST = (location.hostname.endsWith('.hf.space') || location.hostname.endsWith('huggingface.co'))
    ? location.origin
    : location.origin;
  // Body of the bookmarklet — wrapped in (async()=>{...})() so we can await.
  // Keep it compact; readable form lives next to this comment for easy edit.
  /* readable form:
     if (!/ziflow\.io\/proof\//.test(location.href)) {
       alert('Open a Ziflow proof page first, then click this bookmark.');
       return;
     }
     const sleep = ms => new Promise(r => setTimeout(r, ms));
     const outerIframe = () => document.querySelector('iframe[name^="ziflow-iframe-"]');
     const nextBtn = () => document.querySelector('a[data-selector="asset-switcher-next"]');
     const isDisabled = el => el && (el.className||'').toLowerCase().includes('disabled');

     // ── Mode A: "merged rich media" grid (multi-file proofs) ─────────────────
     // The outer iframe loads /api/proof/token/<id>/merged-rich-media which
     // itself contains one nested <iframe src=…/richMedia/<asset_guid>/…>
     // per banner. Same domain, so we can reach into contentDocument.
     function getMergedUrls() {
       const f = outerIframe();
       if (!f || !(f.src||'').includes('merged-rich-media')) return null;
       try {
         const inner = f.contentDocument;
         if (!inner) return null;
         const srcs = [...inner.querySelectorAll('iframe')]
           .map(i => i.src).filter(s => s && s.includes('richMedia'));
         return [...new Set(srcs)];
       } catch (e) { return null; }
     }

     // Wait up to 12s for either layout to populate.
     for (let t = 0; t < 40; t++) {
       const merged = getMergedUrls();
       if (merged && merged.length) break;
       const f = outerIframe();
       if (f && (f.src||'').includes('richMedia')) break;  // single-banner viewer
       await sleep(300);
     }

     const urls = [];
     const merged = getMergedUrls();
     if (merged && merged.length) {
       merged.forEach(u => { if (!urls.includes(u)) urls.push(u); });
     } else {
       // ── Mode B: 1-of-N single-banner viewer ─────────────────────────────
       const ifrSrc = () => {
         const f = outerIframe();
         return f && f.src ? f.src : '';
       };
       for (let i = 0; i < 30; i++) {
         let src = '';
         for (let t = 0; t < 40; t++) {
           src = ifrSrc();
           if (src && src.includes('richMedia')) break;
           await sleep(300);
         }
         if (src && !urls.includes(src)) urls.push(src);
         const nb = nextBtn();
         if (!nb || isDisabled(nb)) break;
         const prev = src;
         nb.click();
         let changed = false;
         for (let t = 0; t < 40; t++) {
           const x = ifrSrc();
           if (x && x !== prev) { changed = true; break; }
           await sleep(300);
         }
         if (!changed) break;
       }
     }

     if (!urls.length) {
       alert("Couldn't find any banners on this Ziflow proof.\n\n" +
             "Make sure: (1) you're on a proof page (URL contains /proof/<id>); " +
             "(2) you're signed in to Ziflow; (3) the proof has at least one banner.");
       return;
     }
     window.open(APP_HOST + '/?ml=' + encodeURIComponent(urls.join('\n')), '_blank');
  */
  const code =
    "(async()=>{if(!/ziflow\\.io\\/proof\\//.test(location.href)){alert('Open a Ziflow proof page first, then click this bookmark.');return;}" +
    "const sleep=ms=>new Promise(r=>setTimeout(r,ms));" +
    "const oI=()=>document.querySelector('iframe[name^=\"ziflow-iframe-\"]');" +
    "const nb=()=>document.querySelector('a[data-selector=\"asset-switcher-next\"]');" +
    "const dis=e=>e&&(e.className||'').toLowerCase().includes('disabled');" +
    "function gM(){const f=oI();if(!f||!(f.src||'').includes('merged-rich-media'))return null;" +
    "try{const ic=f.contentDocument;if(!ic)return null;" +
    "const ss=[...ic.querySelectorAll('iframe')].map(i=>i.src).filter(s=>s&&s.includes('richMedia'));" +
    "return[...new Set(ss)];}catch(e){return null;}}" +
    "for(let t=0;t<40;t++){const m=gM();if(m&&m.length)break;" +
    "const f=oI();if(f&&(f.src||'').includes('richMedia'))break;await sleep(300);}" +
    "const urls=[];const merged=gM();" +
    "if(merged&&merged.length){merged.forEach(u=>{if(!urls.includes(u))urls.push(u);});}" +
    "else{const ifr=()=>{const f=oI();return f&&f.src?f.src:'';};" +
    "for(let i=0;i<30;i++){let src='';" +
    "for(let t=0;t<40;t++){src=ifr();if(src&&src.includes('richMedia'))break;await sleep(300);}" +
    "if(src&&!urls.includes(src))urls.push(src);" +
    "const n=nb();if(!n||dis(n))break;" +
    "const prev=src;n.click();" +
    "let changed=false;for(let t=0;t<40;t++){const x=ifr();if(x&&x!==prev){changed=true;break;}await sleep(300);}" +
    "if(!changed)break;}}" +
    "if(!urls.length){alert(\"Couldn't find any banners on this Ziflow proof.\\n\\nMake sure: (1) you're on a proof page (URL contains /proof/<id>); (2) you're signed in to Ziflow; (3) the proof has at least one banner.\");return;}" +
    "window.open(" + JSON.stringify(APP_HOST) + "+'/?ml='+encodeURIComponent(urls.join('\\n')),'_blank');})();";
  return 'javascript:' + code;
}

document.addEventListener('DOMContentLoaded', () => {
  // Set the bookmarklet href once the host is known.
  const a = document.getElementById('mb-bookmarklet');
  if (a) a.href = _bannerToolBookmarkletSource();
});

async function mbCopyBookmarklet(){
  const code = _bannerToolBookmarkletSource();
  try{
    await navigator.clipboard.writeText(code);
    mbStatus('✓ Bookmarklet copied. Create a new bookmark in your browser, paste this as the URL, name it "Banner Tool: Get from Ziflow".', 'ok');
  }catch(e){
    // Fallback for browsers blocking clipboard without explicit user gesture.
    prompt('Copy this and paste as a new bookmark URL:', code);
  }
}

async function mbLoad(){
  // Note: #mb-url is a <textarea> so it can hold a single URL or multiple
  // newline-separated asset CDN URLs (the manual workaround for private
  // Ziflow proofs). The server's _multi_load auto-detects which mode.
  const url    = document.getElementById('mb-url').value.trim();
  const user   = document.getElementById('mb-user').value.trim();
  const pass   = document.getElementById('mb-pass').value;
  if(!url){ mbStatus('Paste a staging / proof URL (or asset URLs) first.', 'err'); return; }
  mbStatus('Signing in + downloading banners…');

  document.getElementById('mb-views').innerHTML = '';
  document.getElementById('mb-master').style.display = 'none';
  document.getElementById('mb-frame-list-wrap').style.display = 'none';

  let d;
  try{
    const r = await busyFetch('/multi/load', {
      method:'POST', headers:{'Content-Type':'application/json'},
      body: JSON.stringify({url, user, password: pass}),
    }, 'Signing in + downloading banners…');
    d = await r.json();
  }catch(e){
    mbStatus('⚠ ' + e.message, 'err'); return;
  }
  if(!d.ok){ mbStatus('⚠ ' + (d.error || 'Failed'), 'err'); return; }

  mbSlots = d.banners.map(b => ({...b, dur: 20, time: 0}));
  mbBundleName = d.bundle_name || 'klick_capture';
  let msg = '✓ Loaded ' + mbSlots.length + ' banner' + (mbSlots.length === 1 ? '' : 's');
  if(d.errors && d.errors.length) msg += ' (' + d.errors.length + ' failed)';
  mbStatus(msg, 'ok');

  const views = document.getElementById('mb-views');
  views.innerHTML = mbSlots.map(b => `
    <div class="mb-tile">
      <h4>
        <label style="display:inline-flex;align-items:center;gap:6px;cursor:pointer;font-weight:600;">
          <input type="checkbox" id="mb-sel-${b.slot}" checked>
          ${b.name} · ${b.width}×${b.height}
        </label>
      </h4>
      <div class="mb-shell" style="width:${b.width}px;height:${b.height}px;">
        <iframe id="mb-frame-${b.slot}" src="/banner_multi/${b.slot}/index.html"
                style="width:${b.width}px;height:${b.height}px;"></iframe>
      </div>
      <div class="mb-t" id="mb-time-${b.slot}">0.00 s</div>
    </div>
  `).join('');

  document.getElementById('mb-master').style.display = 'flex';
  document.getElementById('mb-frame-list-wrap').style.display = 'block';
}

function mbCmd(cmd, extra){
  mbSlots.forEach(b => {
    const f = document.getElementById('mb-frame-' + b.slot);
    try{ f.contentWindow.postMessage(Object.assign({bc: cmd}, extra || {}), '*'); }catch(e){}
  });
}
function mbOnScrub(val){
  // Proportional scrubbing — each banner seeks to the SAME percentage of its
  // own duration (so a 15s banner and a 22s banner both show the equivalent
  // moment in their animation, instead of the shorter one freezing at the end).
  const pct = val/1000;
  mbSlots.forEach(b => {
    const dur = b.dur || mbDur;
    const t   = pct * dur;
    const f   = document.getElementById('mb-frame-' + b.slot);
    try{ f.contentWindow.postMessage({bc:'seek', time: t}, '*'); }catch(e){}
  });
  document.getElementById('mb-time').textContent =
    't = ' + (pct * mbDur).toFixed(2) + ' s  (' + Math.round(pct*100) + '%)';
}

window.addEventListener('message', e => {
  if(!e.data || !e.data.bc) return;
  // Only handle messages from multi-banner iframes
  const slot = mbSlots.find(b => {
    const f = document.getElementById('mb-frame-' + b.slot);
    return f && e.source === f.contentWindow;
  });
  if(!slot) return;
  if(e.data.bc === 'ready'){
    slot.dur = e.data.duration || 20;
    mbDur = Math.max(20, ...mbSlots.map(s => s.dur || 20));
  }
  if(e.data.bc === 'state'){
    slot.time = e.data.time || 0;
    document.getElementById('mb-time-' + slot.slot).textContent = slot.time.toFixed(2) + ' s';
    if(!mbScrubbing && mbDur > 0){
      // Use the slot with the longest dur as scrubber reference (they're all in lockstep though)
      const ref = mbSlots.find(s => s.dur === mbDur) || slot;
      const pct = (ref.time || 0) / (ref.dur || mbDur);
      document.getElementById('mb-scrubber').value = Math.round(pct * 1000);
      document.getElementById('mb-time').textContent = 't = ' + (ref.time || 0).toFixed(2) + ' s';
    }
  }
});

function mbAddTime(){
  // Pick the master t (whatever the scrubber is showing — proportional to longest banner)
  const v = document.getElementById('mb-scrubber').value;
  const t = parseFloat(((v/1000)*mbDur).toFixed(2));
  mbFrameCounter++;
  mbAddFrameItem('t' + t.toString().replace('.','_') + 's', t);
}
function mbAddFrameItem(name, t){
  const id  = 'mbf-' + Date.now() + '-' + Math.random().toString(36).slice(2);
  const div = document.createElement('div');
  div.className = 'frame-item'; div.id = id;
  div.innerHTML = `
    <input type="text"   value="${name}" placeholder="label">
    <input type="number" value="${t}" step="0.1" min="0" style="width:70px;flex:none;" placeholder="t(s)">
    <button class="frame-del" onclick="document.getElementById('${id}').remove()" title="Remove">×</button>`;
  document.getElementById('mb-frame-list').appendChild(div);
}
function mbGetFrames(){
  return [...document.querySelectorAll('#mb-frame-list .frame-item')].map(el => {
    const inputs = el.querySelectorAll('input');
    return {name: inputs[0].value.trim() || 'frame', t: parseFloat(inputs[1].value) || 0};
  });
}

function mbSelectAll(checked){
  mbSlots.forEach(b => {
    const el = document.getElementById('mb-sel-' + b.slot);
    if(el) el.checked = checked;
  });
}

async function mbRunAll(){
  if(!mbSlots.length){ mbStatus('Load some banners first.', 'err'); return; }

  // Which banners are checked?
  const selectedSlots = mbSlots
    .filter(b => document.getElementById('mb-sel-' + b.slot).checked)
    .map(b => b.slot);
  if(selectedSlots.length === 0){
    mbStatus('Tick at least one banner above to capture.', 'err'); return;
  }

  const frames     = mbGetFrames();
  const doFrames   = document.getElementById('mb-do-frames').checked;
  const doIsiPng   = document.getElementById('mb-do-isi-png').checked;
  const doIsiHtml  = document.getElementById('mb-do-isi-html').checked;
  const doVideo    = document.getElementById('mb-do-video').checked;
  if(!doFrames && !doIsiPng && !doIsiHtml && !doVideo){
    mbStatus('Pick at least one export option (Frames, ISI image, ISI HTML, or MP4).', 'err'); return;
  }
  if(doFrames && frames.length === 0){
    mbStatus('Frames is checked but the list is empty — add a time first or uncheck Frames.', 'err'); return;
  }

  document.getElementById('mb-runlog-wrap').style.display = 'block';
  document.getElementById('mb-log').textContent = '';
  document.getElementById('mb-runstatus').textContent = 'Capturing ' + selectedSlots.length
    + ' banner' + (selectedSlots.length === 1 ? '' : 's') + '…';

  await busyFetch('/multi/capture', {
    method:'POST', headers:{'Content-Type':'application/json'},
    body: JSON.stringify({
      frames:      frames,
      slots:       selectedSlots,
      do_frames:   doFrames,
      do_isi_png:  doIsiPng,
      do_isi_html: doIsiHtml,
      do_video:    doVideo,
    }),
  }, 'Starting capture for ' + selectedSlots.length + ' banner'
   + (selectedSlots.length === 1 ? '' : 's') + '…');

  const es = new EventSource('/progress');
  es.onmessage = (e) => {
    const msg = JSON.parse(e.data);
    if(msg.type === 'log'){
      mbAppendLog(msg.text + '\n');
    } else if(msg.type === 'done'){
      mbAppendLog('\n✓ All banners captured.\n');
      // Hosted: surface a prominent download button. Replace the runstatus
      // text with a button so it's the focal point of the screen — the inline
      // log button was easy to miss. Also add a copy in the log for redundancy.
      const rs = document.getElementById('mb-runstatus');
      if(msg.zip){
        const zname = msg.zip.split('/').pop();
        const url   = '/download/' + encodeURIComponent(zname);
        rs.innerHTML = '';
        const head = document.createElement('span');
        head.textContent = 'Complete ✓ — ';
        head.style.cssText = 'color:var(--accent-success);font-weight:600;margin-right:8px;';
        rs.appendChild(head);
        const btn = document.createElement('button');
        btn.type = 'button';   // explicit so no <form> default-submit behaviour
        btn.className = 'btn btn-p btn-sm';
        btn.textContent = '⬇ Download zip (' + zname + ')';
        btn.addEventListener('click', (ev) => {
          ev.preventDefault();
          ev.stopPropagation();
          downloadFile(url, zname, 'application/zip', '.zip');
        });
        rs.appendChild(btn);
        // Inline log mention so the link is also discoverable from the log.
        mbAppendLog('Bundle zip ready: ' + zname + '\n');
      } else {
        rs.textContent = 'Complete ✓';
      }
      es.close();
    } else if(msg.type === 'error'){
      mbAppendLog('ERROR: ' + msg.text + '\n');
      document.getElementById('mb-runstatus').textContent = 'Error — see log';
      es.close();
    }
  };
}

// ════════════════════════════════════════════════════════════════════════════
// TAB 5 — Storyboard
// ════════════════════════════════════════════════════════════════════════════
let storyBundle  = null;   // path to the currently loaded bundle
let storyBanners = [];     // [{name, frames:[path], isi_png, notes:[str×4]}]
let storyAutoLoaded = false;

function storySetStatus(msg, kind){
  const el = document.getElementById('story-status');
  el.textContent = msg || '';
  el.className   = 'story-status ' + (kind || '');
}

async function storyAutoLoad(){
  // Pull the latest capture bundle from the server (if any).
  storySetStatus('Looking for latest capture…');
  let d;
  try{
    const r = await busyFetch('/storyboard/state',
              {method:'POST', headers:{'Content-Type':'application/json'}, body:'{}'},
              'Loading capture bundle…');
    d = await r.json();
  }catch(e){ storySetStatus('⚠ '+e.message, 'err'); return; }
  if(!d.ok || !d.bundle){
    storyRender(null, []);
    storySetStatus('No recent capture found — run a Multi-Banner capture first, or load a bundle.', '');
    return;
  }
  storyBundle  = d.bundle;
  storyBanners = (d.banners || []).map(b => Object.assign({notes: ['','','','']}, b));
  storyAutoLoaded = true;
  storyRender(d.bundle, storyBanners);
  storySetStatus('✓ Loaded bundle: '+d.name, 'ok');
  // If the Replace tab has pages loaded, re-render so dropdowns pick up the
  // newly-available banners and auto-suggest fires.
  if(replacePages.length) storyReplaceRender();
}

async function storyUploadBundle(input){
  // Hosted equivalent of the old folder-picker. Two paths:
  //   - .zip mode (input #story-bundle-zip-input):    one file, .zip
  //   - folder mode (input #story-bundle-folder-input): every file in the
  //     folder, with webkitRelativePath preserved as the filename
  if(!input.files || !input.files.length) return;
  const isFolder = input.hasAttribute('webkitdirectory');
  let totalBytes = 0;
  for(const f of input.files){ totalBytes += f.size; }
  const sizeMB = (totalBytes/(1024*1024)).toFixed(1);
  if(isFolder){
    storySetStatus('Uploading ' + input.files.length + ' files ('
                 + sizeMB + ' MB)…');
  } else {
    storySetStatus('Uploading ' + input.files[0].name + ' ('
                 + sizeMB + ' MB)…');
  }
  const fd = new FormData();
  for(const f of input.files){
    // webkitRelativePath preserves the folder structure on folder uploads;
    // for single .zip uploads it's empty so f.name is fine.
    const fname = f.webkitRelativePath || f.name;
    fd.append('file', f, fname);
  }
  let d;
  try{
    const r = await busyFetch('/upload/storyboard-bundle',
      {method:'POST', body: fd},
      isFolder ? 'Uploading + scanning folder…' : 'Uploading + extracting bundle…');
    d = await r.json();
  }catch(e){ storySetStatus('⚠ '+e.message, 'err'); return; }
  if(!d.ok){ storySetStatus('⚠ '+(d.error||'Upload failed'), 'err'); return; }
  storyBundle  = d.bundle;
  storyBanners = (d.banners || []).map(b => Object.assign({notes: ['','','','']}, b));
  storyRender(d.bundle, storyBanners);
  storySetStatus('✓ Loaded: '+d.name, 'ok');
  // Reset the input so re-selecting the same bundle re-triggers onchange.
  input.value = '';
}

// Update the bundle thumbnail strip in BOTH storyboard modes — visual confirmation
// that a capture bundle has been loaded and which banners are available.
function bundleStripRender(banners, bundleName){
  const html = (!banners || !banners.length)
    ? `<span class="bundle-empty">No capture bundle loaded yet.</span>`
    : (
        `<div class="bundle-title">✓ Loaded: ${bundleName||'capture bundle'}
           &nbsp;·&nbsp; ${banners.length} banner${banners.length===1?'':'s'}</div>`
        + banners.map(b => {
            // Pick the first frame as a thumb; if none, fall back to the ISI image.
            const thumbPath = (b.frames && b.frames[0]) || b.isi_png || '';
            const thumbSrc  = thumbPath
              ? `/storyboard/frame/${encodeURIComponent(thumbPath)}` : '';
            const fc = (b.frames || []).length;
            const ic = b.isi_png ? 1 : 0;
            return `
              <div class="bundle-pill" title="${b.name}">
                ${thumbSrc ? `<img src="${thumbSrc}" alt="${b.name}">` : ''}
                <div class="nm">${b.name}</div>
                <div class="ct">${fc} frame${fc===1?'':'s'}${ic ? ' · ISI' : ''}</div>
              </div>`;
          }).join('')
      );
  ['bundle-strip-gen', 'bundle-strip-rep'].forEach(id => {
    const el = document.getElementById(id);
    if(el) el.innerHTML = html;
  });
}

function storyRender(bundle, banners){
  // Update the bundle thumb strips first so they appear on both modes
  // even if no PDF is loaded yet for replace mode.
  bundleStripRender(banners, bundle ? bundle.split('/').pop() : null);

  const host = document.getElementById('story-banners');
  if(!banners || !banners.length){
    host.innerHTML = `<p class="hint" style="padding:24px 8px;color:var(--text-secondary);">
      Run a capture on the <strong>Multi-Banner Viewer</strong> tab first.
      When it finishes, come back here — the captured frames will appear automatically.</p>`;
    document.getElementById('story-actions').style.display = 'none';
    return;
  }
  host.innerHTML = banners.map((b, bi) => {
    const tiles = (b.frames || []).slice(0, 4).map((fpath, fi) => {
      const note = (b.notes && b.notes[fi]) || '';
      const safeNote = note.replace(/&/g,'&amp;').replace(/</g,'&lt;').replace(/>/g,'&gt;');
      return `
        <div class="story-frame">
          <div class="label">Frame ${fi+1}</div>
          <img src="/storyboard/frame/${encodeURIComponent(fpath)}" alt="Frame ${fi+1}">
          <textarea data-banner="${bi}" data-frame="${fi}"
            placeholder="Animation notes for Frame ${fi+1} of ${b.name}…"
            oninput="storyOnNoteInput(this)">${safeNote}</textarea>
        </div>`;
    }).join('');
    return `
      <div class="story-banner">
        <h3>${b.name}
          <span class="mute">${(b.frames||[]).length} frames captured</span>
        </h3>
        <div class="story-frames-grid">${tiles}</div>
      </div>`;
  }).join('');
  document.getElementById('story-actions').style.display = 'flex';
}

function storyOnNoteInput(el){
  const bi = parseInt(el.dataset.banner, 10);
  const fi = parseInt(el.dataset.frame,  10);
  if(!storyBanners[bi].notes) storyBanners[bi].notes = ['','','',''];
  storyBanners[bi].notes[fi] = el.value;
}

// Mode switcher (Generate vs Update existing PDF)
function storySetMode(mode){
  const isReplace = (mode === 'replace');
  document.getElementById('story-generate-mode').style.display = isReplace ? 'none' : '';
  document.getElementById('story-replace-mode').style.display  = isReplace ? '' : 'none';
  document.getElementById('story-mode-gen').classList.toggle('active', !isReplace);
  document.getElementById('story-mode-replace').classList.toggle('active', isReplace);
  // When entering replace mode, make sure a capture bundle is loaded so the
  // per-page dropdowns aren't empty. Re-render any already-loaded pages so
  // their dropdowns pick up the latest banners.
  if(isReplace){
    if(!storyBundle) storyAutoLoad();
    else if(replacePages.length) storyReplaceRender();
  }
}

async function storyBuild(){
  if(!storyBanners.length){ storySetStatus('No banners loaded.', 'err'); return; }
  storySetStatus('Generating PDF (one banner per page)…');
  document.getElementById('story-download-link').style.display = 'none';

  const header = document.getElementById('story-header').value;
  const footer = document.getElementById('story-footer').value;

  const payload = {
    header, footer,
    banners: storyBanners.map(b => ({
      name:   b.name,
      frames: (b.frames || []).slice(0, 4),
      notes:  (b.notes  || ['','','','']).slice(0, 4),
    })),
  };
  let d;
  try{
    const r = await busyFetch('/storyboard/build', {method:'POST',
              headers:{'Content-Type':'application/json'},
              body: JSON.stringify(payload)},
              'Generating storyboard PDF…');
    d = await r.json();
  }catch(e){ storySetStatus('⚠ '+e.message, 'err'); return; }
  if(!d.ok){ storySetStatus('⚠ '+(d.error||'Failed'), 'err'); return; }
  storySetStatus('✓ PDF saved to bundle.', 'ok');
  // Reveal the download button — user clicks to get the Save-As dialog.
  document.getElementById('story-download-link').style.display = '';
  // Suggest the bundle name as the default filename.
  window._lastStoryboardName = (storyBanners && storyBundle
                                  ? storyBundle.split('/').pop()
                                  : 'storyboard') + '.pdf';
  // Auto-trigger the download so they get immediate feedback (Save-As dialog
  // pops up where supported).
  storyDownloadLastPdf();
}

async function storyDownloadLastPdf(){
  const url = '/storyboard/download?t=' + Date.now();
  const name = window._lastStoryboardName || 'storyboard.pdf';
  await downloadFile(url, name, 'application/pdf', '.pdf');
}

// ════════════════════════════════════════════════════════════════════════════
// TAB 5b — Storyboard "Update existing PDF" mode
// ════════════════════════════════════════════════════════════════════════════
let replacePages = [];   // [{page_num, thumb, regions, page_w, page_h, ...}]

function replaceSetStatus(msg, kind){
  const el = document.getElementById('story-replace-status');
  el.textContent = msg || '';
  el.className   = 'story-status ' + (kind || '');
}

async function storyReplaceUpload(input){
  if(!input.files || !input.files.length){ return; }
  const file = input.files[0];
  replaceSetStatus('Uploading ' + file.name + '…');
  let d;
  try{
    // Step 1: upload the PDF
    const fd = new FormData(); fd.append('file', file, file.name);
    const u = await busyFetch('/upload/storyboard-existing', {method:'POST', body: fd},
                              'Uploading PDF…');
    const ud = await u.json();
    if(!ud.ok){ replaceSetStatus('⚠ '+(ud.error||'Upload failed'), 'err'); return; }
    // Step 2: ask the server to render + detect regions
    const r = await busyFetch('/storyboard/upload', {method:'POST',
              headers:{'Content-Type':'application/json'},
              body: JSON.stringify({filename: ud.filename})},
              'Reading PDF and detecting regions…');
    d = await r.json();
  }catch(e){ replaceSetStatus('⚠ '+e.message, 'err'); return; }
  if(!d.ok){ replaceSetStatus('⚠ '+(d.error||'Failed'), 'err'); return; }

  document.getElementById('story-replace-name').textContent = '✓ ' + d.name;
  document.getElementById('story-replace-name').className   = 'compare-pill ok';
  replacePages = d.pages;
  // Stash each detected region's original bounds so the server can wipe both
  // the original AND the user-edited rect when the new image gets pasted —
  // kills any leftover frame edges if the user moved/resized the region.
  replacePages.forEach(p => {
    (p.regions || []).forEach(r => {
      r._original = {x: r.x, y: r.y, w: r.w, h: r.h};
    });
    // Cache the pristine page thumbnail so the live-preview canvas can render
    // from a clean base on every drag (server preview thumbs accumulate
    // replacements, which we want to *override* live).
    p._original_thumb = p.thumb;
  });
  replaceSetStatus('✓ Loaded ' + d.pages.length + ' page'
                   + (d.pages.length === 1 ? '' : 's'), 'ok');
  storyReplaceRender();
}

// ── Region editor ─────────────────────────────────────────────────────────────
// Each detected region renders on a per-page canvas with 8 drag handles. Click
// inside a region to drag-move; click on a corner/edge handle to resize. All
// region coords are stored in PAGE pixel space (matches what _detect_banner_regions
// returned), and rendered scaled to thumb space.
// Region-identity colors — kept as plain hex (NOT theme variables) for two
// reasons: (1) the canvas uses the `col + '22'` 8-digit-hex alpha trick to
// draw a 13% translucent fill, which only works on hex strings (a `var(--…)`
// concat would produce invalid CSS and the canvas falls back to solid black),
// and (2) region identity should stay visually consistent across themes —
// "R1" is always pink, "R2" always green, etc.
const REGION_COLORS = ['#f38ba8','#a6e3a1','#f9b34a','#89b4fa','#cba6f7','#fab387'];
const HANDLE_SIZE   = 8;    // px in thumb space (visual size)
const HANDLE_HIT    = 10;   // hit-test radius
const MIN_REGION_PX = 16;   // minimum region size in PAGE coords

function _regionColor(idx){ return REGION_COLORS[idx % REGION_COLORS.length]; }

// The image-options HTML is identical for every dropdown; rebuild only when
// the bundle changes. Each option's value encodes "(frame|isi):<banner>:<path>"
// — the prefix is just for client-side debug; the server only consumes the path.
let _imageOptionsHtml = '<option value="">— No bundle loaded —</option>';
function rebuildImageOptions(){
  if(!storyBanners || !storyBanners.length){
    _imageOptionsHtml = '<option value="">— No bundle loaded —</option>';
    return;
  }
  const groups = storyBanners.map(b => {
    const items = [];
    (b.frames || []).forEach((fp) => {
      const fname = fp.split('/').pop();
      const label = fname.replace(/\.png$/i,'');
      items.push(`<option value="${encodeURIComponent(fp)}" data-kind="frame" data-banner="${b.name}">${label}</option>`);
    });
    if(b.isi_png){
      items.push(`<option value="${encodeURIComponent(b.isi_png)}" data-kind="isi" data-banner="${b.name}">ISI (full)</option>`);
    }
    return `<optgroup label="${b.name}">${items.join('')}</optgroup>`;
  }).join('');
  _imageOptionsHtml = `<option value="">— Don't change this region —</option>` + groups;
}

// Walk the bundle to suggest a sensible default for a freshly-loaded region.
// Banner regions → i-th captured frame from the size-matched banner.
// ISI regions    → that banner's isi_full.png.
function _autoSuggestForRegion(page, regionIdx, region){
  if(!storyBanners || !storyBanners.length) return '';
  if(region.kind === 'isi'){
    let b = page.suggested_banner_size
            ? storyBanners.find(x => x.name === page.suggested_banner_size)
            : null;
    if(!b) b = storyBanners.find(x => x.isi_png);
    if(b && b.isi_png) return encodeURIComponent(b.isi_png);
    return '';
  }
  const target = page.suggested_banner_size
                 ? storyBanners.find(x => x.name === page.suggested_banner_size)
                 : storyBanners[0];
  if(target && target.frames && target.frames.length){
    const fr = target.frames[Math.min(regionIdx, target.frames.length - 1)];
    return encodeURIComponent(fr);
  }
  return '';
}

function storyReplaceRender(){
  rebuildImageOptions();
  const host = document.getElementById('story-replace-pages');
  if(!replacePages.length){
    host.innerHTML = `<p class="hint" style="padding:24px 8px;color:var(--text-secondary);">
      Upload an existing storyboard PDF to begin.</p>`;
    document.getElementById('story-replace-actions').style.display = 'none';
    return;
  }

  // Container for the per-chain ISI slicing editors (rendered after the
  // page list updates, see storyRenderIsiChains below).
  const isiHost = document.getElementById('isi-chain-editors')
                || (() => {
                     const h = document.createElement('div');
                     h.id = 'isi-chain-editors';
                     host.parentNode.insertBefore(h, host);
                     return h;
                   })();

  host.innerHTML = replacePages.map(p => {
    const regions = p.regions || [];
    const isIsi   = regions.length === 1 && regions[0].kind === 'isi';
    const isBanner = regions.length >= 1 && !isIsi;
    let detected;
    if (isBanner) {
      detected = `<span class="replace-detected">${regions.length} region${regions.length===1?'':'s'} detected`
                 + (p.suggested_banner_size ? ' — looks like '+p.suggested_banner_size : '')
                 + ' · drag handles to resize, drag inside to move</span>';
    } else if (isIsi) {
      detected = `<span class="replace-detected">ISI page · drag handles to resize the replacement region</span>`;
    } else {
      detected = `<span class="replace-detected warn">No regions detected on this page — nothing to map.</span>`;
    }

    const pickers = regions.map((r, ri) => `
      <div class="region-picker">
        <span class="swatch" style="background:${_regionColor(ri)};"></span>
        <span class="lbl">Region ${ri+1}</span>
        <select class="region-select" data-page="${p.page_num}" data-region="${ri}">
          ${_imageOptionsHtml}
        </select>
        <button type="button" class="region-snap-btn"
                title="Snap this region to the non-white content inside it"
                onclick="storySnapRegion(${p.page_num}, ${ri})">✂ Snap</button>
        <button type="button" class="region-del-btn"
                title="Remove this region"
                onclick="storyRemoveRegion(${p.page_num}, ${ri})">×</button>
      </div>`).join('');

    return `
      <div class="replace-page" data-page="${p.page_num}">
        <div class="replace-thumb-wrap" style="width:${p.thumb_w}px;height:${p.thumb_h}px;">
          <img src="${p.thumb}" alt="page ${p.page_num}" width="${p.thumb_w}" height="${p.thumb_h}" draggable="false">
          <canvas class="region-canvas" width="${p.thumb_w}" height="${p.thumb_h}"
                  data-page="${p.page_num}"></canvas>
        </div>
        <div class="replace-controls">
          <div class="replace-page-meta">Page ${p.page_num}${isIsi ? ' · ISI' : ''}</div>
          ${detected}
          <div class="region-pickers">
            ${pickers || '<span class="replace-detected warn">No regions yet.</span>'}
          </div>
          <button type="button" class="btn btn-g btn-sm region-add-btn"
                  onclick="storyAddRegion(${p.page_num})">+ Add region</button>
        </div>
      </div>`;
  }).join('');

  // Apply auto-suggest defaults and wire dropdown change handlers.
  replacePages.forEach(p => {
    (p.regions || []).forEach((r, ri) => {
      const sel = host.querySelector(
        `select.region-select[data-page="${p.page_num}"][data-region="${ri}"]`);
      if(!sel) return;
      const initial = r.image_path || _autoSuggestForRegion(p, ri, r);
      if(initial){
        sel.value = initial;
        r.image_path = initial;
      }
      sel.addEventListener('change', () => {
        r.image_path = sel.value || null;
        // If this region is now (or was) part of an ISI chain, rebuild the
        // slicing editor — the set of chained regions may have changed.
        storyRenderIsiChains();
      });
    });
  });

  // Draw overlays + bind drag handlers.
  replacePages.forEach(p => bindRegionEditor(p));

  // ISI slicing editor (only visible when ≥2 regions share the same isi_full.png).
  storyRenderIsiChains();

  document.getElementById('story-replace-actions').style.display = 'flex';
}

function bindRegionEditor(page){
  const host = document.getElementById('story-replace-pages');
  const canvas = host.querySelector(
    `canvas.region-canvas[data-page="${page.page_num}"]`);
  if(!canvas) return;
  const ctx = canvas.getContext('2d');
  const sx = page.thumb_w / page.page_w;   // page-px → thumb-px
  const sy = page.thumb_h / page.page_h;

  function rectInThumb(r){
    return [r.x * sx, r.y * sy, r.w * sx, r.h * sy];
  }
  function handlesFor(r){
    const [x,y,w,h] = rectInThumb(r);
    return {
      'nw':[x,y],         'n':[x+w/2,y],   'ne':[x+w,y],
      'w' :[x,y+h/2],                       'e' :[x+w,y+h/2],
      'sw':[x,y+h],       's':[x+w/2,y+h], 'se':[x+w,y+h],
    };
  }
  function draw(){
    ctx.clearRect(0,0,canvas.width,canvas.height);
    (page.regions || []).forEach((r, idx) => {
      const [x,y,w,h] = rectInThumb(r);
      const col = _regionColor(idx);
      ctx.lineWidth = 2;
      ctx.strokeStyle = col;
      ctx.fillStyle   = col + '22';   // ~13% alpha tint
      ctx.fillRect(x,y,w,h);
      ctx.strokeRect(x,y,w,h);
      // Handles
      const hs = handlesFor(r);
      ctx.fillStyle   = col;
      ctx.strokeStyle = '#fff';
      ctx.lineWidth   = 1;
      Object.values(hs).forEach(([hx,hy]) => {
        ctx.fillRect(hx-HANDLE_SIZE/2, hy-HANDLE_SIZE/2, HANDLE_SIZE, HANDLE_SIZE);
        ctx.strokeRect(hx-HANDLE_SIZE/2, hy-HANDLE_SIZE/2, HANDLE_SIZE, HANDLE_SIZE);
      });
      // Region number badge above the top-left corner.
      const tag = 'R' + (idx + 1);
      ctx.fillStyle = col;
      ctx.fillRect(x, Math.max(0, y - 16), 26, 14);
      ctx.fillStyle = '#0d0d18';
      ctx.font = 'bold 10px Menlo,monospace';
      ctx.fillText(tag, x + 4, Math.max(10, y - 5));
    });
  }
  page._draw = draw;
  draw();

  function hitTest(mx, my){
    // Walk regions back-to-front so an overlapping region's handle wins.
    for(let i = (page.regions||[]).length - 1; i >= 0; i--){
      const r = page.regions[i];
      const hs = handlesFor(r);
      for(const k of Object.keys(hs)){
        const [hx,hy] = hs[k];
        if(Math.abs(mx - hx) <= HANDLE_HIT && Math.abs(my - hy) <= HANDLE_HIT){
          return {idx:i, mode:'resize-'+k};
        }
      }
      const [x,y,w,h] = rectInThumb(r);
      if(mx >= x && mx <= x+w && my >= y && my <= y+h){
        return {idx:i, mode:'move'};
      }
    }
    return null;
  }
  function cursorFor(mode){
    if(!mode) return 'default';
    if(mode === 'move') return 'move';
    const dir = mode.replace('resize-','');
    return ({n:'ns',s:'ns',e:'ew',w:'ew',ne:'nesw',sw:'nesw',nw:'nwse',se:'nwse'})[dir] + '-resize';
  }

  let drag = null;
  canvas.addEventListener('mousedown', e => {
    e.preventDefault();
    const rect = canvas.getBoundingClientRect();
    const mx = e.clientX - rect.left;
    const my = e.clientY - rect.top;
    const hit = hitTest(mx, my);
    if(!hit) return;
    drag = {
      idx: hit.idx, mode: hit.mode,
      startMX: mx, startMY: my,
      startRegion: Object.assign({}, page.regions[hit.idx]),
    };
  });
  canvas.addEventListener('mousemove', e => {
    const rect = canvas.getBoundingClientRect();
    const mx = e.clientX - rect.left;
    const my = e.clientY - rect.top;
    if(!drag){
      const h = hitTest(mx, my);
      canvas.style.cursor = cursorFor(h && h.mode);
      return;
    }
    const dx = (mx - drag.startMX) / sx;   // thumb → page coords
    const dy = (my - drag.startMY) / sy;
    const r0 = drag.startRegion;
    const r  = page.regions[drag.idx];
    if(drag.mode === 'move'){
      r.x = Math.max(0, Math.min(page.page_w - r0.w, r0.x + dx));
      r.y = Math.max(0, Math.min(page.page_h - r0.h, r0.y + dy));
    } else {
      const dir = drag.mode.replace('resize-','');
      let x1 = r0.x, y1 = r0.y, x2 = r0.x + r0.w, y2 = r0.y + r0.h;
      if(dir.includes('n')) y1 = r0.y + dy;
      if(dir.includes('s')) y2 = r0.y + r0.h + dy;
      if(dir.includes('w')) x1 = r0.x + dx;
      if(dir.includes('e')) x2 = r0.x + r0.w + dx;
      if(x2 - x1 < MIN_REGION_PX){
        if(dir.includes('w')) x1 = x2 - MIN_REGION_PX;
        else                  x2 = x1 + MIN_REGION_PX;
      }
      if(y2 - y1 < MIN_REGION_PX){
        if(dir.includes('n')) y1 = y2 - MIN_REGION_PX;
        else                  y2 = y1 + MIN_REGION_PX;
      }
      x1 = Math.max(0, Math.min(page.page_w, x1));
      x2 = Math.max(0, Math.min(page.page_w, x2));
      y1 = Math.max(0, Math.min(page.page_h, y1));
      y2 = Math.max(0, Math.min(page.page_h, y2));
      r.x = x1; r.y = y1; r.w = x2 - x1; r.h = y2 - y1;
    }
    draw();
  });
  const endDrag = () => { drag = null; canvas.style.cursor = 'default'; };
  canvas.addEventListener('mouseup',    endDrag);
  canvas.addEventListener('mouseleave', endDrag);
}

// Add a fresh region to a page in PAGE coords (default 25%×20% of page,
// centered). User can resize/move via handles afterwards.
function storyAddRegion(pageNum){
  const page = replacePages.find(p => p.page_num === pageNum);
  if(!page) return;
  const w = Math.max(60, Math.round(page.page_w * 0.25));
  const h = Math.max(60, Math.round(page.page_h * 0.20));
  const x = Math.round((page.page_w - w) / 2);
  const y = Math.round((page.page_h - h) / 2);
  page.regions = page.regions || [];
  page.regions.push({x, y, w, h, kind: 'banner', image_path: null});
  storyReplaceRender();
}

function storyRemoveRegion(pageNum, regionIdx){
  const page = replacePages.find(p => p.page_num === pageNum);
  if(!page || !page.regions) return;
  page.regions.splice(regionIdx, 1);
  storyReplaceRender();
}

// ── ISI chain slicing editor ─────────────────────────────────────────────────
// When 2+ regions across the document have the same isi_full.png assigned,
// render an editor showing the ISI image with a colored band per region and
// draggable horizontal dividers between bands. The user can drag dividers
// to manually set where each region's slice begins/ends, or click "Reset to
// auto" to fall back to height-weighted automatic boundaries. On confirm,
// the live preview re-renders so the user sees their adjusted slices.

function _safeDecode(s){
  // Dropdown values are produced by encodeURIComponent(absolute_path) and
  // stored on r.image_path that way. Always decode before pattern-matching
  // or comparing — the regex `/isi_full\.png$/` won't see a literal `/` in
  // the encoded form (it has `%2F` instead).
  try { return decodeURIComponent(s); } catch(e) { return s; }
}

function _findIsiChains(){
  // Returns: { '<decoded-path>': [{pi, ri, height}, ...] } for chains of ≥2.
  // Groups keyed by the *decoded* path so subsequent comparisons + URL building
  // stay consistent.
  const groups = {};
  replacePages.forEach((p, pi) => {
    (p.regions || []).forEach((r, ri) => {
      if (!r.image_path) return;
      const decoded = _safeDecode(r.image_path);
      const isIsi   = /(?:^|\/)isi_full\.png$/i.test(decoded);
      if (!isIsi) return;
      groups[decoded] = groups[decoded] || [];
      groups[decoded].push({pi, ri, h: Math.max(1, Math.round(r.h || 1))});
    });
  });
  Object.keys(groups).forEach(k => { if (groups[k].length < 2) delete groups[k]; });
  return groups;
}

function _computeAutoSlices(members){
  // Height-weighted boundaries — matches the server-side fallback.
  const total = members.reduce((s, m) => s + m.h, 0) || 1;
  let cum = 0;
  return members.map(m => {
    const top = cum / total;
    cum += m.h;
    const bot = cum / total;
    return {top, bot};
  });
}

function storyRenderIsiChains(){
  // Render one "Edit ISI slices" button per chain. The actual slicing UI
  // lives in the modal (see openIsiModal), keeping the page list compact.
  const host = document.getElementById('isi-chain-editors');
  if (!host) return;
  const chains = _findIsiChains();
  const paths  = Object.keys(chains);
  if (!paths.length) { host.innerHTML = ''; return; }

  // Persist auto fractions for any region that doesn't already have them so
  // the modal opens with a sensible default before the first drag.
  paths.forEach(path => {
    const members = chains[path];
    const auto    = _computeAutoSlices(members);
    members.forEach((m, i) => {
      const r = replacePages[m.pi].regions[m.ri];
      if (typeof r.isi_top !== 'number' || typeof r.isi_bot !== 'number'){
        r.isi_top = auto[i].top;
        r.isi_bot = auto[i].bot;
      }
    });
  });

  host.innerHTML = paths.map(path => {
    const members  = chains[path];
    const segments = path.split('/');
    // The slot folder name is the second-to-last segment
    // (e.g. "/Users/.../300x250/isi_full.png" -> "300x250").
    const slotName = segments.length >= 2 ? segments[segments.length - 2] : 'ISI';
    return `<button type="button" class="btn btn-g isi-chain-open"
                    onclick="openIsiModal('${encodeURIComponent(path)}')"
                    title="Open the ISI slicing editor for this chain">
              ✂ Edit ISI slices: <strong>${slotName}</strong>
              <span class="muted">(${members.length} regions)</span>
            </button>`;
  }).join('');
}

// ── Modal: open / close / zoom / render / drag ──────────────────────────────
let _currentIsiPath  = null;
let _currentIsiZoom  = null;   // null = "fit"; otherwise a multiplier vs natural
let _currentIsiPage  = null;   // page_num currently shown in the live preview
let _isiImgCache     = {};     // url -> HTMLImageElement (loaded)
let _isiSplitterRaf  = 0;

function openIsiModal(encodedPath){
  _currentIsiPath = _safeDecode(encodedPath);
  _currentIsiZoom = null;     // fit by default
  const modal   = document.getElementById('isi-modal');
  const slot    = (_currentIsiPath.split('/').slice(-2, -1)[0]) || 'ISI';
  document.getElementById('isi-modal-title').textContent = slot;
  document.getElementById('isi-modal-zoom-label').textContent = 'Fit';

  // Pick the first page in the chain as the initial active tab.
  const members  = _findIsiChains()[_currentIsiPath] || [];
  const pageNums = [...new Set(members.map(m => replacePages[m.pi].page_num))]
                   .sort((a, b) => a - b);
  _currentIsiPage = pageNums.length ? pageNums[0] : null;

  modal.hidden = false;
  _wireIsiSplitter();
  isiModalRender();
}

function closeIsiModal(){
  document.getElementById('isi-modal').hidden = true;
  _currentIsiPath = null;
  _currentIsiPage = null;
}

function _wireIsiSplitter(){
  const splitter = document.getElementById('isi-modal-splitter');
  if (!splitter || splitter._wired) return;
  splitter._wired = true;
  splitter.addEventListener('mousedown', ev => {
    ev.preventDefault();
    splitter.classList.add('dragging');
    const body  = splitter.parentElement;
    const left  = document.getElementById('isi-modal-left');
    const right = document.getElementById('isi-modal-right');
    function onMove(e){
      const rect = body.getBoundingClientRect();
      const ratio = (e.clientX - rect.left) / rect.width;
      const clamped = Math.max(0.18, Math.min(0.82, ratio));
      left.style.flex  = `1 1 ${clamped * 100}%`;
      right.style.flex = `1 1 ${(1 - clamped) * 100}%`;
      // Re-render the preview canvas at the new size (debounced via rAF).
      cancelAnimationFrame(_isiSplitterRaf);
      _isiSplitterRaf = requestAnimationFrame(isiModalRenderLivePreview);
    }
    function onUp(){
      splitter.classList.remove('dragging');
      document.removeEventListener('mousemove', onMove);
      document.removeEventListener('mouseup', onUp);
    }
    document.addEventListener('mousemove', onMove);
    document.addEventListener('mouseup', onUp);
  });
}

function isiModalZoom(action){
  // 10% steps from 10% to 500%, plus an explicit "Fit" mode.
  if (action === 'fit') {
    _currentIsiZoom = null;
  } else {
    let cur = (_currentIsiZoom === null) ? 1.0 : _currentIsiZoom;
    if (action === 'in')  cur = Math.min(5.0, Math.round((cur + 0.1) * 100) / 100);
    if (action === 'out') cur = Math.max(0.1, Math.round((cur - 0.1) * 100) / 100);
    _currentIsiZoom = cur;
  }
  isiModalApplyZoom();
}

function isiModalApplyZoom(){
  // Strip image lives inside #isi-modal-strip-wrap (the right pane). The old
  // selector referenced #isi-modal-body which was removed in the split-pane
  // redesign — that's why zoom looked broken.
  const img   = document.querySelector('#isi-modal-strip-wrap .isi-strip-img');
  const label = document.getElementById('isi-modal-zoom-label');
  if (!label) return;
  if (!img) { label.textContent = 'Fit'; return; }
  if (_currentIsiZoom === null) {
    img.style.height    = '';
    img.style.maxHeight = '70vh';
    label.textContent   = 'Fit';
  } else if (img.naturalHeight) {
    img.style.maxHeight = 'none';
    img.style.height    = (img.naturalHeight * _currentIsiZoom) + 'px';
    label.textContent   = Math.round(_currentIsiZoom * 100) + '%';
  }
}

function isiModalRender(){
  if (!_currentIsiPath) return;
  const chains  = _findIsiChains();
  const members = chains[_currentIsiPath];

  // ── Page tabs ────────────────────────────────────────────────────────
  const tabsHost = document.getElementById('isi-modal-tabs');
  if (members && members.length){
    const pageNums = [...new Set(members.map(m => replacePages[m.pi].page_num))]
                     .sort((a, b) => a - b);
    if (_currentIsiPage === null || !pageNums.includes(_currentIsiPage)){
      _currentIsiPage = pageNums[0];
    }
    tabsHost.innerHTML = pageNums.map(pn => {
      const cls = (pn === _currentIsiPage) ? 'isi-modal-tab active' : 'isi-modal-tab';
      return `<button type="button" class="${cls}"
                onclick="isiModalSelectPage(${pn})">Page ${pn}</button>`;
    }).join('');
  } else {
    tabsHost.innerHTML = '';
  }

  // ── Right pane: ISI strip + bands + handles ─────────────────────────
  const stripWrap = document.getElementById('isi-modal-strip-wrap');
  if (!members || !members.length){
    stripWrap.innerHTML = `<p style="color:var(--text-secondary);padding:20px;">
      This ISI is no longer assigned to multiple regions.</p>`;
  } else {
    const imgUrl = '/storyboard/isi-image?path=' + encodeURIComponent(_currentIsiPath);
    const bands = members.map((m, i) => {
      const r = replacePages[m.pi].regions[m.ri];
      const top = r.isi_top * 100, bot = r.isi_bot * 100;
      const c = _regionColor(m.ri);
      const label = `Page ${replacePages[m.pi].page_num} · R${m.ri+1}`;
      return `<div class="isi-band" style="top:${top}%;height:${bot-top}%;
                background:${c}33;border-top:2px solid ${c};border-bottom:2px solid ${c};"
                data-member-idx="${i}" title="${label}">
                <span class="isi-band-label" style="color:${c};">${label}</span>
              </div>`;
    }).join('');
    const handles = members.map((m, i) => {
      const r = replacePages[m.pi].regions[m.ri];
      const c = _regionColor(m.ri);
      return `
        <div class="isi-handle isi-handle-top" data-member-idx="${i}" data-edge="top"
             style="top:${r.isi_top * 100}%;background:${c};color:${c};">
          <span class="isi-handle-label">R${m.ri+1} top</span>
        </div>
        <div class="isi-handle isi-handle-bot" data-member-idx="${i}" data-edge="bot"
             style="top:${r.isi_bot * 100}%;background:${c};color:${c};">
          <span class="isi-handle-label">R${m.ri+1} bot</span>
        </div>`;
    }).join('');
    stripWrap.innerHTML = `
      <div class="isi-modal-strip">
        <img src="${imgUrl}" alt="ISI" class="isi-strip-img" draggable="false"
             onload="isiModalApplyZoom()">
        ${bands}
        ${handles}
      </div>`;
    _wireIsiModalHandles();
  }

  // ── Left pane: live preview canvas ─────────────────────────────────
  isiModalRenderLivePreview();
}

function isiModalSelectPage(pageNum){
  _currentIsiPage = pageNum;
  // Refresh tabs + live preview (slicer doesn't change — chain is global).
  document.querySelectorAll('#isi-modal-tabs .isi-modal-tab').forEach(b => {
    b.classList.toggle('active', b.textContent === `Page ${pageNum}`);
  });
  isiModalRenderLivePreview();
}

// Cache loaded images so live drag stays smooth (no reload per frame).
function _isiLoadImg(url){
  if (_isiImgCache[url] && _isiImgCache[url].complete) {
    return Promise.resolve(_isiImgCache[url]);
  }
  return new Promise((resolve, reject) => {
    const im = new Image();
    im.crossOrigin = 'anonymous';
    im.onload  = () => { _isiImgCache[url] = im; resolve(im); };
    im.onerror = () => reject(new Error('Failed to load ' + url));
    im.src = url;
  });
}

// Render the live preview canvas: page thumbnail (clean upload) + every
// region's current replacement applied client-side. Mirrors the server's
// _apply_per_region_mappings logic so what you see is what export produces.
async function isiModalRenderLivePreview(){
  const canvas = document.getElementById('isi-modal-preview-canvas');
  if (!canvas || _currentIsiPage === null) return;
  const pi = replacePages.findIndex(p => p.page_num === _currentIsiPage);
  if (pi < 0) return;
  const page = replacePages[pi];

  // Use the pristine upload thumbnail as the base — the server-preview thumb
  // may have stale chain content from a previous Apply.
  const baseThumb = page._original_thumb || page.thumb;
  let pageImg;
  try { pageImg = await _isiLoadImg(baseThumb); }
  catch(e){ console.warn(e); return; }

  canvas.width  = pageImg.naturalWidth;
  canvas.height = pageImg.naturalHeight;
  const ctx = canvas.getContext('2d');
  ctx.drawImage(pageImg, 0, 0);

  // Page coords → thumbnail pixels.
  const scale = pageImg.naturalWidth / page.page_w;

  // Walk every region on this page; for those with image_path, draw the
  // replacement live. Frame regions paste whole; ISI regions paste a slice.
  for (const r of (page.regions || [])){
    if (!r.image_path) continue;
    const decodedPath = _safeDecode(r.image_path);
    let img;
    try { img = await _isiLoadImg('/storyboard/isi-image?path=' + encodeURIComponent(decodedPath)); }
    catch(e){ continue; }

    // Compute kind first so we can skip the original_rect wipe for ISI
    // regions (auto-detect often over-includes page titles/surrounding text;
    // wiping that erases content the user wants to keep). Frames still get
    // both wipes so moved frames don't leave ghost edges.
    const isIsi = /(?:^|\/)isi_full\.png$/i.test(decodedPath);
    ctx.fillStyle = '#fff';
    if (r._original && !isIsi){
      ctx.fillRect(Math.round(r._original.x * scale), Math.round(r._original.y * scale),
                   Math.round(r._original.w * scale), Math.round(r._original.h * scale));
    }
    const rx = Math.round(r.x * scale), ry = Math.round(r.y * scale);
    const rw = Math.round(r.w * scale), rh = Math.round(r.h * scale);
    ctx.fillRect(rx, ry, rw, rh);

    // Compute slice from source image — ISI honors top/bot fractions.
    let srcTop = 0, srcBot = img.naturalHeight;
    if (isIsi && typeof r.isi_top === 'number' && typeof r.isi_bot === 'number'){
      srcTop = Math.max(0, Math.round(img.naturalHeight * r.isi_top));
      srcBot = Math.min(img.naturalHeight, Math.max(srcTop + 1, Math.round(img.naturalHeight * r.isi_bot)));
    }
    const sliceW = img.naturalWidth, sliceH = srcBot - srcTop;
    if (sliceH <= 0 || sliceW <= 0) continue;

    // Scale-to-fit (preserve aspect) inside the user's region rect.
    const fit = Math.min(rw / sliceW, rh / sliceH);
    const dw  = Math.max(1, Math.round(sliceW * fit));
    const dh  = Math.max(1, Math.round(sliceH * fit));
    const dx  = rx + Math.round((rw - dw) / 2);
    // ISI top-aligns; frames center vertically.
    const dy  = isIsi ? ry : ry + Math.round((rh - dh) / 2);

    ctx.drawImage(img, 0, srcTop, sliceW, sliceH, dx, dy, dw, dh);

    if (isIsi){
      ctx.strokeStyle = '#aaa';
      ctx.lineWidth   = 1;
      ctx.strokeRect(dx + 0.5, dy + 0.5, dw - 1, dh - 1);
    }
  }
}

function _wireIsiModalHandles(){
  const stripWrap = document.getElementById('isi-modal-strip-wrap');
  if (!stripWrap) return;
  const strip = stripWrap.querySelector('.isi-modal-strip');
  if (!strip) return;
  const members = _findIsiChains()[_currentIsiPath] || [];

  stripWrap.querySelectorAll('.isi-handle').forEach(handle => {
    handle.addEventListener('mousedown', ev => {
      ev.preventDefault();
      const idx  = parseInt(handle.dataset.memberIdx, 10);
      const edge = handle.dataset.edge;
      const m    = members[idx];
      if (!m) return;
      const r    = replacePages[m.pi].regions[m.ri];

      let pendingFrame = 0;
      function onMove(e){
        const rect = strip.getBoundingClientRect();
        let frac = (e.clientY - rect.top) / rect.height;
        frac = Math.max(0, Math.min(1, frac));
        // Top stays strictly above bot (0.001 buffer); otherwise FULLY
        // independent — slices may overlap or have gaps.
        if (edge === 'top') {
          r.isi_top = Math.min(frac, r.isi_bot - 0.001);
          handle.style.top = (r.isi_top * 100) + '%';
        } else {
          r.isi_bot = Math.max(frac, r.isi_top + 0.001);
          handle.style.top = (r.isi_bot * 100) + '%';
        }
        const band = strip.querySelector(`.isi-band[data-member-idx="${idx}"]`);
        if (band) {
          band.style.top    = (r.isi_top * 100) + '%';
          band.style.height = ((r.isi_bot - r.isi_top) * 100) + '%';
        }
        // Live update of the canvas — debounce via requestAnimationFrame so
        // we redraw at most once per frame (~16ms).
        cancelAnimationFrame(pendingFrame);
        pendingFrame = requestAnimationFrame(isiModalRenderLivePreview);
      }
      function onUp(){
        document.removeEventListener('mousemove', onMove);
        document.removeEventListener('mouseup', onUp);
      }
      document.addEventListener('mousemove', onMove);
      document.addEventListener('mouseup', onUp);
    });
  });
}

function isiModalReset(){
  if (!_currentIsiPath) return;
  storyResetIsiChain(encodeURIComponent(_currentIsiPath));
  isiModalRender();   // re-render with fresh auto values
}

// Apply only the *currently visible* page back to the page list thumbnail.
// We send a one-page subset of mappings so the server skips the rest, which
// is faster than a full preview AND avoids stomping unrelated work.
async function isiModalApply(){
  if (_currentIsiPage === null) return;
  await storyReplacePreviewForPage(_currentIsiPage);
  // Refresh the live canvas in case the server applied something the JS
  // replica wouldn't have caught (e.g., color profiles).
  isiModalRenderLivePreview();
}

async function storyReplacePreviewForPage(pageNum){
  if (!replacePages.length) { replaceSetStatus('No PDF loaded.', 'err'); return; }
  if (!storyBanners || !storyBanners.length){
    replaceSetStatus('No capture bundle loaded.', 'err'); return;
  }
  // Build a one-page filtered mapping list: only the active page is sent, so
  // the server only re-renders that one thumbnail.
  const all = storyReplaceCollectMappings();
  const mappings = all.filter(p => p.page_num === pageNum);
  if (!mappings.length || !_hasAnyMapping(mappings)){
    replaceSetStatus('Nothing to apply on this page.', 'err'); return;
  }
  let d;
  try{
    const r = await busyFetch('/storyboard/replace-preview',
      {method:'POST', headers:{'Content-Type':'application/json'},
       body: JSON.stringify({mappings})},
      `Applying page ${pageNum}…`);
    d = await r.json();
  }catch(e){ replaceSetStatus('⚠ '+e.message, 'err'); return; }
  if (!d.ok){ replaceSetStatus('⚠ '+(d.error||'Failed'), 'err'); return; }

  (d.pages || []).forEach(p => {
    const idx = replacePages.findIndex(rp => rp.page_num === p.page_num);
    if (idx < 0) return;
    replacePages[idx].thumb = p.thumb;
    if (replacePages[idx]._draw) replacePages[idx]._draw();
  });
  replaceSetStatus(`✓ Updated page ${pageNum}.`, 'ok');
}

function storyResetIsiChain(encodedPath){
  // `encodedPath` came from data-isi-path (encoded once) — decode it once to
  // get the canonical filesystem path, then decode each r.image_path before
  // comparing (those are also stored encoded).
  const decodedPath = _safeDecode(encodedPath);
  replacePages.forEach(p => {
    (p.regions || []).forEach(r => {
      if (r.image_path && _safeDecode(r.image_path) === decodedPath) {
        delete r.isi_top;
        delete r.isi_bot;
      }
    });
  });
  storyRenderIsiChains();
  replaceSetStatus('↺ ISI slices reset to auto. Click "Apply & preview" to see them.', 'ok');
}

async function storyConfirmIsiChain(){
  // The user has dragged dividers to their liking — push the new slice
  // boundaries through to the page thumbnails by re-running the live preview.
  await storyReplacePreview();
}

async function storySnapRegion(pageNum, regionIdx){
  const page = replacePages.find(p => p.page_num === pageNum);
  if(!page || !page.regions || !page.regions[regionIdx]) return;
  const r = page.regions[regionIdx];
  // Send current (possibly user-edited) bounds; server returns the bbox of
  // non-white pixels inside that area.
  let d;
  try{
    const resp = await busyFetch('/storyboard/snap-region',
      {method:'POST', headers:{'Content-Type':'application/json'},
       body: JSON.stringify({
         page_num: pageNum,
         region: {x: Math.round(r.x), y: Math.round(r.y),
                  w: Math.round(r.w), h: Math.round(r.h)},
         pad: 2,
       })},
      'Snapping to content…');
    d = await resp.json();
  }catch(e){ replaceSetStatus('⚠ '+e.message, 'err'); return; }
  if(!d.ok){ replaceSetStatus('⚠ '+(d.error||'Snap failed'), 'err'); return; }
  Object.assign(r, d.region);
  if(page._draw) page._draw();
  replaceSetStatus('✓ Region snapped to content.', 'ok');
}

function storyReplaceCollectMappings(){
  // New format: per-page, per-region. Sends current (possibly user-edited)
  // region rects + chosen image path. Server pastes each image into its
  // region (scaled to fit). For ISI chains, optional `isi_top` / `isi_bot`
  // override the auto-weighted slice boundaries.
  return replacePages.map(p => ({
    page_num: p.page_num,
    page_w:   p.page_w,
    page_h:   p.page_h,
    regions:  (p.regions || []).map(r => {
      const m = {
        x: Math.round(r.x), y: Math.round(r.y),
        w: Math.round(r.w), h: Math.round(r.h),
        kind: r.kind || 'banner',
        image_path: r.image_path
                    ? decodeURIComponent(r.image_path)
                    : null,
      };
      if (typeof r.isi_top === 'number' && typeof r.isi_bot === 'number') {
        m.isi_top = r.isi_top;
        m.isi_bot = r.isi_bot;
      }
      // Original auto-detected bounds (set on upload). Lets the server wipe
      // both the original area and the user-edited rect, so moved/resized
      // regions don't leave a ghost of the original frame behind.
      if (r._original) {
        m.original_rect = {
          x: Math.round(r._original.x),
          y: Math.round(r._original.y),
          w: Math.round(r._original.w),
          h: Math.round(r._original.h),
        };
      }
      return m;
    }),
  }));
}

function _hasAnyMapping(mappings){
  return mappings.some(p => (p.regions || []).some(r => r.image_path));
}

async function storyReplacePreview(){
  if(!replacePages.length){ replaceSetStatus('No PDF loaded.', 'err'); return; }
  if(!storyBanners || !storyBanners.length){
    replaceSetStatus('No capture bundle loaded — go to Multi-Banner Viewer first, or click "⟳ Refresh capture bundle".', 'err');
    return;
  }
  const mappings = storyReplaceCollectMappings();
  if(!_hasAnyMapping(mappings)){
    replaceSetStatus('Pick at least one region to replace — every dropdown is set to "Don\'t change".', 'err'); return;
  }
  replaceSetStatus('Compositing preview…');
  let d;
  try{
    const r = await busyFetch('/storyboard/replace-preview', {method:'POST',
              headers:{'Content-Type':'application/json'},
              body: JSON.stringify({mappings})},
              'Compositing preview…');
    d = await r.json();
  }catch(e){ replaceSetStatus('⚠ '+e.message, 'err'); return; }
  if(!d.ok){ replaceSetStatus('⚠ '+(d.error||'Preview failed'), 'err'); return; }

  // Update each thumb image in place; preserve user-edited regions on the client.
  d.pages.forEach(p => {
    const idx = replacePages.findIndex(rp => rp.page_num === p.page_num);
    if(idx < 0) return;
    replacePages[idx].thumb = p.thumb;
    const img = document.querySelector(
      `.replace-page[data-page="${p.page_num}"] img`);
    if(img) img.src = p.thumb;
    if(replacePages[idx]._draw) replacePages[idx]._draw();
  });
  replaceSetStatus('✓ Preview updated — review thumbs, then Export.', 'ok');
}

async function storyReplaceExport(){
  if(!replacePages.length){ replaceSetStatus('No PDF loaded.', 'err'); return; }
  if(!storyBanners || !storyBanners.length){
    replaceSetStatus('No capture bundle loaded — go to Multi-Banner Viewer first, or click "⟳ Refresh capture bundle" above.', 'err');
    return;
  }
  const mappings = storyReplaceCollectMappings();
  if(!_hasAnyMapping(mappings)){
    replaceSetStatus('Pick at least one region to replace — every dropdown is set to "Don\'t change".', 'err'); return;
  }
  replaceSetStatus('Generating updated PDF…');
  let d;
  try{
    const r = await busyFetch('/storyboard/replace-export', {method:'POST',
              headers:{'Content-Type':'application/json'},
              body: JSON.stringify({mappings})},
              'Generating updated PDF…');
    d = await r.json();
  }catch(e){ replaceSetStatus('⚠ '+e.message, 'err'); return; }
  if(!d.ok){ replaceSetStatus('⚠ '+(d.error||'Export failed'), 'err'); return; }
  replaceSetStatus('✓ Updated PDF saved.', 'ok');
  document.getElementById('story-replace-download-link').style.display = '';
  // Suggest a sensible default name based on the uploaded source PDF.
  window._lastReplacedName = (window._replaceSourceName || 'storyboard') + '_updated.pdf';
  // Auto-trigger the Save-As dialog now that export succeeded.
  storyDownloadReplacedPdf();
}

async function storyDownloadReplacedPdf(){
  const url = '/storyboard/replace-download?t=' + Date.now();
  const name = window._lastReplacedName || 'storyboard_updated.pdf';
  await downloadFile(url, name, 'application/pdf', '.pdf');
}

// ── CPFP (tab 6): change-proof full-proof ─────────────────────────────────
let _cpfpDiff  = null;                       // last diff result from /cpfp/diff
let _cpfpFiles = {old:null, new:null};       // filenames per slot
let _cpfpMode  = 'visual';                   // 'visual' | 'text' | 'both'

function cpfpStatus(msg, tone){
  const el = document.getElementById('cpfp-status');
  el.textContent = msg || '';
  el.className   = 'bc-status' + (tone ? ' ' + tone : '');
}

function cpfpSetMode(m){
  if (m !== 'visual' && m !== 'text' && m !== 'both') m = 'visual';
  _cpfpMode = m;
  document.querySelectorAll('#cpfp-mode button[data-mode]').forEach(b => {
    b.classList.toggle('active', b.dataset.mode === m);
  });
  // If a comparison has already been computed, re-fetch with the new mode
  // so the user sees results matching the toggle right away. Otherwise
  // just remember the choice and the next Compare click will use it.
  if (_cpfpDiff) cpfpRun();
}

async function cpfpUpload(slot, input){
  if(!input.files || !input.files.length){ return; }
  const file = input.files[0];
  let d;
  try{
    // Step 1: upload the binary
    const fd = new FormData();
    fd.append('slot', slot);
    fd.append('file', file, file.name);
    const u = await busyFetch('/upload/storyboard', {method:'POST', body: fd},
                              'Uploading ' + slot.toUpperCase() + ' PDF…');
    const ud = await u.json();
    if (!ud.ok){ cpfpStatus('⚠ ' + (ud.error || 'Upload failed'), 'err'); return; }
    // Step 2: ask the server to render pages
    const r = await busyFetch('/cpfp/upload',
      {method:'POST', headers:{'Content-Type':'application/json'},
       body: JSON.stringify({slot, filename: ud.filename})},
      slot === 'old' ? 'Reading the OLD PDF…' : 'Reading the NEW PDF…');
    d = await r.json();
  }catch(e){ cpfpStatus('⚠ ' + e.message, 'err'); return; }
  if (!d.ok){ cpfpStatus('⚠ ' + (d.error || 'Failed'), 'err'); return; }

  _cpfpFiles[slot] = d.name;
  const pill = document.getElementById('cpfp-' + slot + '-pill');
  pill.textContent = '✓ ' + d.name + ' · ' + d.page_count + ' page'
                   + (d.page_count === 1 ? '' : 's');
  pill.className   = 'compare-pill ok';
  document.getElementById('cpfp-run-btn').disabled =
    !(_cpfpFiles.old && _cpfpFiles.new);
  cpfpStatus(_cpfpFiles.old && _cpfpFiles.new
    ? 'Both PDFs loaded — click Compare.'
    : 'Loaded ' + slot + '. Now load the other side.', 'ok');
}

async function cpfpRun(){
  const msg = _cpfpMode === 'text'   ? 'Computing text diff for every page…'
            : _cpfpMode === 'both'   ? 'Computing visual + text diff…'
                                     : 'Computing pixel diff for every page…';
  let d;
  try{
    const r = await busyFetch('/cpfp/diff',
      {method:'POST', headers:{'Content-Type':'application/json'},
       body: JSON.stringify({mode: _cpfpMode})}, msg);
    d = await r.json();
  }catch(e){ cpfpStatus('⚠ ' + e.message, 'err'); return; }
  if (!d.ok){ cpfpStatus('⚠ ' + (d.error || 'Failed'), 'err'); return; }
  _cpfpDiff = d;
  // Reset stale Claude analysis whenever the diff is re-run.
  _cpfpClaudeAnalysis = null;
  document.getElementById('cpfp-claude-card').style.display = 'none';
  // Enable Analyze only if Claude is connected — otherwise the click would
  // bounce off the backend with "Click Connect Claude first".
  const claudeBtn = document.getElementById('cpfp-claude-btn');
  claudeBtn.disabled = !_cpfpConnected;
  claudeBtn.title = _cpfpConnected
    ? 'Layer Claude analysis on top of the diff'
    : 'Click 🔌 Connect Claude first';
  cpfpRender();
}

// Claude analysis: layered on top of the deterministic diff. Server-side
// calls the Anthropic Messages API directly using the user's API key (saved
// once via the 🔌 Connect Claude button).
let _cpfpClaudeAnalysis = null;
let _cpfpConnected      = false;

async function cpfpRefreshConnectionStatus(){
  try{
    const r = await fetch('/cpfp/connection-status', {
      method:'POST', headers:{'Content-Type':'application/json'}, body:'{}'});
    const d = await r.json();
    _cpfpConnected = !!(d && d.connected);
    const btn = document.getElementById('cpfp-connect-btn');
    if (btn){
      btn.textContent = _cpfpConnected
        ? `✓ Claude connected${d.key_preview ? ' · '+d.key_preview : ''}`
        : '🔌 Connect Claude';
      btn.title = _cpfpConnected
        ? 'API key on file — click to view / disconnect'
        : 'Connect your Anthropic API key so Claude can analyze the diff';
    }
    // Re-evaluate the Analyze-with-Claude button: it should be enabled only
    // when both (a) a comparison has run AND (b) we're connected to Claude.
    const claudeBtn = document.getElementById('cpfp-claude-btn');
    if (claudeBtn){
      claudeBtn.disabled = !(_cpfpConnected && _cpfpDiff);
      claudeBtn.title = !_cpfpConnected
        ? 'Click 🔌 Connect Claude first'
        : (!_cpfpDiff ? 'Run Compare first' : 'Layer Claude analysis on top of the diff');
    }
  }catch(e){ /* status check is best-effort */ }
}

function cpfpOpenConnectModal(){
  document.getElementById('cpfp-api-key-input').value = '';
  document.getElementById('cpfp-connect-status').textContent = '';
  document.getElementById('cpfp-connect-status').className   = 'bc-status';
  // Show the disconnect button only when we already have a key on file.
  document.getElementById('cpfp-disconnect-btn').style.display =
    _cpfpConnected ? '' : 'none';
  document.getElementById('cpfp-connect-modal').hidden = false;
  setTimeout(() => document.getElementById('cpfp-api-key-input').focus(), 50);
}

function cpfpCloseConnectModal(){
  document.getElementById('cpfp-connect-modal').hidden = true;
}

async function cpfpDoConnect(){
  const key = document.getElementById('cpfp-api-key-input').value.trim();
  const stat = document.getElementById('cpfp-connect-status');
  if (!key){ stat.textContent = 'Paste a key first.'; stat.className='bc-status err'; return; }
  stat.textContent = 'Validating…'; stat.className = 'bc-status';
  let d;
  try{
    const r = await busyFetch('/cpfp/connect',
      {method:'POST', headers:{'Content-Type':'application/json'},
       body: JSON.stringify({api_key: key})},
      'Validating with Anthropic…');
    d = await r.json();
  }catch(e){ stat.textContent = '⚠ '+e.message; stat.className='bc-status err'; return; }
  if (!d.ok){
    stat.textContent = '⚠ ' + (d.error || 'Failed');
    stat.className   = 'bc-status err';
    return;
  }
  stat.textContent = '✓ Connected. Using ' + (d.model || 'default model') + '.';
  stat.className   = 'bc-status ok';
  await cpfpRefreshConnectionStatus();
  setTimeout(cpfpCloseConnectModal, 800);
}

async function cpfpDoDisconnect(){
  await busyFetch('/cpfp/disconnect',
    {method:'POST', headers:{'Content-Type':'application/json'}, body:'{}'},
    'Disconnecting…');
  await cpfpRefreshConnectionStatus();
  cpfpCloseConnectModal();
  cpfpStatus('Disconnected from Claude.', 'ok');
}

// Run a status check whenever the CPFP tab becomes visible (cheap, ~1 line).
document.addEventListener('DOMContentLoaded', cpfpRefreshConnectionStatus);

async function cpfpAnalyzeWithClaude(){
  if (!_cpfpDiff) {
    cpfpStatus('Run Compare first, then Analyze with Claude.', 'err'); return;
  }
  let d;
  try{
    const r = await busyFetch('/cpfp/claude',
      {method:'POST', headers:{'Content-Type':'application/json'}, body:'{}'},
      'Asking Claude to analyze the changes (this can take a minute)…');
    d = await r.json();
  }catch(e){ cpfpStatus('⚠ ' + e.message, 'err'); return; }
  if (!d.ok){
    cpfpStatus('⚠ ' + (d.error || 'Claude analysis failed'), 'err');
    return;
  }
  _cpfpClaudeAnalysis = d;
  cpfpRenderClaude();
  cpfpStatus('✓ Claude analysis complete.', 'ok');
}

function cpfpRenderClaude(){
  const a = _cpfpClaudeAnalysis;
  if (!a) return;
  const card = document.getElementById('cpfp-claude-card');
  const summaryEl = document.getElementById('cpfp-claude-summary');
  const statsEl   = document.getElementById('cpfp-claude-stats');
  card.style.display = '';

  // Whole-document summary — either Claude's structured overall_summary, or
  // (if JSON parsing failed) the raw output dumped into a code-style block.
  if (a.raw) {
    summaryEl.innerHTML = `<div class="cpfp-claude-raw">${escapeHtml(a.overall_summary || '')}</div>`;
    statsEl.innerHTML = '';
    return;
  }
  summaryEl.innerHTML = `<div class="cpfp-claude-overall">${escapeHtml(a.overall_summary || '(no summary)')}</div>`;

  // Severity tally.
  let counts = {critical: 0, content: 0, cosmetic: 0, unknown: 0};
  (a.pages || []).forEach(p => (p.changes || []).forEach(c => {
    const sev = (c.severity || '').toLowerCase();
    counts[sev in counts ? sev : 'unknown']++;
  }));
  statsEl.innerHTML = `
    ${counts.critical ? `<span><span class="cpfp-claude-sev critical">Critical</span> <strong>${counts.critical}</strong></span>` : ''}
    ${counts.content  ? `<span><span class="cpfp-claude-sev content">Content</span> <strong>${counts.content}</strong></span>`   : ''}
    ${counts.cosmetic ? `<span><span class="cpfp-claude-sev cosmetic">Cosmetic</span> <strong>${counts.cosmetic}</strong></span>` : ''}
    ${counts.unknown  ? `<span><span class="cpfp-claude-sev unknown">Unclassified</span> <strong>${counts.unknown}</strong></span>` : ''}
  `;

  // Per-page inline annotations: append to each .cpfp-row[data-page="N"].
  // Wipe any previous Claude blocks first so re-running doesn't stack.
  document.querySelectorAll('.cpfp-row .cpfp-claude-list').forEach(el => el.remove());
  (a.pages || []).forEach(p => {
    const row = document.querySelector(`.cpfp-row[data-page="${p.page_num}"]`);
    if (!row || !p.changes || !p.changes.length) return;
    const list = document.createElement('div');
    list.className = 'cpfp-claude-list';
    list.innerHTML = p.changes.map(c => {
      const sev = (c.severity || '').toLowerCase();
      const sevClass = ['critical','content','cosmetic'].includes(sev) ? sev : 'unknown';
      return `
        <div class="cpfp-claude-row">
          <span class="cpfp-claude-sev ${sevClass}">${sevClass}</span>
          <span class="cpfp-claude-kind">${escapeHtml(c.kind || '')}</span>
          <span style="flex:1;">${escapeHtml(c.summary || '')}</span>
          <span class="cpfp-claude-conf">${escapeHtml(c.confidence || '')}</span>
        </div>`;
    }).join('');
    row.appendChild(list);
  });
}

async function cpfpClear(){
  await busyFetch('/cpfp/clear',
    {method:'POST', headers:{'Content-Type':'application/json'}, body:'{}'},
    'Resetting…');
  _cpfpDiff = null; _cpfpFiles = {old:null, new:null};
  _cpfpClaudeAnalysis = null;
  document.getElementById('cpfp-old-pill').textContent = 'No file loaded';
  document.getElementById('cpfp-new-pill').textContent = 'No file loaded';
  document.getElementById('cpfp-old-pill').className = 'compare-pill';
  document.getElementById('cpfp-new-pill').className = 'compare-pill';
  document.getElementById('cpfp-run-btn').disabled = true;
  document.getElementById('cpfp-claude-btn').disabled = true;
  document.getElementById('cpfp-summary').style.display = 'none';
  document.getElementById('cpfp-claude-card').style.display = 'none';
  document.getElementById('cpfp-pages').innerHTML = '';
  cpfpStatus('Reset.', 'ok');
}

function cpfpRender(){
  if (!_cpfpDiff) return;
  const d = _cpfpDiff;
  const showVisual = (d.mode === 'visual' || d.mode === 'both');
  const showText   = (d.mode === 'text'   || d.mode === 'both');

  // Summary card.
  const sum = document.getElementById('cpfp-summary');
  sum.style.display = '';
  document.getElementById('cpfp-summary-title').textContent =
    d.old_name + '  ⇄  ' + d.new_name;
  const total    = d.pages.length;
  const matches  = d.pages.filter(p => p.status === 'match').length;
  const diffs    = d.pages.filter(p => p.status === 'diff').length;
  const added    = d.pages.filter(p => p.status === 'added').length;
  const removed  = d.pages.filter(p => p.status === 'removed').length;
  const totalRegions     = d.pages.reduce((s, p) => s + (p.change_count   || 0), 0);
  const totalTextChanges = d.pages.reduce((s, p) => s + (p.text_diff_count || 0), 0);
  const modeLabel = d.mode === 'visual' ? 'Visual diff'
                  : d.mode === 'text'   ? 'Text diff'
                                        : 'Visual + text diff';
  document.getElementById('cpfp-summary-stats').innerHTML = `
    <div class="cpfp-stats-line">
      <span><strong>${modeLabel}</strong></span>
      <span><strong>${total}</strong> page${total===1?'':'s'} compared</span>
      <span><strong>${matches}</strong> identical</span>
      <span><strong>${diffs}</strong> with changes</span>
      ${added   ? `<span><strong>${added}</strong> added in new</span>`     : ''}
      ${removed ? `<span><strong>${removed}</strong> removed from new</span>` : ''}
      ${ showVisual ? `<span><strong>${totalRegions}</strong> change region${totalRegions===1?'':'s'}</span>` : '' }
      ${ showText   ? `<span><strong>${totalTextChanges}</strong> word change${totalTextChanges===1?'':'s'}</span>` : '' }
    </div>`;

  // Build per-page rows with the right combination of visual + text panes.
  const host = document.getElementById('cpfp-pages');
  host.innerHTML = d.pages.map(p => {
    const status = p.status;

    // ── Header line: change-count summary tailored to active mode(s) ────
    let summary = '';
    if (status === 'match') {
      summary = (showText && (p.text_diff_count || 0) === 0 && !showVisual)
        ? 'Text identical' : 'No changes';
    } else if (status === 'diff') {
      const parts = [];
      if (showVisual && (p.change_count   || 0) > 0)
        parts.push(p.change_count + ' visual change region'
                   + (p.change_count === 1 ? '' : 's'));
      if (showText && (p.text_diff_count || 0) > 0)
        parts.push(p.text_diff_count + ' word change'
                   + (p.text_diff_count === 1 ? '' : 's'));
      summary = parts.join(' · ');
    } else if (status === 'added') {
      summary = 'Added in new (no counterpart in old)';
    } else {
      summary = 'Removed from new (was in old)';
    }
    const head = `
      <div class="cpfp-row-head">
        <span class="pill">Page ${p.page_num}</span>
        <span class="cpfp-row-meta">${summary}</span>
      </div>`;

    // ── Visual panes (thumbnails + box overlay) ─────────────────────────
    let visualBlock = '';
    if (showVisual || status === 'added' || status === 'removed') {
      const oldSide = (status === 'added')
        ? `<div class="cpfp-thumb-wrap" style="opacity:.4;align-items:center;justify-content:center;padding:40px 12px;color:var(--text-muted);">
             <div class="cpfp-thumb-label">Old · (none)</div>
             <em>Page didn't exist in old version.</em>
           </div>`
        : `<div class="cpfp-thumb-wrap">
             <div class="cpfp-thumb-label">Old · ${escapeHtml(d.old_name || '')}</div>
             <img src="${p.old_thumb}" alt="old page ${p.page_num}" draggable="false">
           </div>`;
      const newSide = (status === 'removed')
        ? `<div class="cpfp-thumb-wrap" style="opacity:.4;align-items:center;justify-content:center;padding:40px 12px;color:var(--text-muted);">
             <div class="cpfp-thumb-label">New · (none)</div>
             <em>Page was removed from the new version.</em>
           </div>`
        : `<div class="cpfp-thumb-wrap">
             <div class="cpfp-thumb-label">New · ${escapeHtml(d.new_name || '')}</div>
             <div style="position:relative;">
               <img src="${p.new_thumb}" alt="new page ${p.page_num}" draggable="false">
               ${ showVisual ? `
                 <div class="cpfp-box-layer" style="top:0;">
                   ${ (p.boxes || []).map((b, i) => `
                     <div class="cpfp-box"
                          style="left:${(b.rx*100).toFixed(2)}%;
                                 top:${(b.ry*100).toFixed(2)}%;
                                 width:${(b.rw*100).toFixed(2)}%;
                                 height:${(b.rh*100).toFixed(2)}%;"
                          title="Change region ${i+1}"></div>
                   `).join('') }
                 </div>` : '' }
             </div>
           </div>`;
      visualBlock = `<div class="cpfp-pair">${oldSide}${newSide}</div>`;
    }

    // ── Text panes (word-diff with <ins>/<del>) ─────────────────────────
    let textBlock = '';
    if (showText && status !== 'added' && status !== 'removed') {
      const leftBody  = p.text_diff_left  || '<em>(no text extracted)</em>';
      const rightBody = p.text_diff_right || '<em>(no text extracted)</em>';
      textBlock = `
        <div class="cpfp-text-pair">
          <div class="cpfp-text-pane">
            <div class="cpfp-thumb-label">Old · text</div>
            <div class="cpfp-text-body">${leftBody}</div>
          </div>
          <div class="cpfp-text-pane">
            <div class="cpfp-thumb-label">New · text</div>
            <div class="cpfp-text-body">${rightBody}</div>
          </div>
        </div>`;
    }

    return `<div class="cpfp-row ${status}" data-page="${p.page_num}">
              ${head}
              ${visualBlock}
              ${textBlock}
            </div>`;
  }).join('');

  cpfpStatus('✓ Compared ' + total + ' page' + (total===1?'':'s') + '.', 'ok');
}

// HTML escape (small util so file names render safely).
function escapeHtml(s){
  return String(s).replace(/[&<>"']/g, ch => ({
    '&':'&amp;','<':'&lt;','>':'&gt;','"':'&quot;',"'":'&#39;'
  }[ch]));
}
</script>
</body>
</html>"""


# ═══════════════════════════════════════════════════════════════════════════════
# HTTP server
# ═══════════════════════════════════════════════════════════════════════════════

def _broadcast(sess, msg):
    """Push a progress message to a single session's SSE queue.

    Note: this no longer fans out to all connected clients — that caused
    cross-user log bleed. Each session has its own queue (sess.sse_queue)
    and only sees its own progress.
    """
    sess.sse_queue.put(msg)


class Handler(BaseHTTPRequestHandler):
    def log_message(self, *a): pass

    # ── Session resolution (runs at the top of every request) ─────────────────
    def _get_cookie(self, name):
        raw = self.headers.get("Cookie") or ""
        for chunk in raw.split(";"):
            chunk = chunk.strip()
            if not chunk or "=" not in chunk:
                continue
            k, v = chunk.split("=", 1)
            if k.strip() == name:
                return v.strip()
        return None

    # New session cookies that should be set on the response. Stored as a list
    # so _send() can pick them up. We only set the cookie once per visitor.
    def _set_cookie(self, name, value):
        # Cookie attributes: HttpOnly so JS can't read it; SameSite=Lax so
        # cross-site nav still sends it; Path=/ so all routes share the same
        # session; no Secure flag because HF Spaces serves HTTPS at the edge
        # but the container itself sees plain HTTP.
        if not hasattr(self, "_pending_cookies"):
            self._pending_cookies = []
        self._pending_cookies.append(
            f"{name}={value}; HttpOnly; SameSite=Lax; Path=/; Max-Age=86400"
        )

    def _resolve_session(self):
        """Find or create the Session for this request, keyed by `bfc_sid`."""
        sid = self._get_cookie("bfc_sid")
        sess = None
        if sid:
            with SESSIONS_LOCK:
                sess = SESSIONS.get(sid)
        if sess is None:
            sess = _new_session()
            self._set_cookie("bfc_sid", sess.id)
        sess.last_seen = time.monotonic()
        return sess

    def do_GET(self):
        self.sess = self._resolve_session()
        # Match the path component only — cache-bust suffixes like `?t=12345`
        # would otherwise make exact-match routes (e.g. `/storyboard/download`)
        # fall through to the 404 branch.
        bare_path = self.path.split("?", 1)[0]
        if bare_path in ("/", ""):
            self._html(HTML)
        elif bare_path == "/progress":
            self._sse()
        elif bare_path.startswith("/banner/") or bare_path == "/banner":
            self._serve_banner(self.sess.banner_dir, "/banner/")
        elif bare_path.startswith("/bannerA/"):
            self._serve_banner(self.sess.p3a.dir, "/bannerA/")
        elif bare_path.startswith("/bannerB/"):
            self._serve_banner(self.sess.p3b.dir, "/bannerB/")
        elif bare_path.startswith("/banner_multi/"):
            self._serve_banner_multi()
        elif bare_path.startswith("/storyboard/frame/"):
            self._serve_storyboard_frame()
        elif bare_path == "/storyboard/download":
            self._serve_storyboard_pdf()
        elif bare_path == "/storyboard/replace-download":
            self._serve_replaced_storyboard_pdf()
        elif bare_path == "/storyboard/isi-image":
            self._serve_storyboard_isi_image()
        elif bare_path.startswith("/cpfp/page/"):
            self._serve_cpfp_page()
        elif bare_path.startswith("/download/"):
            self._serve_session_download()
        else:
            self._send(404, "text/plain", b"Not found")

    def do_POST(self):
        self.sess = self._resolve_session()
        # Match the path component only — query strings (e.g. cache-bust ?t=…)
        # would otherwise make exact-match routes fall through to 404.
        bare_path = self.path.split("?", 1)[0]
        routes = {
            "/validate":       self._validate,
            "/run":             self._run,
            # ── Browser-native uploads (replaces 4 osascript pickers) ──
            "/upload/banner":              self._upload_banner_folder,
            "/upload/isi-doc":             self._upload_isi_doc,
            "/upload/storyboard":          self._upload_storyboard_pdf,
            "/upload/storyboard-existing": self._upload_existing_storyboard_pdf,
            "/upload/storyboard-bundle":   self._upload_storyboard_bundle,
            # ── ISI Compare (page 2) ──
            "/isi/banner":     self._isi_banner_text,
            "/isi/upload":     self._isi_upload_doc,
            "/isi/diff":       self._diff_text,
            # ── Banner Compare (page 3) ──
            "/p3/validate":    self._p3_validate,
            "/p3/info":        self._p3_info,
            "/p3/visdiff":     self._p3_visdiff,
            # ── Multi-Banner Viewer (page 4) ──
            "/multi/load":     self._multi_load,
            "/multi/capture":  self._multi_capture,
            # ── Storyboard (page 5) ──
            "/storyboard/state":           self._storyboard_state,
            "/storyboard/load":            self._storyboard_load,
            "/storyboard/build":           self._storyboard_build,
            "/storyboard/upload":          self._storyboard_upload_existing,
            "/storyboard/replace-preview": self._storyboard_replace_preview,
            "/storyboard/replace-export":  self._storyboard_replace_export,
            "/storyboard/snap-region":     self._storyboard_snap_region,
            # ── CPFP — change-proof full-proof (page 6) ──
            "/cpfp/upload":                self._cpfp_upload,
            "/cpfp/diff":                  self._cpfp_diff,
            "/cpfp/clear":                 self._cpfp_clear,
            "/cpfp/claude":                self._cpfp_claude,
            "/cpfp/connect":               self._cpfp_connect,
            "/cpfp/disconnect":            self._cpfp_disconnect,
            "/cpfp/connection-status":     self._cpfp_connection_status,
        }
        fn = routes.get(bare_path)
        if fn: fn()
        else:  self._send(404, "text/plain", b"Not found")

    # ── Banner file server ─────────────────────────────────────────────────────
    def _serve_banner(self, banner_dir, prefix="/banner/"):
        if not banner_dir:
            self._send(503, "text/plain", b"No banner loaded"); return

        rel = self.path[len(prefix):] or "index.html"
        # strip query strings
        rel = rel.split("?")[0]
        fp  = banner_dir / rel

        if not fp.exists() or not fp.is_file():
            self._send(404, "text/plain", b"Not found"); return

        mime, _ = mimetypes.guess_type(str(fp))
        body = fp.read_bytes()

        if fp.name == "index.html":
            # Inject preview-control script before </body>
            text = body.decode("utf-8", errors="replace")
            text = text.replace("</body>", BANNER_INJECT + "\n</body>")
            body = text.encode("utf-8")
            mime = "text/html; charset=utf-8"

        self._send(200, mime or "application/octet-stream", body)

    # ── Browser-native uploads (replaces 4 osascript pickers) ─────────────────
    def _read_multipart(self):
        """Read the request body and parse it as multipart/form-data.
        Returns list of part dicts (see _parse_multipart) or []."""
        n = int(self.headers.get("Content-Length", 0) or 0)
        if not n:
            return []
        # Cap at 200 MB so a runaway upload can't blow out memory; that's
        # well above any expected banner zip / storyboard PDF.
        if n > 200 * 1024 * 1024:
            return []
        body = self.rfile.read(n)
        return _parse_multipart(body, self.headers.get("Content-Type", "") or "")

    def _save_uploads_to(self, subdir: str, parts: list, single_file: bool = False) -> list:
        """Write each file part into sess.work_dir/<subdir>/, preserving any
        relative folder structure encoded in the filename (for webkitdirectory).
        Returns a list of saved relative paths.
        """
        saved = []
        target = self.sess.work_dir / subdir
        target.mkdir(parents=True, exist_ok=True)
        for p in parts:
            if not p.get("filename"):
                continue
            rel = _safe_subpath(p["filename"])
            if single_file:
                # Strip any folder structure for single-file uploads (we just
                # want the basename — the storyboard PDF, the ISI doc, etc.)
                rel = Path(rel.name)
            dest = target / rel
            dest.parent.mkdir(parents=True, exist_ok=True)
            dest.write_bytes(p["data"])
            saved.append(str(rel))
        return saved

    def _upload_banner_folder(self):
        """Receive a folder of banner files (webkitdirectory) and stash them
        under sess.work_dir/banner[-<slot>]/. Sets the matching slot's banner_dir.

        Optional form field `slot`:
          - omitted / "main"  → Frame Capture banner (sess.banner_dir)
          - "A" / "B"          → Banner Compare slot (sess.p3a / sess.p3b)
        """
        sess = self.sess
        parts = self._read_multipart()
        if not parts:
            self._json({"ok": False, "error": "No files received."}); return
        # Pull slot field out of non-file parts.
        slot = "main"
        file_parts = []
        for p in parts:
            if p.get("filename"):
                file_parts.append(p)
            elif p.get("name") == "slot":
                v = (p["data"] or b"").decode("utf-8", errors="ignore").strip().upper()
                if v in ("A", "B", "MAIN"):
                    slot = v.lower()
        if not file_parts:
            self._json({"ok": False, "error": "No files received."}); return
        # Per-slot subfolder so concurrent uploads don't clobber each other.
        subdir = "banner" if slot == "main" else f"banner-{slot}"
        target_root = sess.work_dir / subdir
        if target_root.exists():
            shutil.rmtree(target_root, ignore_errors=True)
        saved = self._save_uploads_to(subdir, file_parts, single_file=False)
        if not saved:
            self._json({"ok": False, "error": "Couldn't save uploaded files."}); return
        # Find the folder containing index.html (prefer shallowest match).
        candidates = list(target_root.rglob("index.html"))
        if not candidates:
            self._json({"ok": False, "error": "No index.html in the uploaded folder."}); return
        idx = min(candidates, key=lambda p: len(p.parts))
        w, h = _get_banner_size(idx)
        if slot == "main":
            sess.banner_dir = idx.parent
            sess.banner_w, sess.banner_h = w, h
            sess.banner_is_remote  = True
            sess.banner_studio_url = None
        elif slot == "a":
            sess.p3a.dir, sess.p3a.w, sess.p3a.h, sess.p3a.remote = idx.parent, w, h, True
        elif slot == "b":
            sess.p3b.dir, sess.p3b.w, sess.p3b.h, sess.p3b.remote = idx.parent, w, h, True
        self._json({
            "ok": True,
            "slot":      slot,
            "name":      idx.parent.name,
            "width":     w,
            "height":    h,
            "is_remote": True,
            "files":     len(saved),
        })

    def _upload_isi_doc(self):
        """Receive a single PDF/DOCX for ISI Compare. Returns the saved
        filename so /isi/upload can read it back when the user hits Compare."""
        parts = self._read_multipart()
        files = [p for p in parts if p.get("filename")]
        if not files:
            self._json({"ok": False, "error": "No file received."}); return
        # Single-file: strip prior uploads in this slot to avoid clutter
        target = self.sess.work_dir / "isi-doc"
        if target.exists():
            shutil.rmtree(target, ignore_errors=True)
        saved = self._save_uploads_to("isi-doc", [files[0]], single_file=True)
        self._json({"ok": True, "filename": saved[0] if saved else None})

    def _upload_storyboard_pdf(self):
        """Receive a single PDF for CPFP. Body field 'slot' = 'old' or 'new'
        sent as a sibling form field. Stores at sess.work_dir/cpfp/<slot>/."""
        parts = self._read_multipart()
        slot = ""
        files = []
        for p in parts:
            if p.get("filename"):
                files.append(p)
            elif p.get("name") == "slot":
                slot = (p["data"] or b"").decode("utf-8", errors="ignore").strip()
        if slot not in ("old", "new"):
            self._json({"ok": False, "error": "slot must be 'old' or 'new'."}); return
        if not files:
            self._json({"ok": False, "error": "No file received."}); return
        target = self.sess.work_dir / "cpfp" / slot
        if target.exists():
            shutil.rmtree(target, ignore_errors=True)
        saved = self._save_uploads_to(f"cpfp/{slot}", [files[0]], single_file=True)
        self._json({"ok": True, "slot": slot, "filename": saved[0] if saved else None})

    def _upload_existing_storyboard_pdf(self):
        """Receive a single PDF for the storyboard "Update existing" flow."""
        parts = self._read_multipart()
        files = [p for p in parts if p.get("filename")]
        if not files:
            self._json({"ok": False, "error": "No file received."}); return
        target = self.sess.work_dir / "storyboard-existing"
        if target.exists():
            shutil.rmtree(target, ignore_errors=True)
        saved = self._save_uploads_to("storyboard-existing", [files[0]], single_file=True)
        self._json({"ok": True, "filename": saved[0] if saved else None})

    def _upload_storyboard_bundle(self):
        """Receive a previously-captured multi-banner bundle and load it into
        the Storyboard tab.

        Two upload modes are supported (auto-detected from what the browser
        sends):
          1. **Single .zip**       — extract into sess.work_dir, scan
          2. **Folder upload**     — webkitdirectory sends every file in the
                                     folder with the relative path encoded in
                                     the filename; we save them preserving
                                     that structure, then scan

        Either way: sets sess.last_bundle to the discovered bundle root and
        returns the banner list so the Storyboard UI picks it up.
        """
        sess = self.sess
        parts = self._read_multipart()
        files = [p for p in parts if p.get("filename")]
        if not files:
            self._json({"ok": False, "error": "No file(s) received."}); return

        # Single .zip → extract; anything else → treat as folder upload.
        is_zip = (len(files) == 1
                  and (files[0]["filename"] or "").lower().endswith(".zip"))

        # Per-upload sandbox so successive loads don't clobber each other.
        bundles_root = sess.work_dir / "loaded-bundles"
        bundles_root.mkdir(parents=True, exist_ok=True)
        slot = f"b{int(time.monotonic()*1000) & 0xFFFFFF:x}"
        target = bundles_root / slot
        target.mkdir(parents=True, exist_ok=True)

        if is_zip:
            zip_path = target / "_upload.zip"
            zip_path.write_bytes(files[0]["data"])
            try:
                with zipfile.ZipFile(zip_path) as zf:
                    # Defend against zip-slip (paths escaping target).
                    base = target.resolve()
                    for n in zf.namelist():
                        out = (target / n).resolve()
                        try:
                            out.relative_to(base)
                        except ValueError:
                            self._json({"ok": False,
                                        "error": "Bundle contained a path outside the upload folder."})
                            return
                    zf.extractall(target)
            except zipfile.BadZipFile:
                self._json({"ok": False, "error": "Not a valid zip file."}); return
            except Exception as e:
                self._json({"ok": False, "error": f"Couldn't extract zip: {e}"}); return
            finally:
                zip_path.unlink(missing_ok=True)
        else:
            # Folder upload: webkitdirectory sends each file with the relative
            # path encoded in `filename` (e.g. "myCapture/160x600/01_logo.png").
            for p in files:
                rel = _safe_subpath(p["filename"])
                dest = target / rel
                dest.parent.mkdir(parents=True, exist_ok=True)
                dest.write_bytes(p["data"])

        # The bundle layout is:  <bundle_name>/<size>/...   (PNG frames + isi_full.png)
        # First scan the upload root; if that yields nothing, look one level
        # deeper (a single subfolder = the bundle root inside the upload).
        candidate = target
        banners = _scan_capture_bundle(candidate)
        if not banners:
            children = [p for p in target.iterdir() if p.is_dir()]
            if len(children) == 1:
                candidate = children[0]
                banners = _scan_capture_bundle(candidate)
        if not banners:
            self._json({"ok": False,
                        "error": "No banner sub-folders with PNG frames found in the upload."})
            return
        sess.last_bundle = candidate
        self._json({"ok": True, "bundle": str(candidate), "banners": banners,
                    "name": candidate.name})

    # ── Validate ───────────────────────────────────────────────────────────────
    def _validate(self):
        sess = self.sess
        body = self._body()
        raw  = body.get("path", "").strip()
        # Optional: when re-validating to pick a different size from the
        # dropdown, the frontend sends the same input plus `creative_index`.
        # For DoubleClick URLs we use the cached creative list; for Ziflow
        # multi-asset paste we just pick the Nth URL from the same list.
        creative_index = int(body.get("creative_index") or 0)

        dc_info = None  # extras to surface in the response if it's a DoubleClick URL
        studio_url = None  # Studio preview URL for the picked creative (DoubleClick only)
        multi_banners = None  # list of {name, url, creative_id} for the size dropdown

        # ── Multi-asset paste (newline-separated proof-assets.ziflow.io URLs) ──
        # The bookmarklet drops these into the field for private Klick proofs.
        # Same shape as DoubleClick's `creatives` so the frontend can populate
        # the size dropdown identically.
        manual = _parse_ziflow_asset_urls(raw)
        if manual:
            try:
                if not (0 <= creative_index < len(manual)):
                    creative_index = 0
                picked = manual[creative_index]
                folder = _download_banner_authed(picked["url"], "", "")
            except Exception as e:
                self._json({"ok": False, "error": f"Couldn't download asset URL: {e}"})
                return
            html = folder / "index.html"
            is_remote = True
            # Resolve a friendly name (size) per banner so the dropdown shows
            # "160x600" rather than the raw "banner_2cca91a5". Sized banners
            # also need this to render correctly via the existing UI logic.
            sized = []
            for i, b in enumerate(manual):
                # We only know the size for the picked one (we just downloaded
                # it); for the others we'll show the truncated guid.
                if i == creative_index:
                    try:
                        w, h = _get_banner_size(folder / "index.html")
                        sized.append({**b, "name": f"{w}x{h}", "size": f"{w}x{h}"})
                    except Exception:
                        sized.append({**b, "size": None})
                else:
                    sized.append({**b, "size": None})
            multi_banners = {
                "picked_index": creative_index,
                "total_count":  len(manual),
                "creatives":    sized,
                "all_sizes":    [b.get("size") or b["name"] for b in sized],
                "picked_name":  sized[creative_index].get("size")
                                or sized[creative_index]["name"],
            }
        # Ziflow proof URL? Download the assets to a temp folder first.
        elif ZIFLOW_PROOF_RE.match(raw):
            try:
                folder = _ziflow_url_to_local(raw)
            except Exception as e:
                self._json({"ok": False, "error": f"Couldn't load Ziflow banner: {e}"})
                return
            html = folder / "index.html"
            is_remote = True
        # DoubleClick Studio external-preview URL?
        elif DOUBLECLICK_PREVIEW_RE.match(raw):
            try:
                # Reuse the cached creative list when the user is just
                # switching size on the same URL — much faster than another
                # Playwright cycle through the size dropdown.
                cached = None
                cache = getattr(sess, "_dc_cache", None) or {}
                if cache.get("url") == raw:
                    cached = cache.get("banners")
                if cached:
                    folder, dc_info = _doubleclick_url_to_local(
                        raw, picked_index=creative_index, cached_banners=cached)
                else:
                    with PLAYWRIGHT_SEM:
                        folder, dc_info = _doubleclick_url_to_local(
                            raw, picked_index=creative_index)
                    # Cache discovery for subsequent dropdown picks.
                    sess._dc_cache = {"url": raw,
                                      "banners": dc_info.get("creatives") or []}
            except Exception as e:
                self._json({"ok": False, "error": f"Couldn't load DoubleClick preview: {e}"})
                return
            html = folder / "index.html"
            is_remote = True
            studio_url = dc_info.get("studio_url") if dc_info else None
        else:
            folder = Path(raw)
            html   = (folder/"index.html") if folder.is_dir() else (
                folder if folder.suffix == ".html" else folder/"index.html"
            )
            is_remote = False

        if html.exists():
            sess.banner_dir = html.parent
            sess.banner_w, sess.banner_h = _get_banner_size(html)
            sess.banner_is_remote = is_remote
            sess.banner_studio_url = studio_url
            payload = {"ok": True, "name": html.parent.name,
                       "width": sess.banner_w, "height": sess.banner_h,
                       "is_remote": is_remote}
            if dc_info:
                payload["doubleclick"] = dc_info
            if multi_banners:
                # Same shape as `doubleclick` — frontend's _renderDcPicker
                # treats both keys identically so the size dropdown shows.
                payload["multi_banners"] = multi_banners
            self._json(payload)
        else:
            self._json({"ok": False, "error": f"No index.html in: {raw}"})

    # ── Run ────────────────────────────────────────────────────────────────────
    def _run(self):
        sess = self.sess
        body = self._body()

        if not sess.banner_dir:
            self._json({"status": "error", "error": "No banner loaded"})
            return

        # Banner files always live in sess.banner_dir (local folders, Ziflow
        # temp folders, or DoubleClick temp folders, or session uploads).
        banner = sess.banner_dir / "index.html"

        # Hosted: output always goes to a per-session temp folder; the user
        # downloads the resulting zip via the URL we send back in the
        # `done` SSE event.
        out = sess.work_dir / "banner_frames"
        if out.exists():
            shutil.rmtree(out, ignore_errors=True)
        out.mkdir(parents=True, exist_ok=True)
        sess.output_dir = out

        raw_frames  = body.get("frames", [])
        frames      = [(f["name"], float(f["t"])) for f in raw_frames] if raw_frames else DEFAULT_FRAMES
        do_frames   = body.get("do_frames",   True)
        do_video    = body.get("do_video",    True)
        # Independent ISI controls — match the Multi-Banner Viewer's 4-checkbox
        # layout. Default each to do_frames if not provided so older clients
        # that only send {do_frames, do_video} still get the original behaviour.
        do_isi_png  = body.get("do_isi_png",  do_frames)
        do_isi_html = body.get("do_isi_html", do_frames)

        # Capture session reference for the worker thread (self goes out of
        # scope when the request returns).
        captured_sess = sess
        studio_url    = sess.banner_studio_url

        def _work():
            try:
                with PLAYWRIGHT_SEM:
                    asyncio.run(run_all(
                        banner, out, frames, do_frames, do_video,
                        lambda m: _broadcast(captured_sess, {"type":"log","text":m}),
                        do_isi_png=do_isi_png,
                        do_isi_html=do_isi_html,
                        studio_url=studio_url,
                    ))
                # Zip the output for download.
                zip_path = captured_sess.work_dir / "banner_frames.zip"
                _broadcast(captured_sess, {"type":"log","text":f"Zipping → {zip_path.name}"})
                with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
                    for f in out.rglob("*"):
                        if f.is_file():
                            zf.write(f, f.relative_to(out.parent))
                captured_sess.last_bundle = out
                _broadcast(captured_sess, {
                    "type":"done", "path":str(out),
                    "zip": str(zip_path),
                    "zip_url": f"/download/{zip_path.name}",
                })
            except Exception as exc:
                _broadcast(captured_sess, {"type":"error","text":str(exc)})

        threading.Thread(target=_work, daemon=True).start()
        self._json({"status":"started"})

    # ═════════════════════════════════════════════════════════════════════════
    # ISI Compare (page 2)
    # ═════════════════════════════════════════════════════════════════════════

    def _isi_banner_text(self):
        """Return the ISI text from the banner currently loaded on Page 1."""
        sess = self.sess
        if not sess.banner_dir:
            self._json({"ok": False, "error": "No banner loaded on the Frame Capture page yet."})
            return
        try:
            with PLAYWRIGHT_SEM:
                info = asyncio.run(extract_banner_info(sess.banner_dir))
        except Exception as e:
            self._json({"ok": False, "error": f"Couldn't read banner ISI: {e}"})
            return
        self._json({"ok": True, "name": sess.banner_dir.name,
                    "header": info["header_text"], "isi": info["isi_text"]})

    def _isi_upload_doc(self):
        """Read a previously-uploaded ISI doc by name and return its text.

        The browser uploads the doc via /upload/isi-doc; that endpoint stores
        the file in sess.work_dir and returns the filename. This route reads
        it back when the user clicks "Compare".
        """
        body = self._body()
        fname = body.get("filename", "").strip()
        if not fname:
            self._json({"ok": False, "error": "No file uploaded."})
            return
        fpath = self.sess.work_dir / "isi-doc" / fname
        if not fpath.exists():
            self._json({"ok": False, "error": f"Upload not found: {fname}"})
            return
        try:
            text = _extract_text_from_doc(fpath)
        except Exception as e:
            self._json({"ok": False, "error": f"Couldn't read document: {e}"})
            return
        self._json({"ok": True, "name": fpath.name, "text": _normalize_text(text)})

    def _diff_text(self):
        """Generic word-level text diff endpoint."""
        body = self._body()
        a = body.get("a", "")
        b = body.get("b", "")
        result = _word_diff_html(a, b)
        self._json({"ok": True, **result})

    # ═════════════════════════════════════════════════════════════════════════
    # Banner Compare (page 3)
    # ═════════════════════════════════════════════════════════════════════════

    def _p3_validate(self):
        """Validate a banner for slot 'A' or 'B' on the compare page."""
        sess = self.sess
        body = self._body()
        slot = body.get("slot", "A").upper()
        raw  = body.get("path", "").strip()

        if ZIFLOW_PROOF_RE.match(raw):
            try:
                folder = _ziflow_url_to_local(raw)
            except Exception as e:
                self._json({"ok": False, "error": f"Couldn't load Ziflow banner: {e}"})
                return
            html = folder / "index.html"
            is_remote = True
        else:
            folder = Path(raw)
            html = (folder/"index.html") if folder.is_dir() else (
                folder if folder.suffix == ".html" else folder/"index.html"
            )
            is_remote = False

        if not html.exists():
            self._json({"ok": False, "error": f"No index.html in: {raw}"})
            return

        w, h = _get_banner_size(html)
        target = sess.p3a if slot == "A" else sess.p3b
        target.dir, target.w, target.h, target.remote = html.parent, w, h, is_remote

        self._json({"ok": True, "slot": slot, "name": html.parent.name,
                    "width": w, "height": h, "is_remote": is_remote})

    def _p3_info(self):
        """Pull metadata + ISI text for both slots and compute diffs."""
        sess = self.sess
        if not sess.p3a.dir or not sess.p3b.dir:
            self._json({"ok": False, "error": "Load both banners first."})
            return
        try:
            with PLAYWRIGHT_SEM:
                info_a = asyncio.run(extract_banner_info(sess.p3a.dir))
                info_b = asyncio.run(extract_banner_info(sess.p3b.dir))
        except Exception as e:
            self._json({"ok": False, "error": f"Couldn't read banner: {e}"})
            return
        isi_diff = _word_diff_html(info_a["isi_text"], info_b["isi_text"])
        self._json({
            "ok": True,
            "a": {"name": sess.p3a.dir.name, **info_a},
            "b": {"name": sess.p3b.dir.name, **info_b},
            "isi_diff": isi_diff,
        })

    def _p3_visdiff(self):
        """Capture both banners at time `t` and return a diff overlay PNG (base64)."""
        sess = self.sess
        body = self._body()
        try:
            t = float(body.get("t", 0))
        except Exception:
            t = 0.0
        if not sess.p3a.dir or not sess.p3b.dir:
            self._json({"ok": False, "error": "Load both banners first."})
            return
        try:
            with PLAYWRIGHT_SEM:
                shot_a = asyncio.run(screenshot_at(sess.p3a.dir, t))
                shot_b = asyncio.run(screenshot_at(sess.p3b.dir, t))
            diff   = _visual_diff_image(shot_a, shot_b)
        except Exception as e:
            self._json({"ok": False, "error": f"Visual diff failed: {e}"})
            return
        import base64
        self._json({
            "ok":   True,
            "t":    t,
            "a":    "data:image/png;base64," + base64.b64encode(shot_a).decode(),
            "b":    "data:image/png;base64," + base64.b64encode(shot_b).decode(),
            "diff": "data:image/png;base64," + base64.b64encode(diff).decode(),
        })

    # ═════════════════════════════════════════════════════════════════════════
    # Multi-Banner Viewer (page 4) — load N banners from a Klick staging URL
    # ═════════════════════════════════════════════════════════════════════════

    def _serve_banner_multi(self):
        """Serve files from the matching slot's downloaded banner folder.

        URL pattern:  /banner_multi/<slot>/<rel-path>
        """
        # /banner_multi/<slot>/<rest>
        path = self.path.split("?", 1)[0]
        parts = path[len("/banner_multi/"):].split("/", 1)
        if not parts or not parts[0]:
            self._send(404, "text/plain", b"slot missing"); return
        slot = parts[0]
        rel  = parts[1] if len(parts) > 1 else "index.html"
        slot_info = self.sess.multi_slots.get(slot)
        if not slot_info:
            self._send(404, "text/plain", b"unknown slot"); return
        # Reuse existing _serve_banner by faking the prefix
        self._serve_banner(slot_info["dir"], f"/banner_multi/{slot}/")

    def _multi_load(self):
        """Discover every banner on the source URL and download each.

        URL flavors supported (auto-detected):
          1. **Klick / AdPiler staging** — form-based password gate, banners
             served via iframes on the same page (Playwright drives auth).
          2. **DoubleClick Studio external-preview** — no auth, dropdown of
             creatives (Playwright cycles through sizes).
          3. **Ziflow proof URL** — public proofs work via cheap regex; for
             private (Klick) proofs the user can paste their Ziflow Cookie
             into `cookie` to authenticate.
          4. **Multi-URL paste** — for private proofs without cookies, the
             user can paste each banner's `proof-assets.ziflow.io/.../
             richMedia/<guid>/index.html` URL on separate lines (or comma-
             separated) and skip discovery entirely.
        """
        sess = self.sess
        body = self._body()
        url    = (body.get("url") or "").strip()
        user   = (body.get("user") or "").strip()
        pwd    = body.get("password") or ""
        cookie = (body.get("cookie") or "").strip()

        if not url:
            self._json({"ok": False, "error": "Staging URL is required."}); return

        # Route to the appropriate discovery handler based on URL shape.
        try:
            # ── Multi-URL paste (newline / comma-separated asset CDN URLs) ──
            manual = _parse_ziflow_asset_urls(url)
            if manual:
                banners = manual
            elif ZIFLOW_PROOF_RE.match(url):
                # Ziflow has two flavors:
                #   1. Old-style proofs that ship every banner's richMedia
                #      reference statically — the cheap regex helper finds
                #      them in <1s, no Playwright needed.
                #   2. Klick's modern proofs (klickhealth.ziflow.io) which
                #      are SPAs — only the FIRST banner's guid is in the
                #      static HTML. For these we drive the SPA with
                #      Playwright (slower, ~10-20s per proof).
                # Try the cheap path first; fall back to async discovery if
                # we found fewer than 2 banners (means it's an SPA or
                # genuinely single-banner). Pass the Ziflow cookie if the
                # user provided one — needed for private Klick proofs.
                banners = _discover_ziflow_banners(url)
                if len(banners) < 2:
                    with PLAYWRIGHT_SEM:
                        banners = asyncio.run(
                            _discover_ziflow_banners_async(url, cookie_header=cookie))
            elif DOUBLECLICK_PREVIEW_RE.match(url):
                with PLAYWRIGHT_SEM:
                    banners = asyncio.run(_discover_doubleclick_banners(url))
            else:
                with PLAYWRIGHT_SEM:
                    banners = asyncio.run(_discover_klick_banners(url, user, pwd))
        except ZiflowAuthGateError as e:
            # The proof page redirected us to Ziflow's identity-verification
            # screen. Surface the helpful message verbatim — it tells the user
            # exactly which fallback to use.
            self._json({"ok": False, "error": str(e)}); return
        except Exception as e:
            err = str(e)
            if "401" in err or "Unauthorized" in err.lower():
                msg = "Wrong username or password (401)."
            elif "timeout" in err.lower():
                msg = "Timed out reaching the staging URL."
            else:
                msg = f"Couldn't load staging page: {e}"
            self._json({"ok": False, "error": msg}); return

        if not banners:
            self._json({"ok": False, "error": "No banners found on that page."}); return

        # Reset previous slots, download each banner, register
        sess.multi_slots = {}
        registered  = []
        errors      = []
        name_counts = {}   # tracks duplicate dimension labels for de-dup
        is_doubleclick = bool(DOUBLECLICK_PREVIEW_RE.match(url))
        for i, b in enumerate(banners):
            slot = f"s{i+1}"
            try:
                folder = _download_banner_authed(b["url"], user, pwd)
                w, h   = _get_banner_size(folder / "index.html")
                # Prefer the detected dimensions for the slot's friendly name.
                base_name = f"{w}x{h}"
                name_counts[base_name] = name_counts.get(base_name, 0) + 1
                disp = base_name if name_counts[base_name] == 1 else f"{base_name}_{name_counts[base_name]}"
                # For DoubleClick slots, build the per-creative Studio URL so ISI
                # capture can source markup from the live Studio preview.
                studio_url = (_studio_url_for_creative(url, b["creative_id"])
                              if is_doubleclick and b.get("creative_id") else None)
                sess.multi_slots[slot] = {"dir": folder, "w": w, "h": h, "name": disp,
                                          "studio_url": studio_url}
                registered.append({"slot": slot, "name": disp,
                                   "width": w, "height": h, "url": b["url"]})
            except Exception as e:
                errors.append(f"{b['name']}: {e}")

        if not registered:
            self._json({"ok": False, "error": "Couldn't download any banner. " + " ; ".join(errors)})
            return

        # Use the staging page's path tail as a friendly bundle name. Strip
        # the query / fragment / trailing slashes first, then keep only safe
        # filename chars — otherwise we end up with bundle names like
        # `index.html?1778011457958` which produce a zip filename containing
        # `?`, which then breaks the /download/<file> URL on the frontend
        # (encodeURIComponent → %3F → server can't find the file).
        # For multi-line inputs (paste of asset URLs), use the LAST line.
        last_line = (url.splitlines() or [url])[-1].rstrip("/")
        raw_tail = (last_line.split("?")[0].split("#")[0].split("/")[-1]
                    or "klick_capture")
        bundle_name = re.sub(r"[^A-Za-z0-9._-]+", "_", raw_tail).strip("._-") \
                      or "klick_capture"
        sess.multi_meta = {"source_url": url, "user": user, "name": bundle_name}

        self._json({
            "ok": True, "banners": registered,
            "errors": errors,
            "bundle_name": bundle_name,
        })

    def _multi_capture(self):
        """Run the existing capture pipeline once per loaded slot, then zip the bundle."""
        sess         = self.sess
        body         = self._body()
        raw_frames   = body.get("frames", [])
        do_frames    = body.get("do_frames",   True)
        do_isi_png   = body.get("do_isi_png",  True)
        do_isi_html  = body.get("do_isi_html", True)
        do_video     = body.get("do_video",    True)
        # Optional list of slots to include — if omitted/empty, do everything.
        selected     = body.get("slots") or []
        selected_set = set(selected) if selected else None

        if not sess.multi_slots:
            self._json({"status": "error", "error": "No banners loaded."}); return
        if not (do_frames or do_isi_png or do_isi_html or do_video):
            self._json({"status": "error", "error": "Pick at least one export option."}); return

        # Hosted: write the multi-bundle into the per-session work_dir; the
        # user gets the resulting zip via the existing download mechanism.
        bundle_name = (sess.multi_meta.get("name") or "klick_capture").replace("/", "_")
        bundle_root = sess.work_dir / f"{bundle_name}_capture"
        bundle_root.mkdir(parents=True, exist_ok=True)

        frames = [(f["name"], float(f["t"])) for f in raw_frames if f.get("name")]

        slots_snapshot = [(s, i) for s, i in sess.multi_slots.items()
                          if selected_set is None or s in selected_set]
        if not slots_snapshot:
            self._json({"status": "error", "error": "No banners selected."}); return

        captured_sess = sess

        def _work():
            try:
                with PLAYWRIGHT_SEM:
                    for slot, info in slots_snapshot:
                        slot_out = bundle_root / info["name"]
                        slot_out.mkdir(parents=True, exist_ok=True)
                        # Push the per-slot dimensions onto the session so
                        # capture functions that runtime-read these (initial
                        # viewport sizing, legacy paths) see the right values.
                        captured_sess.banner_w = int(info.get("w") or 160)
                        captured_sess.banner_h = int(info.get("h") or 600)
                        _broadcast(captured_sess, {"type": "log",
                                    "text": f"── {info['name']} ({info['w']}×{info['h']}) ─────"})
                        asyncio.run(run_all(
                            info["dir"] / "index.html",
                            slot_out,
                            frames,
                            do_frames=do_frames,
                            do_video=do_video,
                            log=lambda m, sname=info["name"]: _broadcast(
                                captured_sess, {"type": "log", "text": f"[{sname}] {m}"}
                            ),
                            do_isi_png=do_isi_png,
                            do_isi_html=do_isi_html,
                            studio_url=info.get("studio_url"),
                        ))

                # Remember this bundle so the Storyboard tab can auto-pull its frames.
                captured_sess.last_bundle = bundle_root

                # Bundle as one zip alongside the output folder (in work_dir)
                zip_path = sess.work_dir / f"{bundle_name}_capture.zip"
                _broadcast(captured_sess, {"type": "log", "text": f"Zipping bundle → {zip_path}"})
                with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
                    for f in bundle_root.rglob("*"):
                        if f.is_file():
                            zf.write(f, f.relative_to(bundle_root.parent))
                _broadcast(captured_sess, {"type": "done", "path": str(bundle_root), "zip": str(zip_path),
                            "bundle": str(bundle_root)})
            except Exception as exc:
                _broadcast(captured_sess, {"type": "error", "text": str(exc)})

        sess.output_dir = bundle_root
        threading.Thread(target=_work, daemon=True).start()
        self._json({"status": "started", "bundle": str(bundle_root)})

    # ═════════════════════════════════════════════════════════════════════════
    # Storyboard (page 5)
    # ═════════════════════════════════════════════════════════════════════════

    def _serve_storyboard_frame(self):
        """Serve a captured-frame PNG by its absolute path (path-encoded)."""
        sess = self.sess
        rel = self.path[len("/storyboard/frame/"):].split("?", 1)[0]
        try:
            decoded = urllib.parse.unquote(rel)
        except Exception:
            self._send(400, "text/plain", b"bad path"); return
        # Restrict to files inside this session's last_bundle (basic safety).
        if not sess.last_bundle:
            self._send(404, "text/plain", b"no bundle loaded"); return
        try:
            p = Path(decoded).resolve()
            if not p.is_file() or sess.last_bundle.resolve() not in p.parents:
                self._send(404, "text/plain", b"not in bundle"); return
        except Exception:
            self._send(404, "text/plain", b"not found"); return
        mime, _ = mimetypes.guess_type(str(p))
        self._send(200, mime or "image/png", p.read_bytes())

    def _serve_storyboard_pdf(self):
        """Stream the most recently generated storyboard PDF."""
        sess = self.sess
        if not sess.storyboard_pdf_path or not Path(sess.storyboard_pdf_path).exists():
            self._send(404, "text/plain", b"no pdf yet"); return
        data = Path(sess.storyboard_pdf_path).read_bytes()
        self.send_response(200)
        self.send_header("Content-Type",        "application/pdf")
        self.send_header("Content-Length",      str(len(data)))
        self.send_header("Content-Disposition",
                         f'attachment; filename="{Path(sess.storyboard_pdf_path).name}"')
        self.end_headers()
        self.wfile.write(data)

    def _storyboard_state(self):
        """Return the most recent capture bundle so the UI can pre-fill itself."""
        sess = self.sess
        if not sess.last_bundle or not sess.last_bundle.exists():
            self._json({"ok": True, "bundle": None, "banners": []}); return
        banners = _scan_capture_bundle(sess.last_bundle)
        self._json({"ok": True, "bundle": str(sess.last_bundle),
                    "banners": banners, "name": sess.last_bundle.name})

    def _storyboard_load(self):
        """Point the storyboard tab at a folder of pre-captured frames
        (e.g. a bundle from a different session — only works inside this
        session's work_dir on the hosted build)."""
        sess = self.sess
        body = self._body()
        path = (body.get("path") or "").strip()
        if not path:
            self._json({"ok": False, "error": "Folder path required."}); return
        folder = Path(path)
        if not folder.is_dir():
            self._json({"ok": False, "error": f"Not a folder: {path}"}); return
        # Sandbox: only allow folders inside this session's work_dir
        try:
            if sess.work_dir.resolve() not in folder.resolve().parents \
                    and folder.resolve() != sess.work_dir.resolve():
                self._json({"ok": False, "error": "Folder is outside this session."}); return
        except Exception:
            self._json({"ok": False, "error": "Could not resolve path."}); return
        banners = _scan_capture_bundle(folder)
        if not banners:
            self._json({"ok": False, "error": "No banner sub-folders with PNG frames found in that folder."})
            return
        sess.last_bundle = folder
        self._json({"ok": True, "bundle": str(folder), "banners": banners,
                    "name": folder.name})

    def _storyboard_build(self):
        """Generate the storyboard PDF from the user-supplied notes + frames."""
        sess = self.sess
        body = self._body()
        if not sess.last_bundle:
            self._json({"ok": False, "error": "No capture bundle loaded."}); return

        # Build the spec: header + footer text + an entry per banner with its
        # frame paths and the user's per-frame animation notes.
        header = body.get("header", "")
        footer = body.get("footer", "")
        in_banners = body.get("banners", [])
        if not in_banners:
            self._json({"ok": False, "error": "No banners selected."}); return

        spec_banners = []
        for b in in_banners:
            spec_banners.append({
                "name":   b.get("name", ""),
                "frames": b.get("frames", []),
                "notes":  b.get("notes",  []),
            })

        spec = {"header": header, "footer": footer, "banners": spec_banners}
        out  = sess.last_bundle / "storyboard.pdf"
        try:
            asyncio.run(_generate_storyboard_pdf(spec, out))
        except Exception as e:
            self._json({"ok": False, "error": f"PDF generation failed: {e}"}); return
        sess.storyboard_pdf_path = str(out)
        self._json({"ok": True, "pdf": str(out), "url": "/storyboard/download"})

    # ── "Update existing PDF" sub-mode ─────────────────────────────────────────

    def _serve_replaced_storyboard_pdf(self):
        """Stream the most recently exported updated-storyboard PDF."""
        sess = self.sess
        if not sess.replaced_storyboard_path or not Path(sess.replaced_storyboard_path).exists():
            self._send(404, "text/plain", b"no updated pdf yet"); return
        data = Path(sess.replaced_storyboard_path).read_bytes()
        self.send_response(200)
        self.send_header("Content-Type",        "application/pdf")
        self.send_header("Content-Length",      str(len(data)))
        self.send_header("Content-Disposition",
                         f'attachment; filename="{Path(sess.replaced_storyboard_path).name}"')
        self.end_headers()
        self.wfile.write(data)

    def _serve_storyboard_isi_image(self):
        """Serve an ISI capture (isi_full.png) by absolute path so the JS
        chain editor can render it. Restricted to files under the session's
        last_bundle to prevent arbitrary-file disclosure.
        """
        sess = self.sess
        q = self.path.split("?", 1)[1] if "?" in self.path else ""
        params = urllib.parse.parse_qs(q)
        path = (params.get("path") or [""])[0]
        if not path:
            self._send(400, "text/plain", b"missing path"); return
        if not sess.last_bundle:
            self._send(403, "text/plain", b"no bundle loaded"); return
        try:
            p = Path(urllib.parse.unquote(path)).resolve()
            if sess.last_bundle.resolve() not in p.parents:
                self._send(403, "text/plain", b"path not in bundle"); return
        except Exception:
            self._send(404, "text/plain", b"not found"); return
        if not p.is_file():
            self._send(404, "text/plain", b"not found"); return
        mime, _ = mimetypes.guess_type(str(p))
        self._send(200, mime or "image/png", p.read_bytes())

    # ── CPFP (page 6) ─────────────────────────────────────────────────────
    def _cpfp_upload(self):
        """Read a previously-uploaded PDF (via /upload/storyboard or a similar
        multipart upload) keyed by filename, render every page to a PIL image,
        and stash it in sess.cpfp_state[slot].
        Body: {"slot":"old"|"new", "filename": "..."}
        """
        sess = self.sess
        body = self._body()
        slot = (body.get("slot") or "").strip()
        fname = (body.get("filename") or "").strip()
        if slot not in ("old", "new"):
            self._json({"ok": False, "error": "slot must be 'old' or 'new'"}); return
        if not fname:
            self._json({"ok": False, "error": "No file uploaded."}); return

        path = sess.work_dir / "cpfp" / slot / fname
        if not path.exists():
            self._json({"ok": False, "error": f"Upload not found: {fname}"}); return

        try:
            page_imgs = _extract_pdf_pages_to_pngs(path)
        except Exception as e:
            self._json({"ok": False, "error": f"Couldn't read PDF: {e}"}); return

        # Per-page text extraction for the CPFP "Text" / "Both" diff modes.
        page_texts = _extract_pdf_pages_text(path)
        while len(page_texts) < len(page_imgs):
            page_texts.append("")
        any_text = any(t.strip() for t in page_texts)

        sess.cpfp_state[slot] = {
            "name": path.name,
            "pages": [
                {"page_num": i + 1, "full_image": img,
                 "w": img.size[0], "h": img.size[1],
                 "text": page_texts[i] if i < len(page_texts) else ""}
                for i, img in enumerate(page_imgs)
            ],
        }
        self._json({
            "ok": True,
            "slot": slot,
            "name": path.name,
            "page_count": len(page_imgs),
            "has_extractable_text": any_text,
        })

    def _cpfp_diff(self):
        """Compute the strict-by-index diff for both uploaded PDFs and return
        per-page payloads. Body: {"mode": "visual" | "text" | "both"}.

        Visual mode: per-page bounding boxes around regions that differ.
        Text   mode: per-page word-level <ins>/<del>-marked HTML for each side.
        Both   mode: both of the above.

        Pages where one side is missing are tagged "added" or "removed".
        """
        body = self._body() or {}
        mode = body.get("mode", "visual")
        if mode not in ("visual", "text", "both"):
            mode = "visual"
        do_visual = mode in ("visual", "both")
        do_text   = mode in ("text", "both")

        sess = self.sess
        old_pages = sess.cpfp_state["old"]["pages"]
        new_pages = sess.cpfp_state["new"]["pages"]
        if not old_pages and not new_pages:
            self._json({"ok": False, "error": "Upload both PDFs first."}); return

        n = max(len(old_pages), len(new_pages))
        out = []
        for i in range(n):
            old = old_pages[i] if i < len(old_pages) else None
            new = new_pages[i] if i < len(new_pages) else None
            entry = {"page_num": i + 1}
            if old and new:
                # Always include thumbs — even in text mode the frontend can
                # show them as visual context if it likes; "text" mode just
                # opts out of computing the bounding boxes (the slow part).
                old_thumb, otw, oth = _page_thumb_b64(old["full_image"])
                new_thumb, ntw, nth = _page_thumb_b64(new["full_image"])
                entry.update({
                    "old_thumb": old_thumb, "old_tw": otw, "old_th": oth,
                    "new_thumb": new_thumb, "new_tw": ntw, "new_th": nth,
                })
                # Visual diff (only when requested)
                ratio_boxes = []
                if do_visual:
                    try:
                        boxes = _diff_page_boxes(old["full_image"], new["full_image"])
                    except Exception as e:
                        boxes = []
                        entry["error"] = str(e)
                    ow, ohp = old["w"], old["h"]
                    ratio_boxes = [{
                        "rx": b["x"] / max(1, ow),
                        "ry": b["y"] / max(1, ohp),
                        "rw": b["w"] / max(1, ow),
                        "rh": b["h"] / max(1, ohp),
                    } for b in boxes]
                entry["boxes"] = ratio_boxes
                entry["change_count"] = len(ratio_boxes)

                # Text diff (only when requested)
                if do_text:
                    td = _word_diff_html(old.get("text", ""),
                                         new.get("text", ""))
                    entry["text_diff_left"]  = td["left"]
                    entry["text_diff_right"] = td["right"]
                    entry["text_diff_count"] = td["diffs"]
                    entry["text_identical"]  = td["identical"]
                    entry["old_has_text"]    = bool(old.get("text", "").strip())
                    entry["new_has_text"]    = bool(new.get("text", "").strip())

                # Status reflects whatever was checked: a page is "diff" if
                # ANY active mode found a change, otherwise "match".
                visually_changed  = do_visual and entry["change_count"] > 0
                textually_changed = do_text   and not entry.get("text_identical", True)
                entry["status"]   = "diff" if (visually_changed or textually_changed) else "match"
            elif new and not old:
                new_thumb, ntw, nth = _page_thumb_b64(new["full_image"])
                entry.update({"status": "added",
                              "new_thumb": new_thumb, "new_tw": ntw, "new_th": nth,
                              "boxes": [], "change_count": 0})
            else:                       # old and not new
                old_thumb, otw, oth = _page_thumb_b64(old["full_image"])
                entry.update({"status": "removed",
                              "old_thumb": old_thumb, "old_tw": otw, "old_th": oth,
                              "boxes": [], "change_count": 0})
            out.append(entry)
        self._json({
            "ok": True,
            "mode":       mode,
            "old_name":   sess.cpfp_state["old"]["name"],
            "new_name":   sess.cpfp_state["new"]["name"],
            "old_count":  len(old_pages),
            "new_count":  len(new_pages),
            "pages":      out,
        })

    def _cpfp_clear(self):
        """Reset both slots so the user can start over."""
        self.sess.cpfp_state["old"] = {"name": None, "pages": []}
        self.sess.cpfp_state["new"] = {"name": None, "pages": []}
        self._json({"ok": True})

    def _cpfp_claude(self):
        """Layer Claude analysis on top of the deterministic CPFP diff.

        Builds a prompt summarizing every page that has changes (text +
        visual region count) and shells out to the locally installed
        `claude` CLI (Claude Code) — so we use the user's existing Claude
        login (Pro/Max plan or org-issued API key) without managing keys
        in this app. Returns structured per-change annotations + an
        overall summary, so the UI can render inline severity badges
        next to each red box / text-diff mark.
        """
        sess = self.sess
        old_pages = sess.cpfp_state["old"]["pages"]
        new_pages = sess.cpfp_state["new"]["pages"]
        if not old_pages or not new_pages:
            self._json({"ok": False, "error": "Upload both PDFs first."}); return

        # Walk pages in strict-by-index order, collect changes for the prompt.
        # Skip pages where neither text nor visual diff finds anything.
        pages_data = []
        n = max(len(old_pages), len(new_pages))
        for i in range(n):
            old = old_pages[i] if i < len(old_pages) else None
            new = new_pages[i] if i < len(new_pages) else None
            if old and new:
                td = _word_diff_html(old.get("text", ""), new.get("text", ""))
                try:
                    boxes = _diff_page_boxes(old["full_image"], new["full_image"])
                except Exception:
                    boxes = []
                if td["identical"] and not boxes:
                    continue
                pages_data.append({
                    "page_num": i + 1,
                    "old_text": (old.get("text", "") or "")[:3000],
                    "new_text": (new.get("text", "") or "")[:3000],
                    "visual_change_count": len(boxes),
                    "text_change_count":   td["diffs"],
                })
            elif new and not old:
                pages_data.append({
                    "page_num": i + 1, "added": True,
                    "new_text": (new.get("text", "") or "")[:1500],
                })
            elif old and not new:
                pages_data.append({
                    "page_num": i + 1, "removed": True,
                    "old_text": (old.get("text", "") or "")[:1500],
                })

        if not pages_data:
            self._json({"ok": True,
                        "overall_summary": "Old and new are identical — no changes to analyze.",
                        "pages": []})
            return

        prompt = _build_cpfp_claude_prompt(
            sess.cpfp_state["old"]["name"] or "old.pdf",
            sess.cpfp_state["new"]["name"] or "new.pdf",
            pages_data)

        # Resolve the saved Anthropic API key (set by /cpfp/connect).
        api_key = _cpfp_get_api_key()
        if not api_key:
            self._json({"ok": False, "error":
                "Click '🔌 Connect Claude' first to add your Anthropic API key."}); return

        # Talk to api.anthropic.com directly — no CLI dependency.
        try:
            raw = _cpfp_anthropic_call(api_key, prompt)
        except urllib.error.HTTPError as e:
            err_body = e.read().decode("utf-8", errors="replace")[:400]
            self._json({"ok": False, "error":
                f"Anthropic API error {e.code}: {err_body}"}); return
        except Exception as e:
            self._json({"ok": False, "error":
                f"Couldn't reach the Anthropic API: {e}"}); return

        analysis = _extract_cpfp_json_payload(raw)
        if analysis is None:
            # Fall back to surfacing Claude's raw text as the summary so the
            # user still sees something useful even when JSON parsing fails.
            self._json({"ok": True, "overall_summary": raw.strip(),
                        "pages": [], "raw": True})
            return
        self._json({"ok": True, **analysis})

    def _cpfp_connect(self):
        """Validate an Anthropic API key by making a tiny test call, then
        persist it to ~/.banner_app/config.json (0600 perms). The key is
        used by /cpfp/claude for subsequent analysis runs.
        """
        body = self._body() or {}
        key = (body.get("api_key") or "").strip()
        if not key:
            self._json({"ok": False, "error": "Paste an API key first."}); return
        if not key.startswith("sk-ant-"):
            self._json({"ok": False, "error":
                "That doesn't look like an Anthropic API key. "
                "Keys start with 'sk-ant-'. Get one at console.anthropic.com."}); return
        # Validate with a 5-token round-trip — fast, cheap, fails clearly on
        # bad keys / no network / wrong permissions.
        try:
            _cpfp_anthropic_call(key, "Reply with just OK.", max_tokens=5, timeout=20)
        except urllib.error.HTTPError as e:
            err_body = e.read().decode("utf-8", errors="replace")[:200]
            self._json({"ok": False, "error":
                f"Anthropic rejected the key ({e.code}): {err_body}"}); return
        except Exception as e:
            self._json({"ok": False, "error":
                f"Couldn't reach Anthropic: {e}"}); return
        cfg = _cpfp_load_config()
        cfg["anthropic_api_key"] = key
        _cpfp_save_config(cfg)
        self._json({"ok": True, "model": _CPFP_DEFAULT_MODEL})

    def _cpfp_disconnect(self):
        """Remove the saved API key from the config file."""
        cfg = _cpfp_load_config()
        cfg.pop("anthropic_api_key", None)
        _cpfp_save_config(cfg)
        self._json({"ok": True})

    def _cpfp_connection_status(self):
        """Return whether a key is on file, plus the model that will be used
        for analysis. Frontend calls this on load to decide which buttons to
        enable/disable.
        """
        key = _cpfp_get_api_key()
        # Don't return the key itself — only whether one exists, plus a
        # masked preview so the user can confirm it's the right one.
        masked = (key[:7] + "…" + key[-4:]) if key and len(key) > 12 else None
        self._json({"ok": True, "connected": bool(key),
                    "model": _CPFP_DEFAULT_MODEL,
                    "key_preview": masked})

    def _serve_cpfp_page(self):
        """Stream a full-resolution rendered page so the user can zoom in.
        URL: /cpfp/page/<slot>/<page_num>
        """
        sess = self.sess
        bare = self.path.split("?", 1)[0]
        parts = bare.split("/")
        # ['', 'cpfp', 'page', '<slot>', '<n>']
        if len(parts) < 5:
            self._send(400, "text/plain", b"bad path"); return
        slot, n_s = parts[3], parts[4]
        try:
            n = int(n_s)
        except ValueError:
            self._send(400, "text/plain", b"bad page number"); return
        if slot not in sess.cpfp_state:
            self._send(404, "text/plain", b"unknown slot"); return
        pages = sess.cpfp_state[slot]["pages"]
        if n < 1 or n > len(pages):
            self._send(404, "text/plain", b"page out of range"); return
        import io
        buf = io.BytesIO()
        pages[n - 1]["full_image"].save(buf, format="PNG", optimize=True)
        self._send(200, "image/png", buf.getvalue())

    def _storyboard_upload_existing(self):
        """Read a previously-uploaded storyboard PDF (via /upload/storyboard-existing)
        and process it. Body: {"filename": "..."}.
        """
        sess = self.sess
        body = self._body()
        fname = (body.get("filename") or "").strip()
        if not fname:
            self._json({"ok": False, "error": "No file uploaded."}); return
        path = sess.work_dir / "storyboard-existing" / fname
        if not path.exists():
            self._json({"ok": False, "error": f"Upload not found: {fname}"}); return

        try:
            page_imgs = _extract_pdf_pages_to_pngs(path)
        except Exception as e:
            self._json({"ok": False, "error": f"Couldn't read PDF: {e}"}); return

        sess.existing_storyboard_pages = []
        out_pages = []
        for idx, pil in enumerate(page_imgs, start=1):
            regions = _detect_banner_regions(pil)
            sess.existing_storyboard_pages.append({
                "page_num": idx, "full_image": pil, "regions": regions,
            })
            thumb_url, tw, th = _page_thumb_b64(pil)
            out_pages.append({
                "page_num": idx,
                "thumb": thumb_url,
                "thumb_w": tw, "thumb_h": th,
                "page_w": pil.size[0],
                "page_h": pil.size[1],
                "regions": regions,
                "suggested_banner_size": _suggest_banner_size(regions),
            })

        sess.existing_storyboard_name = path.stem
        self._json({"ok": True, "name": path.name, "pages": out_pages})

    def _storyboard_replace_preview(self):
        """Composite new banners onto each page (per the user's mappings) and
        return updated thumbs."""
        sess = self.sess
        body = self._body()
        mappings = body.get("mappings", [])
        if not sess.existing_storyboard_pages:
            self._json({"ok": False, "error": "No storyboard PDF uploaded yet."}); return
        try:
            updated = _route_replace_mappings(sess, mappings, return_thumbs=True)
        except Exception as e:
            self._json({"ok": False, "error": f"Preview failed: {e}"}); return
        self._json({"ok": True, "pages": updated})

    def _storyboard_replace_export(self):
        """Generate the final updated PDF using full-resolution composited pages."""
        sess = self.sess
        body = self._body()
        mappings = body.get("mappings", [])
        if not sess.existing_storyboard_pages:
            self._json({"ok": False, "error": "No storyboard PDF uploaded yet."}); return
        try:
            full_pages = _route_replace_mappings(sess, mappings, return_thumbs=False)
        except Exception as e:
            self._json({"ok": False, "error": f"Export failed: {e}"}); return

        # Save as multi-page PDF inside this session's work_dir.
        out_dir = sess.last_bundle if sess.last_bundle else sess.work_dir
        out_dir.mkdir(parents=True, exist_ok=True)
        name = (sess.existing_storyboard_name or "storyboard") + "_updated.pdf"
        out  = out_dir / name
        try:
            full_pages[0].save(str(out), "PDF",
                               save_all=True, append_images=full_pages[1:])
        except Exception as e:
            self._json({"ok": False, "error": f"Couldn't write PDF: {e}"}); return
        sess.replaced_storyboard_path = str(out)
        self._json({"ok": True, "pdf": str(out), "url": "/storyboard/replace-download"})

    def _storyboard_snap_region(self):
        """Tighten a region's bounds to the non-white content within it.

        Body: {page_num, region: {x,y,w,h}, [pad: int = 0]}
        Returns: {ok, region: {x,y,w,h}}
        """
        sess = self.sess
        body = self._body()
        page_num = body.get("page_num")
        region   = body.get("region") or {}
        pad      = int(body.get("pad", 0))
        if not sess.existing_storyboard_pages:
            self._json({"ok": False, "error": "No storyboard PDF uploaded yet."}); return
        srv = next((e for e in sess.existing_storyboard_pages
                    if e["page_num"] == page_num), None)
        if not srv:
            self._json({"ok": False, "error": f"Page {page_num} not loaded."}); return

        snapped = _snap_region_to_content(srv["full_image"], region, pad=pad)
        if not snapped:
            self._json({"ok": False, "error": "No non-white content in that region."}); return
        self._json({"ok": True, "region": snapped})

    # ── Session download (zip artifacts the capture pipeline produced) ────────
    def _serve_session_download(self):
        """Serve a file from sess.work_dir by name. URL: /download/<filename>.

        Sandboxed: only files directly under sess.work_dir, no subpath
        traversal. Sends Content-Disposition so the browser saves it.
        """
        sess = self.sess
        bare = self.path.split("?", 1)[0]
        # The frontend uses encodeURIComponent() to build the URL — decode
        # back to the actual filename before looking on disk. Without this,
        # any zip whose name contains characters that need encoding (e.g.
        # %3F for ?, %20 for space) would 404.
        name = urllib.parse.unquote(bare[len("/download/"):])
        # Strip any path separators — only allow direct children of work_dir.
        if not name or "/" in name or "\\" in name or name in (".", ".."):
            self._send(400, "text/plain", b"bad name"); return
        try:
            target = (sess.work_dir / name).resolve()
            if sess.work_dir.resolve() != target.parent:
                self._send(403, "text/plain", b"out of scope"); return
        except Exception:
            self._send(404, "text/plain", b"not found"); return
        if not target.is_file():
            self._send(404, "text/plain", b"not found"); return
        data = target.read_bytes()
        mime, _ = mimetypes.guess_type(str(target))
        self.send_response(200)
        self.send_header("Content-Type",   mime or "application/octet-stream")
        self.send_header("Content-Length", str(len(data)))
        self.send_header("Content-Disposition",
                         f'attachment; filename="{target.name}"')
        for c in getattr(self, "_pending_cookies", []):
            self.send_header("Set-Cookie", c)
        self.end_headers()
        self.wfile.write(data)

    # ── SSE ────────────────────────────────────────────────────────────────────
    def _sse(self):
        q = self.sess.sse_queue
        try:
            self.send_response(200)
            self.send_header("Content-Type",  "text/event-stream")
            self.send_header("Cache-Control", "no-cache")
            self.send_header("Connection",    "keep-alive")
            for c in getattr(self, "_pending_cookies", []):
                self.send_header("Set-Cookie", c)
            self.end_headers()
            while True:
                try:
                    msg = q.get(timeout=25)
                    self.wfile.write(f"data: {json.dumps(msg)}\n\n".encode())
                    self.wfile.flush()
                    if msg.get("type") in ("done","error"): break
                except queue.Empty:
                    self.wfile.write(b": ping\n\n"); self.wfile.flush()
        except (BrokenPipeError, ConnectionResetError):
            pass

    # ── Helpers ────────────────────────────────────────────────────────────────
    def _body(self):
        n = int(self.headers.get("Content-Length", 0))
        return json.loads(self.rfile.read(n)) if n else {}

    def _send(self, code, ct, body):
        self.send_response(code)
        self.send_header("Content-Type",   ct)
        self.send_header("Content-Length", str(len(body)))
        for c in getattr(self, "_pending_cookies", []):
            self.send_header("Set-Cookie", c)
        self.end_headers()
        self.wfile.write(body)

    def _html(self, s):  self._send(200, "text/html; charset=utf-8", s.encode())
    def _json(self, d):  self._send(200, "application/json", json.dumps(d).encode())


# ═══════════════════════════════════════════════════════════════════════════════

def _session_janitor():
    """Background thread: every 5 min, evict idle sessions and wipe their
    work_dir. With 5-15 concurrent users at ~50 MB scratch each, this
    matters — without GC, the Hugging Face Space's disk fills up.
    """
    IDLE_LIMIT = 60 * 60       # 1 hour
    SLEEP      = 5 * 60        # 5 min between sweeps
    while True:
        try:
            time.sleep(SLEEP)
            now = time.monotonic()
            stale = []
            with SESSIONS_LOCK:
                for sid, s in list(SESSIONS.items()):
                    if now - s.last_seen > IDLE_LIMIT:
                        stale.append(s)
                        del SESSIONS[sid]
            for s in stale:
                try:
                    shutil.rmtree(s.work_dir, ignore_errors=True)
                except Exception:
                    pass
        except Exception:
            # Never let the janitor crash — it'd quietly leak forever.
            pass


def main():
    # Hosted: bind to 0.0.0.0 and use the standard PORT env var.
    # Hugging Face Spaces sets PORT=7860; locally the user can override it.
    port = int(os.environ.get("PORT", "7860"))
    server = ThreadingHTTPServer(("0.0.0.0", port), Handler)
    print(f"Banner Tool listening on http://0.0.0.0:{port}",
          flush=True)

    # Background session janitor for hosted multi-user lifetimes.
    threading.Thread(target=_session_janitor, daemon=True).start()

    try:
        server.serve_forever()
    except KeyboardInterrupt:
        server.shutdown()

if __name__ == "__main__":
    main()
